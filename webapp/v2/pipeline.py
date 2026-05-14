"""
V2 — Pipeline Orchestrator (Fase 8).

Cuce insieme le Fasi 1-7 in un'unica pipeline end-to-end:
    ZIP → estrazione → triage → classifier → cache prompt → OCR → analyze
        → parse YAML → builder docx incrementale → merge → output

Caratteristiche:
- API uniforme: process_zip_v2(zip_bytes, api_key, ...) -> dict
- SSE typed events emessi attraverso ogni fase (Fase 6)
- Token meter integrato (Fase 7.5)
- Dry-run mode: process_zip_v2(..., dry_run=True) → nessuna chiamata Gemini,
  pipeline produce comunque docx valido per smoke E2E
- Cleanup automatico delle risorse (sezioni partial, OCR manifest, file
  uploaded a Gemini Files API)
- Mai eccezione propagata: errori finiscono in event SSE error + return dict
  con success=False
"""
from __future__ import annotations

import os
import threading
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


# ──────────────────────────────────────────────────────────────────────────────
# Stub Fase 0 — mantenuto per compat con /api/v2/health
# ──────────────────────────────────────────────────────────────────────────────

def process_v2_stub() -> Dict[str, Any]:
    """Stub originale (Fase 0). Ora la pipeline reale è process_zip_v2."""
    return {
        "status": "v2_stub_alive",
        "version": "0.8.0-alpha",
        "phase": "8_orchestrator",
        "timestamp": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "message": "Pipeline V2 orchestrator attivo. Endpoint reale: /api/v2/report/process",
    }


# ──────────────────────────────────────────────────────────────────────────────
# Auto-tuning leve V2 (Leva 2C compact + Leva 4 model mix)
# ──────────────────────────────────────────────────────────────────────────────
#
# Le leve compact_mode (Fase 2C) e model_mix (Leva 4) ottimizzano i costi
# consolidando i documenti AGGREGABLE in batch dedicati con tier MINIMO +
# flash-lite. Sono efficaci solo quando il volume aggregable è alto.
#
# Su pratiche piccole (es. ZIP ISO 27001 con poche decine di file dove ogni
# documento è una procedura SGSI individuale), il consolidamento penalizza la
# qualità dell'output (vedi confronto SIRIH 28 MB: V1 93 schede vs V2 con leve
# attive solo 30 schede).
#
# La logica auto-tuning:
# - Se classified contiene >= MIN_AGGREGABLE_FOR_COMPACT file con audit_role
#   AGGREGABLE → leve abilitate (comportamento ottimizzato)
# - Altrimenti → leve disabilitate (comportamento V1-like, qualità preservata)
#
# Override manuale via env var:
#   V2_LEVA2_AGGREGABLE_COMPACT={true|false|auto}  (default "auto")
#   V2_LEVA4_MODEL_MIX={true|false|auto}            (default "auto")

# Soglia minima di file AGGREGABLE per attivazione auto delle leve compact/lite
DEFAULT_MIN_AGGREGABLE_FOR_COMPACT = 50


def _resolve_leva_flag(env_name: str, auto_default: bool) -> Tuple[bool, str]:
    """
    Risolve flag leva tri-stato: 'true' / 'false' / 'auto' (default).

    Args:
        env_name: nome variabile ambiente (es. "V2_LEVA2_AGGREGABLE_COMPACT")
        auto_default: valore di default quando il flag è 'auto' o non settato

    Returns:
        (enabled, source) dove source ∈ {"manual_on", "manual_off", "auto_on",
        "auto_off"}. Utile per logging.
    """
    raw = os.environ.get(env_name, "auto").strip().lower()
    if raw == "true":
        return (True, "manual_on")
    if raw == "false":
        return (False, "manual_off")
    # "auto" o valore non riconosciuto → applica auto_default
    return (auto_default, "auto_on" if auto_default else "auto_off")


def _count_aggregable_documents(
    documents: List[Dict[str, Any]],
    role_by_filename: Dict[str, str],
) -> int:
    """
    Conta i documenti effettivamente eligibili per il compact_mode.

    Un documento è "aggregable" se:
    - è presente nei `documents` (post triage + OCR + dedup safety net)
    - il suo audit_role mappato è "AGGREGABLE"

    Documenti con role None, CORE, SUPPORT, NOISE non contano.
    """
    n = 0
    for d in documents:
        fname = d.get("filename")
        if fname and role_by_filename.get(fname) == "AGGREGABLE":
            n += 1
    return n


def _build_role_by_filename(
    classified: List[Any],
    skipped_filenames: set,
) -> Dict[str, str]:
    """
    Costruisce mappa filename → audit_role (stringa) dai classified, escludendo
    i file già marcati come skipped dalla safety net (Leva 2 Fase B).

    Quando audit_role è None nel ClassifiedFile, viene derivato dal default
    deterministico (CLASS_TO_DEFAULT_ROLE in schemas/classification.py).
    """
    out: Dict[str, str] = {}
    if not classified:
        return out
    from v2.schemas.classification import (
        DocumentClass,
        default_audit_role_for,
    )
    for cf in classified:
        if cf.filename in skipped_filenames:
            continue
        role = cf.audit_role
        if role is None:
            try:
                classe_str = cf.classe if isinstance(cf.classe, str) else cf.classe.value
                role = default_audit_role_for(DocumentClass(classe_str)).value
            except Exception:
                role = "SUPPORT"
        role_str = role if isinstance(role, str) else role.value
        out[cf.filename] = role_str
    return out


# ──────────────────────────────────────────────────────────────────────────────
# Smart batching (replica V1 First Fit Decreasing)
# ──────────────────────────────────────────────────────────────────────────────

def _create_smart_batches(
    documents: List[Dict[str, Any]],
    max_files: int = 4,
    max_chars: int = 50_000,
) -> List[List[Dict[str, Any]]]:
    """
    First Fit Decreasing per char count, replica V1 senza importare da V1.
    Documents con `content` (o `extracted_text`) come testo da contare.
    """
    def _content_len(d: Dict[str, Any]) -> int:
        return len(d.get("content") or d.get("extracted_text") or "")

    sorted_docs = sorted(documents, key=_content_len, reverse=True)
    batches: List[List[Dict[str, Any]]] = []
    batch_chars: List[int] = []

    for doc in sorted_docs:
        cl = _content_len(doc)
        placed = False
        for i, batch in enumerate(batches):
            if len(batch) < max_files and (batch_chars[i] + cl) <= max_chars:
                batch.append(doc)
                batch_chars[i] += cl
                placed = True
                break
        if not placed:
            batches.append([doc])
            batch_chars.append(cl)
    return batches


# ──────────────────────────────────────────────────────────────────────────────
# Dry-run mock data
# ──────────────────────────────────────────────────────────────────────────────

_DRY_RUN_YAML = """
meta:
  azienda:
    nome: "DRY RUN COMPANY SRL"
    piva: "00000000000"
    sede: "Via Test 1"
  audit:
    data_estrazione: "01/05/2026"
  indice:
    - {n: 1, tipo: "Visura Camerale", titolo: "Visura demo", categoria: "08 Legale"}

sezioni:
  - id: "08"
    nome: "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA"
    documenti:
      - tipo: "Visura Camerale"
        titolo: "Visura demo"
        categoria: "08 · LEGALE"
        riferimento: "DRY-RUN"
        data_doc: "01/01/2025"
        emesso_da: "DRY-RUN"
        soggetto: "DRY RUN COMPANY SRL"
"""


# ──────────────────────────────────────────────────────────────────────────────
# Leva 2 — dry-run audit_role persistence
# ──────────────────────────────────────────────────────────────────────────────

def _persist_audit_role_dryrun(
    session_id: str,
    classified: List[Any],
    classify_metrics: Dict[str, Any],
) -> None:
    """
    Persiste in `temp/audit_role_dryrun/{session_id}.json` la distribuzione
    audit_role del classifier + l'elenco dei candidati NOISE con i loro
    metadati. Non altera alcun comportamento del pipeline (Fase A è solo
    osservazione).
    """
    import json as _json

    base_dir = Path(__file__).resolve().parent.parent.parent / "temp" / "audit_role_dryrun"
    base_dir.mkdir(parents=True, exist_ok=True)

    candidates_noise: List[Dict[str, Any]] = []
    candidates_aggregable: List[Dict[str, Any]] = []
    candidates_support: List[Dict[str, Any]] = []
    role_with_low_arc: List[Dict[str, Any]] = []
    for cf in classified:
        try:
            classe_str = cf.classe if isinstance(cf.classe, str) else cf.classe.value
            role = cf.audit_role
            if role is None:
                from v2.schemas.classification import (
                    DocumentClass,
                    default_audit_role_for,
                )
                role = default_audit_role_for(DocumentClass(classe_str)).value
            role_str = role if isinstance(role, str) else role.value
            arc = cf.audit_role_confidence
            entry = {
                "filename": cf.filename,
                "classe": classe_str,
                "audit_role": role_str,
                "audit_role_confidence": arc,
                "confidence": cf.confidence,
            }
            if role_str == "NOISE":
                candidates_noise.append(entry)
                if arc is None or arc < 0.90:
                    role_with_low_arc.append(entry)
            elif role_str == "AGGREGABLE":
                candidates_aggregable.append(entry)
            elif role_str == "SUPPORT":
                candidates_support.append(entry)
        except Exception:
            continue

    payload = {
        "session_id": session_id,
        "summary": classify_metrics,
        "noise_candidates": candidates_noise,
        "aggregable_candidates": candidates_aggregable,
        "support_candidates": candidates_support,
        "noise_with_low_audit_role_confidence": role_with_low_arc,
    }
    out = base_dir / f"{session_id}.json"
    tmp = out.with_suffix(".json.tmp")
    tmp.write_text(_json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    os.replace(tmp, out)


# ──────────────────────────────────────────────────────────────────────────────
# Pipeline orchestrator
# ──────────────────────────────────────────────────────────────────────────────

def process_zip_v2(
    zip_bytes: bytes,
    session_id: str,
    api_key: Optional[str] = None,
    emitter=None,
    dry_run: bool = False,
    output_dir: Optional[Path] = None,
    _client=None,  # injectable per test
    provider: Optional[str] = None,
    zip_filename: str = "",
    output_mode: Optional[str] = None,  # "schematic" | "narrative" | None (default schematic)
) -> Dict[str, Any]:
    """
    Pipeline V2 completa.

    Args:
        zip_bytes: contenuto ZIP da elaborare
        session_id: identificatore univoco run (alphanumerico_-)
        api_key: chiave Gemini. Se None, legge da env GEMINI_API_KEY.
                 In dry_run è ignorata.
        emitter: SSEEmitter istanziato dal caller (None per uso programmatico)
        dry_run: se True, NESSUNA chiamata Gemini reale; usa mock pre-canned
        output_dir: directory per il docx finale (default temp/v2_output/)
        _client: genai.Client injectable (per test); altrimenti costruito da api_key
        provider: chiave provider LLM ("gemini-2.5-flash" default,
                  "gpt-4.1-mini-azure" alternativa). Se None legge da env
                  V2_PROVIDER, fallback a "gemini-2.5-flash".

    Returns:
        Dict con success/error e metadati (output_path, company_name, stats).
        Mai eccezione propagata.
    """
    from v2 import token_meter
    from v2.docx_merger import merge_session_sections
    from v2.incremental_docx_builder import (
        build_all_sections, builder_summary, cleanup_session_sections,
    )
    from v2.schemas.events import (
        DoneEvent, ErrorKind, PipelinePhase,
    )
    from v2.yaml_parser import (
        extract_company_name,
        get_last_parse_failures,
        parse_aggregated_yaml,
    )
    from v2.zip_extractor import (
        cleanup_extraction, extract_summary, extract_zip_bytes,
    )

    pipeline_start = time.monotonic()

    # Tracker dei file non elaborati lungo le fasi (per la sezione finale
    # "DOCUMENTI NON ELABORATI" del docx). Ogni voce: {filename, reason, phase}.
    unprocessed_files: List[Dict[str, Any]] = []

    def _track_unprocessed(filename: str, reason: str, phase: str) -> None:
        if filename:
            unprocessed_files.append({
                "filename": str(filename),
                "reason": str(reason),
                "phase": str(phase),
            })

    # Default output dir
    if output_dir is None:
        webapp_dir = Path(__file__).resolve().parent.parent
        output_dir = webapp_dir.parent / "temp" / "v2_output"
    output_dir.mkdir(parents=True, exist_ok=True)

    # No-op emitter se non fornito (rende la chiamata sicura senza SSE)
    class _NoopEmitter:
        def __getattr__(self, name):
            return lambda *a, **kw: True
        session_id = ""

    if emitter is None:
        emitter = _NoopEmitter()
        emitter.session_id = session_id  # type: ignore

    # ── Validazione iniziale ─────────────────────────────────────────────
    if not isinstance(zip_bytes, (bytes, bytearray)) or not zip_bytes:
        emitter.emit_error(ErrorKind.UNKNOWN, "zip_bytes vuoto o non bytes")
        return {"success": False, "error": "invalid_zip_bytes"}

    # ── FASE: ingestion (estrazione ZIP) ─────────────────────────────────
    extract_dir = None
    files: List[Dict[str, Any]] = []
    try:
        emitter.emit_phase_start(PipelinePhase.INGESTION)

        # Callback per emettere il nome di ogni file estratto in tempo reale.
        # pct sale da 0.05 a 0.90 con una curva logaritmica (i primi file
        # avanzano veloce, poi si satura) così la barra si muove fin dal primo file.
        import math as _math
        def _zip_on_file(fname, count):
            try:
                # log(count+1)/log(201) ≈ 0→1 per count 0→200; clampiamo a 0.90
                estimated_pct = min(0.90, _math.log(count + 1) / _math.log(201))
                emitter.emit_phase_tick(
                    PipelinePhase.INGESTION,
                    pct=estimated_pct,
                    detail={"current_file": fname[:60], "extracted": count},
                )
            except Exception:
                pass

        files, extract_dir = extract_zip_bytes(zip_bytes, session_id, on_file=_zip_on_file)
        ext_summary = extract_summary(files)
        emitter.emit_phase_end(
            PipelinePhase.INGESTION,
            duration_seconds=round(time.monotonic() - pipeline_start, 2),
            metrics=ext_summary,
        )
        if not files:
            emitter.emit_error(ErrorKind.UNKNOWN, "ZIP vuoto o nessun file estratto")
            return {"success": False, "error": "empty_zip"}
    except ValueError as e:
        emitter.emit_error(ErrorKind.UNKNOWN, f"estrazione_zip_fallita: {e}")
        return {"success": False, "error": f"extraction_failed: {e}"}
    except Exception as e:
        emitter.emit_error(ErrorKind.UNKNOWN, f"errore_imprevisto: {e}")
        return {"success": False, "error": f"unexpected_extraction: {e}"}

    # ── Client Gemini (skip in dry_run) ─────────────────────────────────
    client = _client
    if not dry_run and client is None:
        try:
            from v2.genai_factory_v2 import create_genai_client_v2
            client = create_genai_client_v2(api_key=api_key)
        except Exception as e:
            emitter.emit_error(ErrorKind.API_AUTH, f"client_creation: {e}")
            cleanup_extraction(extract_dir)
            return {"success": False, "error": f"no_client: {e}"}

    # ── FASE: triage (Fase 1) ────────────────────────────────────────────
    from v2.file_triage import triage_files, triage_summary
    phase_start = time.monotonic()
    emitter.emit_phase_start(PipelinePhase.TRIAGE, total_items=len(files))
    triaged = triage_files(files)
    triage_sum = triage_summary(triaged)
    emitter.emit_phase_end(
        PipelinePhase.TRIAGE,
        duration_seconds=round(time.monotonic() - phase_start, 2),
        metrics=triage_sum,
    )

    # ── FASE: estrazione testo non-PDF (Limitazione 1) ──────────────────
    # Per file docx/xlsx/txt/etc., estraggo testo qui prima del classifier.
    # Le immagini (jpg/png/heic/...) hanno method="deferred_to_ocr" e vengono
    # promosse al bucket needs_ocr per il processing via Gemini Vision Files
    # API (che gia' supporta image/jpeg, image/png, image/heic out-of-box).
    from v2.text_handlers import extract_text_for_category
    non_pdf_files = triaged.get("non_pdf", [])
    non_pdf_with_text: List[Dict[str, Any]] = []
    images_for_ocr: List[Dict[str, Any]] = []
    for f in non_pdf_files:
        text, method = extract_text_for_category(f)
        if text and len(text.strip()) >= 10:
            f["extracted_text"] = text
            f["extraction_method"] = method
            non_pdf_with_text.append(f)
        elif method == "deferred_to_ocr":
            # Immagini: non hanno text-layer, ma Vision OCR le legge bene.
            # Vengono accodate ai PDF scansionati per la fase OCR Files API.
            f["extraction_method"] = method
            images_for_ocr.append(f)
        else:
            # Niente testo estraibile → resta "non_pdf" senza contenuto
            f["extraction_method"] = method or "unsupported_format"
            _track_unprocessed(
                f.get("filename"),
                f"non-PDF senza testo estraibile (method={f['extraction_method']})",
                "triage",
            )

    # Promuove le immagini al bucket needs_ocr cosi' la fase OCR esistente
    # le elabora con lo stesso flow dei PDF scansionati.
    if images_for_ocr:
        triaged.setdefault("needs_ocr", []).extend(images_for_ocr)
        print(
            f"[V2 PIPELINE] {len(images_for_ocr)} immagine(/i) promosse a OCR "
            f"Vision (jpg/png/heic supportati da Gemini Files API)"
        )

    # File irrecuperabili dal triage (password, corrotti, troppo grandi)
    for f in triaged.get("unrecoverable", []):
        _track_unprocessed(
            f.get("filename"),
            f"PDF irrecuperabile (method={f.get('extraction_method', 'unknown')})",
            "triage",
        )

    # ── Resolve provider EARLY: serve per decidere se lo step CLASSIFY ──
    # ── e' necessario. Su path Azure il classifier produce solo il file ──
    # ── di osservazione audit_role_dryrun.json e dati per leve 2C/4 che ──
    # ── sono auto-OFF su Azure (vedi piu' sotto). Quindi su Azure salta. ──
    from v2.provider_profiles_v2 import get_profile, resolve_provider_key
    _provider_key_early = resolve_provider_key(provider)
    try:
        _profile_early = get_profile(_provider_key_early)
    except KeyError:
        _profile_early = None  # gestito poi nel blocco analyze

    # Skip classify SOLO se provider e' Azure E nessuno ha attivato la
    # Leva 2 Fase B (skip NOISE) che invece richiede audit_role calcolato.
    _force_classify = (
        os.environ.get("V2_LEVA2_SKIP_NOISE", "false").lower() == "true"
        or os.environ.get("V2_FORCE_CLASSIFY", "false").lower() == "true"
    )
    _skip_classify = (
        _profile_early is not None
        and _profile_early.api_kind == "azure_openai"
        and not _force_classify
    )

    # ── FASE: classify (Fase 2) ──────────────────────────────────────────
    # Tutti i file con o senza testo (anche pre-OCR) vengono classificati
    files_for_classify = (
        triaged.get("native_text", [])
        + triaged.get("needs_ocr", [])
        + non_pdf_with_text
    )

    classified: List = []
    if files_for_classify and not dry_run and not _skip_classify:
        from v2.document_classifier import classify_files_batch, classifier_summary
        phase_start = time.monotonic()
        emitter.emit_phase_start(PipelinePhase.CLASSIFY, total_items=len(files_for_classify))
        try:
            classified = classify_files_batch(
                files_for_classify,
                api_key=api_key,
                _client=client,
                meter_session_id=session_id,
            )
        except Exception as e:
            emitter.emit_error(ErrorKind.UNKNOWN, f"classifier_error: {e}")
            classified = []
        classify_metrics = classifier_summary(classified) if classified else {}
        emitter.emit_phase_end(
            PipelinePhase.CLASSIFY,
            duration_seconds=round(time.monotonic() - phase_start, 2),
            metrics=classify_metrics,
        )

        # Leva 2 — Fase A (dry-run audit_role): persisti la distribuzione e
        # i candidati NOISE per ispezione successiva. NIENTE filtraggio attivo.
        if classified:
            try:
                _persist_audit_role_dryrun(session_id, classified, classify_metrics)
            except Exception as e:
                print(f"[V2 PIPELINE] dry-run audit_role persist fallito: {e}")
    elif _skip_classify and files_for_classify and not dry_run:
        # Su path Azure: emetto comunque lo phase event come "skipped"
        # cosi' il frontend ha uno step pulito invece che mancante.
        emitter.emit_phase_start(PipelinePhase.CLASSIFY, total_items=0)
        emitter.emit_phase_end(
            PipelinePhase.CLASSIFY,
            duration_seconds=0.0,
            metrics={"skipped": True, "reason": "azure_provider_no_leve_active",
                     "n_files_skipped": len(files_for_classify)},
        )
        print(
            f"[V2 PIPELINE] Classify SKIP (provider={_provider_key_early}, "
            f"{len(files_for_classify)} file, saving ~80s). Override con "
            "V2_LEVA2_SKIP_NOISE=true o V2_FORCE_CLASSIFY=true."
        )

    # ── FASE: OCR (Fase 4) ───────────────────────────────────────────────
    # FIX 13: engine OCR selezionabile via env var V2_OCR_ENGINE.
    # - "azure_di" (default raccomandato): Azure Document Intelligence Read API
    #   2-3s/pagina, 1 round-trip, niente Files API, niente bug encoding accenti
    # - "gemini_vision": legacy Gemini Vision Files API (fallback)
    needs_ocr_files = triaged.get("needs_ocr", [])
    ocr_results: List = []
    if needs_ocr_files and not dry_run:
        ocr_engine = os.environ.get("V2_OCR_ENGINE", "azure_di").strip().lower()
        if ocr_engine == "azure_di":
            from v2.azure_document_intelligence import (
                ocr_extract_files, ocr_summary,
            )
            print(f"[V2 PIPELINE] OCR engine: Azure Document Intelligence "
                  f"({len(needs_ocr_files)} file)")
        else:
            from v2.gemini_ocr_v2 import ocr_extract_files, ocr_summary
            print(f"[V2 PIPELINE] OCR engine: Gemini Vision Files API (legacy) "
                  f"({len(needs_ocr_files)} file)")

        phase_start = time.monotonic()
        emitter.emit_phase_start(PipelinePhase.OCR, total_items=len(needs_ocr_files))
        try:
            # Heartbeat timer: emette tick simulati ogni 3s finché OCR non
            # produce completamenti reali — evita freeze del frontend a 0%
            # durante la lunga fase di upload+analisi Azure DI.
            _ocr_completed_ref = [0]
            _ocr_stop_timer = threading.Event()

            def _ocr_heartbeat_timer():
                import time as _time
                _elapsed = 0
                while not _ocr_stop_timer.wait(timeout=3.0):
                    _elapsed += 3
                    if _ocr_completed_ref[0] == 0:
                        # Nessun completamento ancora: simula avanzamento lento
                        # (max 25% stimato, per non "mentire" troppo)
                        est_pct = min(0.25, _elapsed / (len(needs_ocr_files) * 4.0))
                        try:
                            emitter.emit_phase_tick(
                                PipelinePhase.OCR,
                                pct=est_pct,
                                detail={"engine": ocr_engine, "waiting": True},
                            )
                        except Exception:
                            pass

            _hb_thread = threading.Thread(target=_ocr_heartbeat_timer, daemon=True)
            _hb_thread.start()

            # Callback granulare → emit_phase_tick per ogni file OCR completato
            def _ocr_progress(completed, total, current_file):
                _ocr_completed_ref[0] = completed
                try:
                    pct = completed / max(total, 1)
                    emitter.emit_phase_tick(
                        PipelinePhase.OCR,
                        pct=pct,
                        detail={
                            "completed": completed,
                            "total": total,
                            "current_file": str(current_file or "")[:60],
                            "engine": ocr_engine,
                        },
                    )
                except Exception:
                    pass

            ocr_results = ocr_extract_files(
                client,
                needs_ocr_files,
                session_id=session_id,
                cleanup_after=True,
                meter_session_id=session_id,
                on_progress=_ocr_progress,
            )
            _ocr_stop_timer.set()  # ferma heartbeat timer OCR
            # Aggiorna i file con il testo estratto
            ocr_text_by_filename = {
                r.filename: r.text for r in ocr_results if r.success
            }
            for f in needs_ocr_files:
                if f["filename"] in ocr_text_by_filename:
                    f["extracted_text"] = ocr_text_by_filename[f["filename"]]
                else:
                    # OCR fallito: file scansionato senza text-layer e Vision
                    # non ha prodotto testo utile → tracciato come unprocessed
                    _track_unprocessed(
                        f.get("filename"),
                        "OCR Vision fallito o senza testo utile",
                        "ocr",
                    )
        except Exception as e:
            emitter.emit_error(ErrorKind.OCR_FAILED, f"ocr_pipeline: {e}")
        emitter.emit_phase_end(
            PipelinePhase.OCR,
            duration_seconds=round(time.monotonic() - phase_start, 2),
            metrics=ocr_summary(ocr_results) if ocr_results else {},
        )

    # ── Leva 2 Fase B: applica safety net ai classified ────────────────
    # Filtra i file marcati audit_role=NOISE (con 4 layer di safety net) PRIMA
    # di costruire i documents per analyze. Attivata via env var
    # `V2_LEVA2_SKIP_NOISE=true`. Quando spenta, comportamento invariato.
    skipped_filenames: set = set()
    relevance_ledger: List[Dict[str, Any]] = []
    relevance_stats: Dict[str, Any] = {}
    if (
        classified
        and os.environ.get("V2_LEVA2_SKIP_NOISE", "false").lower() == "true"
        and not dry_run
    ):
        from v2.relevance_safetynet import apply_safety_net

        # files_index serve al dedup hash: mappa filename → file_info con
        # extracted_text. Lo costruiamo dai 3 set di file processati.
        files_index = {
            f["filename"]: f
            for f in (triaged.get("native_text", []) + needs_ocr_files + non_pdf_with_text)
        }
        sn_result = apply_safety_net(classified, files_index)
        skipped_filenames = {cf.filename for cf in sn_result["skipped"]}
        relevance_ledger = sn_result["ledger"]
        relevance_stats = sn_result["stats"]
        if skipped_filenames:
            print(
                f"[V2 PIPELINE] Leva 2 Fase B: {len(skipped_filenames)} file "
                f"marcati NOISE e skippati (cap {relevance_stats.get('skipped_pct', 0)}%). "
                f"Whitelist override: {relevance_stats.get('whitelist_overrides_to_core', 0)}, "
                f"dedup: {relevance_stats.get('duplicates_marked_noise', 0)}"
            )
        # Persisti il ledger su disco (il rendering nel docx arriverà in Fase C)
        try:
            import json as _json
            ledger_dir = Path(__file__).resolve().parent.parent.parent / "temp" / "relevance_ledger"
            ledger_dir.mkdir(parents=True, exist_ok=True)
            payload = {
                "session_id": session_id,
                "stats": relevance_stats,
                "ledger": relevance_ledger,
                "details": sn_result.get("details", {}),
            }
            out_path = ledger_dir / f"{session_id}.json"
            tmp = out_path.with_suffix(".json.tmp")
            tmp.write_text(_json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
            os.replace(tmp, out_path)
        except Exception as e:
            print(f"[V2 PIPELINE] relevance_ledger persist fallito: {e}")

    # ── Costruisci documents finali per analyze ─────────────────────────
    # Includiamo tutti i file con extracted_text non vuoto:
    # PDF nativi + OCR PDF + non-PDF (Word, Excel, ecc.)
    # Esclusi: file in skipped_filenames (Leva 2 Fase B).
    documents = []
    for f in (triaged.get("native_text", []) + needs_ocr_files + non_pdf_with_text):
        if f["filename"] in skipped_filenames:
            continue
        text = (f.get("extracted_text") or "").strip()
        if text:
            documents.append({
                "filename": f["filename"],
                "content": text,
            })

    if not documents and not dry_run:
        emitter.emit_error(ErrorKind.UNKNOWN, "Nessun documento con testo estratto")
        cleanup_extraction(extract_dir)
        return {"success": False, "error": "no_extractable_documents"}

    # ── FASE: cache prompt + analyze (Fase 3 + 5) ───────────────────────
    raw_yamls: List[str] = []
    n_fallback_batches = 0
    fallback_batch_idxs: List[int] = []
    provider_key_used = ""
    analyze_phase_metrics: Dict[str, Any] = {}
    output_mode_narrative = False  # settato a True nel path Azure se V2_OUTPUT_MODE=narrative
    narrative_paragraphs_by_idx: Dict[int, List[Dict[str, Any]]] = {}

    if dry_run:
        # Dry-run: usa mock YAML pre-canned
        raw_yamls = [_DRY_RUN_YAML]
    else:
        # ── Resolve provider + build dispatch (eventuale fallback Gemini) ──
        from v2.provider_profiles_v2 import get_profile, resolve_provider_key
        from v2.llm_dispatch import build_dispatch, fallback_disabled

        provider_key = resolve_provider_key(provider)
        provider_key_used = provider_key
        try:
            primary_profile = get_profile(provider_key)
        except KeyError as e:
            emitter.emit_error(ErrorKind.UNKNOWN, f"provider_unknown: {e}")
            cleanup_extraction(extract_dir)
            return {"success": False, "error": f"provider_unknown: {e}"}

        try:
            primary_dispatch = build_dispatch(
                provider_key, api_key=api_key, gemini_client=client,
            )
        except Exception as e:
            emitter.emit_error(ErrorKind.API_AUTH, f"dispatch_build: {e}")
            cleanup_extraction(extract_dir)
            return {"success": False, "error": f"dispatch_build: {e}"}

        # Fallback Gemini disponibile solo se provider primario è Azure
        # E V2_DISABLE_FALLBACK != "1". Costruito SEMPRE eager per evitare
        # latenze sul primo 429 (saturazione concorrente).
        fallback_dispatch = None
        if primary_profile.api_kind == "azure_openai" and not fallback_disabled():
            try:
                fallback_dispatch = build_dispatch(
                    "gemini-2.5-flash", api_key=api_key, gemini_client=client,
                )
                print(
                    f"[V2 PIPELINE] Fallback Gemini ARMATO per "
                    f"{provider_key} (V2_DISABLE_FALLBACK=0)"
                )
            except Exception as e:
                print(f"[V2 PIPELINE] Fallback Gemini non disponibile: {e}")

        # ── PATH GEMINI: cached_id + auto-tuning leve V2 PROD ──────────────
        if primary_profile.api_kind == "gemini_v2_native":
            from v2.cache_manager import get_cached_prompt
            cached_id = get_cached_prompt(client)

            role_by_filename = _build_role_by_filename(classified, skipped_filenames)
            n_aggregable = _count_aggregable_documents(documents, role_by_filename)

            threshold = int(
                os.environ.get(
                    "V2_MIN_AGGREGABLE_THRESHOLD",
                    str(DEFAULT_MIN_AGGREGABLE_FOR_COMPACT),
                )
            )
            auto_should_compact = (
                bool(classified) and n_aggregable >= threshold
            )

            aggregable_compact_enabled, compact_source = _resolve_leva_flag(
                "V2_LEVA2_AGGREGABLE_COMPACT", auto_default=auto_should_compact
            )
            aggregable_compact_enabled = aggregable_compact_enabled and bool(classified)

            model_mix_enabled, model_mix_source = _resolve_leva_flag(
                "V2_LEVA4_MODEL_MIX", auto_default=auto_should_compact
            )
            model_mix_enabled = model_mix_enabled and aggregable_compact_enabled

            print(
                f"[V2 PIPELINE] Auto-tuning leve: aggregable_count={n_aggregable} "
                f"(soglia={threshold}). "
                f"compact_mode={'ON' if aggregable_compact_enabled else 'OFF'} "
                f"({compact_source}), "
                f"model_mix={'ON' if model_mix_enabled else 'OFF'} "
                f"({model_mix_source})"
            )

            documents_standard: List[Dict[str, Any]] = []
            documents_aggregable: List[Dict[str, Any]] = []
            if aggregable_compact_enabled:
                for d in documents:
                    if role_by_filename.get(d["filename"]) == "AGGREGABLE":
                        documents_aggregable.append(d)
                    else:
                        documents_standard.append(d)
            else:
                documents_standard = documents

            batches_standard = _create_smart_batches(
                documents_standard, max_files=4, max_chars=50_000,
            )
            batches_aggregable = (
                _create_smart_batches(documents_aggregable, max_files=8, max_chars=50_000)
                if documents_aggregable
                else []
            )

            from v2.gemini_client_v2 import ANALYZE_MODEL_LITE
            lite_model_for_aggregable = (
                ANALYZE_MODEL_LITE if model_mix_enabled else None
            )

            all_batches: List[Tuple[List[Dict[str, Any]], bool, Optional[str]]] = (
                [(b, False, None) for b in batches_standard]
                + [(b, True, lite_model_for_aggregable) for b in batches_aggregable]
            )

            print(
                f"[V2 PIPELINE] Batch split (Gemini): docs_total={len(documents)}, "
                f"standard={len(documents_standard)}, aggregable={len(documents_aggregable)}, "
                f"batches_standard={len(batches_standard)}, "
                f"batches_aggregable={len(batches_aggregable)}, "
                f"model_aggregable={lite_model_for_aggregable or 'flash'}"
            )
            max_workers = min(primary_profile.max_workers, max(1, len(all_batches)))
            analyze_phase_metrics = {
                "batches_total": len(all_batches),
                "batches_standard": len(batches_standard),
                "batches_aggregable_compact": len(batches_aggregable),
                "model_mix_active": model_mix_enabled,
            }
        else:
            # ── PATH AZURE: niente cache_id, batch profile-driven, no leve ─
            cached_id = None
            azure_batches = _create_smart_batches(
                documents,
                max_files=primary_profile.batch_max_files,
                max_chars=primary_profile.batch_max_chars,
            )
            all_batches = [(b, False, None) for b in azure_batches]
            print(
                f"[V2 PIPELINE] Provider Azure {primary_profile.name}: "
                f"docs_total={len(documents)}, batches={len(all_batches)} "
                f"(max_files={primary_profile.batch_max_files}, "
                f"max_chars={primary_profile.batch_max_chars}), "
                f"max_workers={primary_profile.max_workers}"
            )
            max_workers = min(primary_profile.max_workers, max(1, len(all_batches)))
            analyze_phase_metrics = {
                "batches_total": len(all_batches),
            }

        # ── Modalità output (per-request + env fallback + default schematic) ──
        # schematic  → prosa schematica telegrafica (DEFAULT)
        # narrative  → prosa narrativa discorsiva (modalità storica PROD)
        # Precedenza: arg per-request (output_mode) > env var V2_OUTPUT_MODE > default "schematic"
        _arg_mode = (output_mode or "").strip().lower() if output_mode else ""
        _env_mode = os.environ.get("V2_OUTPUT_MODE", "").strip().lower()
        _effective_mode = _arg_mode or _env_mode or "schematic"
        if _effective_mode not in ("schematic", "narrative"):
            _effective_mode = "schematic"

        output_mode_schematic = (
            _effective_mode == "schematic"
            and primary_profile.api_kind == "azure_openai"
        )
        # output_mode_narrative resta True anche per schematic, perche' il flow
        # downstream e' identico (JSON narrativo): cambia solo il prompt e
        # l'analyze function. Mantiene compat con tutti i check esistenti.
        output_mode_narrative = (
            (_effective_mode == "narrative" and primary_profile.api_kind == "azure_openai")
            or output_mode_schematic
        )
        _output_mode_label = (
            "schematic" if output_mode_schematic
            else ("narrative" if output_mode_narrative else "yaml")
        )
        if output_mode_narrative:
            print(
                f"[V2 PIPELINE] OUTPUT_MODE={_output_mode_label} attivato su {provider_key} "
                f"(arg={_arg_mode!r}, env={_env_mode!r}). "
                "Path YAML/incremental-builder BYPASSATO."
            )

        # ── Loop analyze comune (Gemini + Azure con fallback) ──────────────
        from v2.azure_openai_client_v2 import AzureRateLimitExhausted

        phase_start = time.monotonic()
        emitter.emit_phase_start(PipelinePhase.ANALYZE, total_items=len(all_batches))

        from concurrent.futures import ThreadPoolExecutor, as_completed
        completed = 0
        results_by_idx: Dict[int, str] = {}
        # PATH NARRATIVO: raccoglie paragrafi al posto di stringhe YAML
        fallback_used: List[int] = []

        # Setup path narrativo: carica prompt e calcola para_start per ogni batch
        _narrative_prompt: str = ""
        _batch_para_starts: List[int] = []
        analyze_batch_narrative = None
        analyze_batch_narrative_gemini = None
        if output_mode_narrative:
            if output_mode_schematic:
                # SCHEMATIC: client ora in webapp/v2/ (promosso a PROD default)
                from v2.schematic_client_v2 import (
                    analyze_batch_schematic as analyze_batch_narrative,
                    analyze_batch_schematic_gemini as analyze_batch_narrative_gemini,
                    _load_schematic_prompt as _load_prompt_fn,
                )
                print("[V2 PIPELINE] Schematic client caricato da webapp/v2/")
            else:
                # NARRATIVE: client modalità discorsiva storica
                from v2.narrative_client_v2 import (
                    analyze_batch_narrative,
                    analyze_batch_narrative_gemini,
                    _load_narrative_prompt as _load_prompt_fn,
                )
            _narrative_prompt = _load_prompt_fn()
            # Calcola para_start cumulativo per ogni batch (1-indexed)
            _running = 1
            for _b, _c, _m in all_batches:
                _batch_para_starts.append(_running)
                _running += len(_b)

        def _try_fallback(batch_docs, idx, compact):
            """Tenta il batch su fallback_dispatch. Ritorna (result|None)."""
            if fallback_dispatch is None:
                return None
            try:
                return fallback_dispatch.analyze(
                    client=fallback_dispatch.client,
                    batch_docs=batch_docs,
                    batch_idx=idx,
                    total_docs=len(documents),
                    cached_content_id=None,
                    meter_session_id=session_id,
                    compact_mode=compact,
                    model_override=None,
                )
            except Exception as e:
                print(f"[V2 PIPELINE] Fallback batch {idx} sollevato: {e}")
                return None

        _force_fallback = os.environ.get("V2_FORCE_NARRATIVE_FALLBACK", "").strip() == "1"

        def _analyze_one(idx: int, batch_docs: List[Dict[str, Any]],
                          compact: bool, model_override: Optional[str]):
            # ── PATH NARRATIVO Azure ────────────────────────────────────────
            if output_mode_narrative:
                para_start = _batch_para_starts[idx]
                verb_idx = idx  # ciclo verbi per batch

                # Forza fallback Gemini senza chiamare Azure (solo per test)
                if _force_fallback and fallback_dispatch is not None:
                    print(f"[V2 PIPELINE] FORCE_FALLBACK batch {idx} → Gemini narrativo")
                    try:
                        fb_paras, _ = analyze_batch_narrative_gemini(
                            gemini_client=fallback_dispatch.client,
                            batch_docs=batch_docs,
                            batch_idx=idx,
                            total_docs=len(documents),
                            para_start=para_start,
                            verb_idx=verb_idx,
                            narrative_prompt=_narrative_prompt,
                            meter_session_id=session_id,
                        )
                        return idx, fb_paras, True, True
                    except Exception as fb_e:
                        print(f"[V2 PIPELINE] FORCE_FALLBACK batch {idx} Gemini fallito: {fb_e}")
                        return idx, None, False, True

                try:
                    paras, _fb = analyze_batch_narrative(
                        client=primary_dispatch.client,
                        batch_docs=batch_docs,
                        batch_idx=idx,
                        total_docs=len(documents),
                        para_start=para_start,
                        verb_idx=verb_idx,
                        narrative_prompt=_narrative_prompt,
                        meter_session_id=session_id,
                    )
                    return idx, paras, False, True  # (idx, paras, fallback_used, is_narrative)
                except AzureRateLimitExhausted as e:
                    print(
                        f"[V2 PIPELINE] Narrative batch {idx} Azure 429 esaurito "
                        f"({e.n_retries} retry) → fallback Gemini NARRATIVE"
                    )
                    if fallback_dispatch is None:
                        print(f"[V2 PIPELINE] Fallback non disponibile, batch {idx} perso")
                        return idx, None, False, True
                    # Chiama Gemini con lo stesso prompt narrativo (non YAML)
                    try:
                        fb_paras, _ = analyze_batch_narrative_gemini(
                            gemini_client=fallback_dispatch.client,
                            batch_docs=batch_docs,
                            batch_idx=idx,
                            total_docs=len(documents),
                            para_start=para_start,
                            verb_idx=verb_idx,
                            narrative_prompt=_narrative_prompt,
                            meter_session_id=session_id,
                        )
                        print(
                            f"[V2 PIPELINE] Narrative fallback Gemini batch {idx}: "
                            f"{len(fb_paras)} paragrafi"
                        )
                        return idx, fb_paras, True, True
                    except Exception as fb_e:
                        print(
                            f"[V2 PIPELINE] Narrative fallback Gemini batch {idx} fallito: {fb_e}"
                        )
                        return idx, None, False, True

            # ── PATH YAML (Gemini + Azure YAML) ────────────────────────────
            # Trigger 1: rate limit Azure 429 cumulativi → fallback obbligatorio
            try:
                result = primary_dispatch.analyze(
                    client=primary_dispatch.client,
                    batch_docs=batch_docs,
                    batch_idx=idx,
                    total_docs=len(documents),
                    cached_content_id=cached_id,
                    meter_session_id=session_id,
                    compact_mode=compact,
                    model_override=model_override,
                )
            except AzureRateLimitExhausted as e:
                print(
                    f"[V2 PIPELINE] Batch {idx} Azure 429 esaurito "
                    f"({e.n_retries} retry) → fallback Gemini"
                )
                if fallback_dispatch is None:
                    print(f"[V2 PIPELINE] Fallback non disponibile, batch {idx} perso")
                    return idx, None, False, False
                fb = _try_fallback(batch_docs, idx, compact)
                if fb is None:
                    print(f"[V2 PIPELINE] Fallback batch {idx} fallito: nessun result")
                    return idx, None, False, False
                return idx, fb, True, False

            # Trigger 2: output troncato (Azure 32K cap finish_reason=length).
            # Il batch ha risposta HTTP 200 ma YAML probabilmente rotto a metà.
            # Fallback Gemini sfrutta il 64K cap (più ampio) per recuperarlo.
            # `truncated_output` viene settato SOLO da azure_openai_client_v2;
            # su path Gemini puro è sempre False → no fallback inutile.
            if getattr(result, "truncated_output", False) and fallback_dispatch is not None:
                print(
                    f"[V2 PIPELINE] Batch {idx} Azure output troncato "
                    f"(finish_reason=length, 32K cap) → fallback Gemini"
                )
                fb = _try_fallback(batch_docs, idx, compact)
                if fb is not None and fb.text and not fb.error:
                    return idx, fb, True, False
                # Fallback fallito o senza output utile → conservo il primary
                # parziale (parser potrebbe ancora estrarre qualcosa di utile)
                print(f"[V2 PIPELINE] Fallback truncation batch {idx} non utile, conservo Azure parziale")

            return idx, result, False, False

        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {
                pool.submit(_analyze_one, i, b, c, m): i
                for i, (b, c, m) in enumerate(all_batches)
            }
            for fut in as_completed(futures):
                try:
                    idx, result, used_fb, _is_narr = fut.result()
                    if _is_narr:
                        # result è List[Dict] paragrafi
                        if result:
                            narrative_paragraphs_by_idx[idx] = result
                    else:
                        if result and result.text and not result.error:
                            results_by_idx[idx] = result.text
                    if used_fb:
                        fallback_used.append(idx)
                except Exception as e:
                    emitter.emit_error(ErrorKind.UNKNOWN, f"batch_analyze: {e}")
                completed += 1
                emitter.emit_phase_tick(
                    PipelinePhase.ANALYZE,
                    pct=completed / max(len(all_batches), 1),
                    detail={"completed": completed, "total": len(all_batches)},
                )

        raw_yamls = [results_by_idx[i] for i in sorted(results_by_idx.keys())]
        n_fallback_batches = len(fallback_used)
        fallback_batch_idxs = sorted(fallback_used)

        # Tracking dei file dei batch persi dopo eventuali fallback.
        _resolved_idxs = (
            set(narrative_paragraphs_by_idx.keys())
            if output_mode_narrative
            else set(results_by_idx.keys())
        )
        batches_lost = sorted(
            i for i in range(len(all_batches)) if i not in _resolved_idxs
        )
        for lost_idx in batches_lost:
            batch_docs_lost = all_batches[lost_idx][0]
            for d in batch_docs_lost:
                _track_unprocessed(
                    d.get("filename"),
                    f"batch_{lost_idx} non recuperabile (Azure error + fallback fallito)",
                    "analyze",
                )

        _batches_ok = (
            len(narrative_paragraphs_by_idx) if output_mode_narrative else len(raw_yamls)
        )
        analyze_phase_metrics["batches_ok"] = _batches_ok
        analyze_phase_metrics["batches_lost"] = len(batches_lost)
        analyze_phase_metrics["provider"] = provider_key
        analyze_phase_metrics["n_fallback_batches"] = n_fallback_batches

        emitter.emit_phase_end(
            PipelinePhase.ANALYZE,
            duration_seconds=round(time.monotonic() - phase_start, 2),
            metrics=analyze_phase_metrics,
        )

    # ══════════════════════════════════════════════════════════════════════
    # PATH NARRATIVO: parse JSON → generate_structured_evidence_docx
    # Bypassato: YAML parse, incremental docx builder, merge_session_sections
    # ══════════════════════════════════════════════════════════════════════
    if output_mode_narrative:
        if not narrative_paragraphs_by_idx:
            emitter.emit_error(ErrorKind.UNKNOWN, "Nessun batch narrativo ha prodotto output")
            cleanup_extraction(extract_dir)
            return {"success": False, "error": "no_narrative_output"}

        # Raduna tutti i paragrafi in ordine di batch, poi rinumera 1..N
        all_paragraphs: List[Dict[str, Any]] = []
        for idx in sorted(narrative_paragraphs_by_idx.keys()):
            all_paragraphs.extend(narrative_paragraphs_by_idx[idx])
        for i, p in enumerate(all_paragraphs, start=1):
            p["numero"] = i

        # Conta macroaree distinte
        _macroaree = {p.get("categoria", "ALTRO") for p in all_paragraphs}

        _summary_msg = (
            f"Analisi completata: {len(documents)} documenti analizzati, "
            f"{len(all_paragraphs)} paragrafi generati, "
            f"{len(_macroaree)} macroaree"
        )
        print(f"[V2 PIPELINE] {_summary_msg}")
        emitter.emit_phase_tick(
            PipelinePhase.ANALYZE,
            pct=1.0,
            detail={"msg": _summary_msg, "paragraphs": len(all_paragraphs), "macroaree": len(_macroaree)},
        )

        # Estrae company_name con la stessa logica del sistema legacy
        # (visura → ente_auditato con blacklist → regex forme giuridiche)
        from modules.report_generator import extract_company_name, generate_report_word
        docs_estratti = len(documents)
        docs_vuoti = max(0, len(files) - docs_estratti)
        company_name = extract_company_name(
            documents=documents,
            paragraphs=all_paragraphs,
            zip_filename=zip_filename,
        )

        # Build docx narrativo con generate_report_word (paragrafi liberi, NON tabelle)
        phase_start = time.monotonic()
        emitter.emit_phase_start(PipelinePhase.DOCX_BUILD)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company = "".join(c for c in company_name if c.isalnum() or c == " ")[:30]
        safe_company = safe_company.replace(" ", "_") or "AUDIT"
        final_filename = f"Audit_V2_NAR_{safe_company}_{timestamp}.docx"
        final_path = output_dir / f"{session_id}_final.docx"

        import re as _re_docx
        import traceback as _tb_docx

        def _sanitize_para(p: dict) -> dict:
            """Copia il paragrafo rimuovendo ogni char XML-illegale da tutti i campi stringa."""
            def _s(v):
                if isinstance(v, str):
                    return _re_docx.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", v)
                return v
            return {k: _s(v) for k, v in p.items()}

        stats_for_word = {
            "total_paragraphs": len(all_paragraphs),
            "docs_estratti": docs_estratti,
            "docs_vuoti": docs_vuoti,
        }
        clean_company = _re_docx.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", company_name)
        clean_paragraphs = [_sanitize_para(p) for p in all_paragraphs]

        try:
            docx_bytes = generate_report_word(
                paragraphs=clean_paragraphs,
                stats=stats_for_word,
                company_name=clean_company,
            )
            final_path.write_bytes(docx_bytes)
        except Exception as e:
            print(f"[V2 PIPELINE] narrative_docx_build TRACEBACK:\n{_tb_docx.format_exc()}")
            # Fallback: docx minimale con solo i testi grezzi concatenati
            try:
                from docx import Document as _Doc
                from io import BytesIO as _BytesIO
                _fb_doc = _Doc()
                _fb_doc.add_heading(f"AUDIT - {clean_company} (output parziale)", 0)
                _fb_doc.add_paragraph(
                    f"ATTENZIONE: la generazione normale ha restituito un errore ({e}). "
                    "Di seguito il testo grezzo dei paragrafi analizzati."
                )
                for p in clean_paragraphs:
                    _fb_doc.add_heading(str(p.get("sottotitolo", ""))[:200], 2)
                    _fb_doc.add_paragraph(str(p.get("contenuto", ""))[:5000])
                _buf = _BytesIO()
                _fb_doc.save(_buf)
                final_path.write_bytes(_buf.getvalue())
                emitter.emit_error(ErrorKind.UNKNOWN, f"narrative_docx_build (fallback usato): {e}")
            except Exception as e2:
                print(f"[V2 PIPELINE] narrative_docx_build FALLBACK FAILED:\n{_tb_docx.format_exc()}")
                emitter.emit_error(ErrorKind.UNKNOWN, f"narrative_docx_build: {e}")
                cleanup_extraction(extract_dir)
                return {"success": False, "error": f"narrative_docx_build: {e}"}

        emitter.emit_phase_end(
            PipelinePhase.DOCX_BUILD,
            duration_seconds=round(time.monotonic() - phase_start, 2),
            metrics={"paragraphs": len(all_paragraphs)},
        )

        # Cleanup
        phase_start = time.monotonic()
        emitter.emit_phase_start(PipelinePhase.CLEANUP)
        cleanup_extraction(extract_dir)
        emitter.emit_phase_end(
            PipelinePhase.CLEANUP,
            duration_seconds=round(time.monotonic() - phase_start, 2),
        )

        token_report = token_meter.get_session_report(session_id)
        total_duration = round(time.monotonic() - pipeline_start, 2)
        stats = {
            "total_files_in_zip": len(files),
            "documents_with_text": docs_estratti,
            "duration_seconds": total_duration,
            "company_name": company_name,
            "tokens": {
                "input_total": token_report.get("total_input", 0),
                "cached_total": token_report.get("total_cached", 0),
                "output_total": token_report.get("total_output", 0),
                "cost_eur": token_report.get("total_cost_eur", 0.0),
                "saved_by_caching_eur": token_report.get("saved_by_caching_eur", 0.0),
            },
            "dry_run": dry_run,
            "provider": provider_key_used,
            "n_fallback_batches": n_fallback_batches,
            "fallback_batch_idxs": fallback_batch_idxs,
            "n_unprocessed_files": len(unprocessed_files),
            "unprocessed_files": unprocessed_files,
            "output_mode": "narrative",
            "n_paragraphs": len(all_paragraphs),
        }
        emitter.emit_done(
            output_filename=final_filename,
            company_name=company_name,
            duration_seconds=total_duration,
            stats=stats,
        )
        return {
            "success": True,
            "output_path": str(final_path),
            "output_filename": final_filename,
            "company_name": company_name,
            "stats": stats,
            "merge_used_fallback": False,
            "provider": provider_key_used,
            "n_fallback_batches": n_fallback_batches,
            "fallback_batch_idxs": fallback_batch_idxs,
            "unprocessed_files": unprocessed_files,
        }

    # ══════════════════════════════════════════════════════════════════════
    # PATH YAML (Gemini + Azure YAML): flusso originale invariato
    # ══════════════════════════════════════════════════════════════════════
    if not raw_yamls:
        emitter.emit_error(ErrorKind.UNKNOWN, "Nessun batch ha prodotto YAML utile")
        cleanup_extraction(extract_dir)
        return {"success": False, "error": "no_yaml_output"}

    # ── Parse YAML aggregato ─────────────────────────────────────────────
    full_yaml = "\n\n---\n\n".join(raw_yamls)
    parsed_data = parse_aggregated_yaml(full_yaml)
    company_name = extract_company_name(parsed_data)

    # Leva 3 — segnala batch persi durante il parse YAML (errori non più
    # silenziosi). Ogni batch fallito diventa un evento PARSE_FAILED con
    # batch_id e prima riga.
    parse_failures = get_last_parse_failures()
    for failure in parse_failures:
        try:
            emitter.emit_error(
                ErrorKind.PARSE_FAILED,
                f"yaml_batch_skipped: {failure.get('batch_id')} → "
                f"{failure.get('error')}",
            )
        except Exception:
            pass
    if parse_failures:
        print(
            f"[V2 PIPELINE] {len(parse_failures)} batch YAML saltati "
            f"durante l'aggregazione (vedi eventi PARSE_FAILED)"
        )

    # ── FASE: docx_build (Fase 7) ────────────────────────────────────────
    phase_start = time.monotonic()
    emitter.emit_phase_start(PipelinePhase.DOCX_BUILD)
    docs_estratti = len(documents)
    docs_vuoti = max(0, len(files) - docs_estratti)
    build_results = build_all_sections(
        parsed_data,
        session_id=session_id,
        docs_estratti=docs_estratti,
        docs_vuoti=docs_vuoti,
        unprocessed_files=unprocessed_files,
    )
    build_sum = builder_summary(build_results)
    emitter.emit_phase_end(
        PipelinePhase.DOCX_BUILD,
        duration_seconds=round(time.monotonic() - phase_start, 2),
        metrics=build_sum,
    )

    # ── Merge finale ─────────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_company = "".join(c for c in company_name if c.isalnum() or c == " ")[:30]
    safe_company = safe_company.replace(" ", "_") or "AUDIT"
    final_filename = f"Audit_V2_{safe_company}_{timestamp}.docx"
    final_path = output_dir / f"{session_id}_final.docx"

    merge_result = merge_session_sections(session_id, final_path)
    if not merge_result.success:
        emitter.emit_error(ErrorKind.UNKNOWN, f"merge_failed: {merge_result.error}")
        cleanup_session_sections(session_id)
        cleanup_extraction(extract_dir)
        return {"success": False, "error": f"merge_failed: {merge_result.error}"}

    # ── Cleanup ──────────────────────────────────────────────────────────
    phase_start = time.monotonic()
    emitter.emit_phase_start(PipelinePhase.CLEANUP)
    cleanup_session_sections(session_id)
    cleanup_extraction(extract_dir)
    emitter.emit_phase_end(
        PipelinePhase.CLEANUP,
        duration_seconds=round(time.monotonic() - phase_start, 2),
    )

    # ── Token report finale ──────────────────────────────────────────────
    token_report = token_meter.get_session_report(session_id)

    total_duration = round(time.monotonic() - pipeline_start, 2)
    stats = {
        "total_files_in_zip": len(files),
        "documents_with_text": docs_estratti,
        "duration_seconds": total_duration,
        "company_name": company_name,
        "tokens": {
            "input_total": token_report.get("total_input", 0),
            "cached_total": token_report.get("total_cached", 0),
            "output_total": token_report.get("total_output", 0),
            "cost_eur": token_report.get("total_cost_eur", 0.0),
            "saved_by_caching_eur": token_report.get("saved_by_caching_eur", 0.0),
        },
        "dry_run": dry_run,
        "provider": provider_key_used,
        "n_fallback_batches": n_fallback_batches,
        "fallback_batch_idxs": fallback_batch_idxs,
        "n_unprocessed_files": len(unprocessed_files),
        "unprocessed_files": unprocessed_files,
    }

    emitter.emit_done(
        output_filename=final_filename,
        company_name=company_name,
        duration_seconds=total_duration,
        stats=stats,
    )

    return {
        "success": True,
        "output_path": str(final_path),
        "output_filename": final_filename,
        "company_name": company_name,
        "stats": stats,
        "merge_used_fallback": merge_result.used_fallback,
        "provider": provider_key_used,
        "n_fallback_batches": n_fallback_batches,
        "fallback_batch_idxs": fallback_batch_idxs,
        "unprocessed_files": unprocessed_files,
    }
