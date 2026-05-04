"""
V2 — DOCX Merger (Fase 7).

Unisce N file .docx parziali (prodotti da incremental_docx_builder.py) in un
unico documento finale via docxcompose.

Caratteristiche:
- Ordinamento deterministico tramite prefix numerico del filename (00, 01, ..)
- File con prefix "00_" diventa la BASE (contiene titolo+meta), gli altri
  vengono accodati
- Sezioni vuote (output_path=None nel SectionBuildResult) vengono skippate
- Fallback su python-docx atomico se docxcompose fallisce
- Mai eccezione: ritorna MergeResult con success/error
- Cleanup parziali opzionale (default off; orchestrator gestisce a fine pipeline)

API pubblica:
    merge_session_sections(session_id, output_path, ...) -> MergeResult
"""
from __future__ import annotations

import os
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docxcompose.composer import Composer
    HAS_DOCXCOMPOSE = True
except ImportError:
    HAS_DOCXCOMPOSE = False
    Composer = None  # type: ignore

from v2.incremental_docx_builder import (
    SECTIONS_BASE_DIR,
    _section_dir,
    _validate_session_id,
)


# Pattern dei file parziali: NN_slug.docx
_PARTIAL_FILE_RE = re.compile(r"^(\d{2,4})_[A-Za-z0-9_]+\.docx$")


# ──────────────────────────────────────────────────────────────────────────────
# Output
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class MergeResult:
    """Esito del merge incrementale."""
    success: bool
    output_path: Optional[Path] = None
    sections_merged: int = 0
    file_size_bytes: int = 0
    used_fallback: bool = False
    error: Optional[str] = None
    duration_seconds: float = 0.0
    skipped_files: List[str] = field(default_factory=list)


# ──────────────────────────────────────────────────────────────────────────────
# Discovery dei file parziali
# ──────────────────────────────────────────────────────────────────────────────

def list_session_partials(
    session_id: str,
    base_dir: Optional[Path] = None,
) -> List[Path]:
    """
    Lista i file parziali di una session, ordinati per prefix numerico.

    Returns:
        Lista di Path. Vuota se directory inesistente o nessun file valido.
    """
    try:
        _validate_session_id(session_id)
    except ValueError:
        return []

    section_dir = _section_dir(session_id, base_dir)
    if not section_dir.exists():
        return []

    candidates = []
    for path in section_dir.glob("*.docx"):
        m = _PARTIAL_FILE_RE.match(path.name)
        if m:
            try:
                idx = int(m.group(1))
            except ValueError:
                continue
            candidates.append((idx, path))

    candidates.sort(key=lambda x: x[0])
    return [p for _, p in candidates]


# ──────────────────────────────────────────────────────────────────────────────
# Merge primario via docxcompose
# ──────────────────────────────────────────────────────────────────────────────

def _merge_via_docxcompose(
    partial_paths: List[Path],
    output_path: Path,
) -> MergeResult:
    """
    Merge primario: docxcompose. Solleva ValueError/RuntimeError su fail
    perché il caller possa attivare il fallback.
    """
    if not partial_paths:
        return MergeResult(
            success=False, error="no_partials_to_merge",
        )

    if not HAS_DOCXCOMPOSE:
        raise RuntimeError("docxcompose_not_installed")

    # Apri il primo come base
    base_doc = Document(str(partial_paths[0]))
    composer = Composer(base_doc)

    for p in partial_paths[1:]:
        composer.append(Document(str(p)))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))

    return MergeResult(
        success=True,
        output_path=output_path,
        sections_merged=len(partial_paths),
        file_size_bytes=output_path.stat().st_size,
        used_fallback=False,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Fallback via python-docx atomico
# ──────────────────────────────────────────────────────────────────────────────

def _merge_via_atomic_fallback(
    partial_paths: List[Path],
    output_path: Path,
) -> MergeResult:
    """
    Fallback: copia il primo file come "base", poi appende il contenuto
    body-by-body degli altri tramite manipolazione XML diretta. Meno
    sofisticato di docxcompose ma non dipende da nessuna libreria oltre
    python-docx (che V2 ha già).
    """
    if not partial_paths:
        return MergeResult(success=False, error="no_partials", used_fallback=True)

    if not HAS_DOCX:
        return MergeResult(
            success=False,
            error="python_docx_not_installed",
            used_fallback=True,
        )

    try:
        from docx.oxml.ns import qn

        base = Document(str(partial_paths[0]))
        base_body = base.element.body

        for p in partial_paths[1:]:
            other = Document(str(p))
            other_body = other.element.body
            # Copia tutti gli elementi top-level eccetto il sectPr finale
            for child in list(other_body):
                if child.tag == qn("w:sectPr"):
                    continue
                # Importazione: clona ed appendi
                # NOTE: questa è una soluzione "best effort"; gli stili
                # non sono fusi (eredità da base). Sufficiente come fallback
                # quando docxcompose non funziona.
                base_body.append(child)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        base.save(str(output_path))

        return MergeResult(
            success=True,
            output_path=output_path,
            sections_merged=len(partial_paths),
            file_size_bytes=output_path.stat().st_size,
            used_fallback=True,
        )
    except Exception as e:
        return MergeResult(
            success=False,
            error=f"fallback_failed: {e}",
            used_fallback=True,
        )


# ──────────────────────────────────────────────────────────────────────────────
# API pubblica
# ──────────────────────────────────────────────────────────────────────────────

def merge_session_sections(
    session_id: str,
    output_path: Path,
    base_dir: Optional[Path] = None,
    force_fallback: bool = False,
) -> MergeResult:
    """
    Unisce le sezioni parziali di una session in un unico .docx finale.

    Args:
        session_id: identificatore sessione
        output_path: path del file finale .docx
        base_dir: opzionale per test
        force_fallback: se True, salta docxcompose e usa direttamente fallback

    Returns:
        MergeResult con success/error e metadati. Mai eccezione.
    """
    t0 = time.monotonic()

    try:
        _validate_session_id(session_id)
    except ValueError as e:
        return MergeResult(success=False, error=str(e))

    partials = list_session_partials(session_id, base_dir)

    if not partials:
        return MergeResult(
            success=False,
            error="no_partials_found",
            duration_seconds=round(time.monotonic() - t0, 3),
        )

    # Tentativo primario: docxcompose
    if not force_fallback and HAS_DOCXCOMPOSE:
        try:
            result = _merge_via_docxcompose(partials, output_path)
            result.duration_seconds = round(time.monotonic() - t0, 3)
            return result
        except Exception as e:
            print(f"[V2 MERGER] docxcompose fallito ({e}), uso fallback atomico")

    # Fallback
    result = _merge_via_atomic_fallback(partials, output_path)
    result.duration_seconds = round(time.monotonic() - t0, 3)
    return result


def merge_summary(result: MergeResult) -> dict:
    """Snapshot del risultato per logging/SSE."""
    return {
        "success": result.success,
        "sections_merged": result.sections_merged,
        "size_kb": round(result.file_size_bytes / 1024, 1),
        "used_fallback": result.used_fallback,
        "duration_seconds": result.duration_seconds,
        "error": result.error,
    }
