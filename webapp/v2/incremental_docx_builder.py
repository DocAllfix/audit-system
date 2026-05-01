"""
V2 — Incremental DOCX Builder (Fase 7).

Costruisce il documento Word di audit IN MODO INCREMENTALE: 1 file .docx
parziale per ogni macroarea SOP. Il merger (docx_merger.py) li unirà in 1
documento finale via docxcompose.

Risolve il caso "Batch 64 boss":
- V1 costruiva 1 albero XML monolitico → lxml segfault su parsed_data grandi
- V2 costruisce N alberi piccoli, ognuno proporzionale alla singola macroarea
  → memoria 1/N, niente più crash

Caratteristiche:
- File `00_header.docx` speciale: titolo + subtitle "Audit - {company_name}"
  (CRITICO per Tab 2 che legge la prima riga per identificare l'azienda)
- File `01_*.docx` ... `10_*.docx`: una macroarea SOP per file
- Ogni builder è isolato: failure su una sezione non blocca le altre
- Mai propagation di eccezioni: ritorna SectionBuildResult con success=False
- Replica pattern di V1 structured_evidence_generator senza importarlo

API pubblica:
    build_header_section(parsed_data, session_id, ...) -> SectionBuildResult
    build_macroarea_section(parsed_data, macroarea, idx, session_id, ...) -> SectionBuildResult
    build_all_sections(parsed_data, session_id, ...) -> List[SectionBuildResult]
"""
from __future__ import annotations

import os
import re
import shutil
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ──────────────────────────────────────────────────────────────────────────────
# Config
# ──────────────────────────────────────────────────────────────────────────────

# Macroaree SOP nell'ordine canonico (replica V1 report_generator.MACROAREA_ORDER)
MACROAREA_ORDER = [
    "DOCUMENTAZIONE LEGALE E SOCIETARIA",
    "REGOLARITÀ CONTRIBUTIVA E FISCALE",
    "SICUREZZA SUL LAVORO",
    "SORVEGLIANZA SANITARIA",
    "FORMAZIONE E ADDESTRAMENTO",
    "GESTIONE RISORSE UMANE",
    "GESTIONE MEZZI E ATTREZZATURE",
    "GESTIONE FORNITORI E APPALTI",
    "GESTIONE AMBIENTALE E RIFIUTI",
    "ALTRO",
]

# Path base sezioni
_WEBAPP_DIR = Path(__file__).resolve().parent.parent
SECTIONS_BASE_DIR = _WEBAPP_DIR.parent / "temp" / "sections"

# Validazione anti path-traversal sul session_id
_VALID_SESSION_ID_RE = re.compile(r"^[A-Za-z0-9_\-]{1,200}$")

# Limite caratteri per cella tabella (anti-blob)
MAX_CELL_CHARS = 5_000


# ──────────────────────────────────────────────────────────────────────────────
# Output
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class SectionBuildResult:
    """Esito della costruzione di una sezione parziale."""
    section_index: int                      # 00, 01, 02, ...
    section_name: str                       # "header" | macroarea
    success: bool
    output_path: Optional[Path] = None
    file_size_bytes: int = 0
    documents_count: int = 0
    error: Optional[str] = None
    duration_seconds: float = 0.0


# ──────────────────────────────────────────────────────────────────────────────
# Helpers (replica pattern V1, ma NON importati da V1)
# ──────────────────────────────────────────────────────────────────────────────

def _validate_session_id(session_id: str) -> None:
    if not _VALID_SESSION_ID_RE.match(session_id or ""):
        raise ValueError(f"session_id non valido: {session_id!r}")


def _section_dir(session_id: str, base_dir: Optional[Path] = None) -> Path:
    base = base_dir or SECTIONS_BASE_DIR
    return base / session_id


def _slugify(name: str) -> str:
    """Converte nome macroarea in slug filesystem-safe."""
    slug = name.lower()
    slug = re.sub(r"[àáâã]", "a", slug)
    slug = re.sub(r"[èéê]", "e", slug)
    slug = re.sub(r"[ìíî]", "i", slug)
    slug = re.sub(r"[òóô]", "o", slug)
    slug = re.sub(r"[ùúû]", "u", slug)
    slug = re.sub(r"[^a-z0-9]+", "_", slug)
    slug = slug.strip("_")
    return slug[:60] or "section"


def _safe_str(v) -> str:
    """Converte qualsiasi valore in stringa leggibile capped."""
    if v is None:
        return "n.d."
    if isinstance(v, (list, tuple)):
        return " | ".join(str(x) for x in v) if v else "n.d."
    if isinstance(v, dict):
        return " | ".join(f"{k}: {vv}" for k, vv in v.items()) if v else "n.d."
    s = str(v).strip()
    return (s[:MAX_CELL_CHARS] + "…") if len(s) > MAX_CELL_CHARS else (s or "n.d.")


def _set_cell_bg(cell, hex_color: str) -> None:
    """Imposta colore sfondo cella (hex senza #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_header_row(table, headers: List[str], bg_hex: str = "2E4057") -> None:
    """Aggiunge riga header colorata."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = str(header)
            _set_cell_bg(cell, bg_hex)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)


def _dict_to_table(doc, data: Dict[str, Any]) -> None:
    """Crea tabella 2 colonne (chiave/valore) da dict."""
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(data.items()):
        row = table.rows[i]
        row.cells[0].text = str(k)
        row.cells[1].text = _safe_str(v)
        # Stile chiave (grassetto, sfondo grigio chiaro)
        _set_cell_bg(row.cells[0], "F0F0F0")
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)


def _setup_margins(doc) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _extract_company_name(parsed_data: Dict[str, Any]) -> str:
    """Estrae nome azienda dal META (con fallback)."""
    meta = parsed_data.get("meta") or {}
    azienda = meta.get("azienda") or {}
    name = str(azienda.get("nome", "") or "").strip().upper()
    if not name or name in ("N.D.", "N/A"):
        return "AZIENDA NON IDENTIFICATA"
    return name


def _docs_in_macroarea(parsed_data: Dict[str, Any], macroarea: str) -> List[Dict[str, Any]]:
    """
    Estrae tutti i documenti del parsed_data appartenenti a una macroarea.

    Logica: il prompt universale produce sezioni con `nome` o `categoria`
    che possono mappare a una macroarea. Cerca match esatto o fuzzy.
    """
    macroarea_up = macroarea.upper()
    out = []
    for sezione in parsed_data.get("sezioni") or []:
        if not isinstance(sezione, dict):
            continue
        sec_name = str(sezione.get("nome", "")).upper()
        # Match esatto o substring (es. "01 · DOCUMENTAZIONE LEGALE" contiene
        # "DOCUMENTAZIONE LEGALE")
        if macroarea_up in sec_name or sec_name in macroarea_up:
            for doc in sezione.get("documenti") or []:
                if isinstance(doc, dict):
                    out.append(doc)
    return out


# ──────────────────────────────────────────────────────────────────────────────
# Builder: HEADER (file 00)
# ──────────────────────────────────────────────────────────────────────────────

def build_header_section(
    parsed_data: Dict[str, Any],
    session_id: str,
    docs_estratti: int = 0,
    docs_vuoti: int = 0,
    base_dir: Optional[Path] = None,
) -> SectionBuildResult:
    """
    Genera il file 00_header.docx con:
    - Titolo "RELAZIONE DI EVIDENZE STRUTTURATE - AUDIT-OS"
    - Subtitle "Audit - {company_name}" (CRITICO per Tab 2)
    - Intestazione manuale (campi auditor lead)
    - Tabella dati azienda
    - Statistiche di elaborazione
    - Indice documenti
    """
    import time
    t0 = time.monotonic()

    if not HAS_DOCX:
        return SectionBuildResult(
            section_index=0, section_name="header", success=False,
            error="python_docx_not_installed",
        )

    try:
        _validate_session_id(session_id)
    except ValueError as e:
        return SectionBuildResult(
            section_index=0, section_name="header", success=False, error=str(e),
        )

    if not isinstance(parsed_data, dict):
        return SectionBuildResult(
            section_index=0, section_name="header", success=False,
            error="parsed_data_not_dict",
        )

    section_dir = _section_dir(session_id, base_dir)
    section_dir.mkdir(parents=True, exist_ok=True)
    output_path = section_dir / "00_header.docx"

    try:
        company_name = _extract_company_name(parsed_data)
        meta = parsed_data.get("meta") or {}
        azienda_dict = meta.get("azienda") or {}
        audit_dict = meta.get("audit") or {}

        doc = Document()
        _setup_margins(doc)

        # Titolo
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("RELAZIONE DI EVIDENZE STRUTTURATE - AUDIT-OS")
        title_run.bold = True
        title_run.font.size = Pt(16)

        # Subtitle CRITICO per Tab 2
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(f"Audit - {company_name}")
        sub_run.bold = True
        sub_run.font.size = Pt(13)

        doc.add_paragraph()

        # Intestazione manuale
        doc.add_heading("INTESTAZIONE AUDIT — A COMPILAZIONE MANUALE", level=1)
        note = doc.add_paragraph()
        note_run = note.add_run(
            "I campi sotto sono lasciati volutamente vuoti: l'auditor lead li "
            "compila a mano sul documento prima della consegna."
        )
        note_run.italic = True
        note_run.font.size = Pt(9)

        intestazione = {
            "Tipo audit": "__________________________________",
            "Auditor lead": "__________________________________",
            "Data audit": "__/__/____",
            "Commessa / riferimento": "__________________________________",
            "Note intestazione": "__________________________________",
        }
        _dict_to_table(doc, intestazione)
        doc.add_paragraph()

        # Dati azienda
        if azienda_dict:
            doc.add_heading("DATI AZIENDA", level=1)
            _dict_to_table(doc, azienda_dict)
            doc.add_paragraph()

        # Statistiche
        total_docs = sum(
            len(s.get("documenti", []) or [])
            for s in (parsed_data.get("sezioni") or [])
            if isinstance(s, dict)
        )
        data_estrazione = (
            audit_dict.get("data_estrazione")
            or datetime.now().strftime("%d/%m/%Y")
        )
        stats = {
            "Data estrazione": str(data_estrazione),
            "Documenti estratti": str(docs_estratti),
            "Documenti vuoti": str(docs_vuoti),
            "Schede generate": str(total_docs),
        }
        periodo = audit_dict.get("periodo_copertura")
        if periodo:
            stats["Periodo copertura"] = str(periodo)

        doc.add_heading("STATISTICHE DI ELABORAZIONE", level=1)
        _dict_to_table(doc, stats)
        doc.add_paragraph()

        # Indice documenti
        indice = meta.get("indice") or []
        if indice and isinstance(indice, list):
            p = doc.add_paragraph("Indice Documenti")
            if p.runs:
                p.runs[0].bold = True
            valid_entries = [e for e in indice if isinstance(e, dict)]
            if valid_entries:
                idx_table = doc.add_table(rows=1 + len(valid_entries), cols=4)
                idx_table.style = "Table Grid"
                _add_header_row(idx_table, ["N", "Tipo", "Titolo", "Categoria"])
                for i, entry in enumerate(valid_entries):
                    row = idx_table.rows[i + 1]
                    row.cells[0].text = str(entry.get("n", i + 1))
                    row.cells[1].text = _safe_str(entry.get("tipo", ""))
                    row.cells[2].text = _safe_str(entry.get("titolo", ""))
                    row.cells[3].text = _safe_str(entry.get("categoria", ""))
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
            doc.add_paragraph()

        doc.save(str(output_path))
        size = output_path.stat().st_size

        return SectionBuildResult(
            section_index=0,
            section_name="header",
            success=True,
            output_path=output_path,
            file_size_bytes=size,
            documents_count=total_docs,
            duration_seconds=round(time.monotonic() - t0, 3),
        )
    except Exception as e:
        return SectionBuildResult(
            section_index=0,
            section_name="header",
            success=False,
            error=f"build_failed: {e}",
            duration_seconds=round(time.monotonic() - t0, 3),
        )


# ──────────────────────────────────────────────────────────────────────────────
# Builder: MACROAREA (file 01..10)
# ──────────────────────────────────────────────────────────────────────────────

def _render_doc_card(doc_obj, doc_entry: Dict[str, Any]) -> None:
    """Renderizza la scheda completa di un singolo documento."""
    tipo = _safe_str(doc_entry.get("tipo", "N/A"))
    titolo = _safe_str(doc_entry.get("titolo", "Documento senza titolo"))

    doc_obj.add_heading(f"{tipo} — {titolo}", level=2)

    # Header doc (campi anagrafici)
    header_fields = {
        "Tipo": tipo,
        "Categoria": doc_entry.get("categoria", ""),
        "Riferimento": doc_entry.get("riferimento", "n.d."),
        "Data documento": doc_entry.get("data_doc", "n.d."),
        "Data scadenza": doc_entry.get("data_scadenza", "n.d."),
        "Emesso da": doc_entry.get("emesso_da", "n.d."),
        "Soggetto": doc_entry.get("soggetto", "n.d."),
    }
    cat_sec = doc_entry.get("categorie_secondarie") or []
    if cat_sec:
        header_fields["Categorie secondarie"] = " | ".join(str(c) for c in cat_sec)
    _dict_to_table(doc_obj, header_fields)
    doc_obj.add_paragraph()

    # Cluster tematici (key/value libere dal prompt)
    SKIP_KEYS = {
        "tipo", "titolo", "categoria", "riferimento", "data_doc",
        "data_scadenza", "emesso_da", "soggetto", "categorie_secondarie",
    }
    extra_fields = {
        k: v for k, v in doc_entry.items()
        if k not in SKIP_KEYS and v not in (None, "", [], {})
    }
    if extra_fields:
        _dict_to_table(doc_obj, extra_fields)
        doc_obj.add_paragraph()


def build_macroarea_section(
    parsed_data: Dict[str, Any],
    macroarea: str,
    section_index: int,
    session_id: str,
    base_dir: Optional[Path] = None,
) -> SectionBuildResult:
    """
    Genera 1 file .docx parziale per una macroarea SOP.

    Args:
        parsed_data: dict completo (META + sezioni)
        macroarea: nome macroarea (es. "DOCUMENTAZIONE LEGALE E SOCIETARIA")
        section_index: numero d'ordine 1-10
        session_id: identificatore sessione
        base_dir: opzionale per test

    Returns:
        SectionBuildResult. Se la macroarea non ha documenti, success=True
        ma documents_count=0 e output_path=None (file NON creato).
    """
    import time
    t0 = time.monotonic()

    if not HAS_DOCX:
        return SectionBuildResult(
            section_index=section_index, section_name=macroarea, success=False,
            error="python_docx_not_installed",
        )

    try:
        _validate_session_id(session_id)
    except ValueError as e:
        return SectionBuildResult(
            section_index=section_index, section_name=macroarea,
            success=False, error=str(e),
        )

    if not isinstance(parsed_data, dict):
        return SectionBuildResult(
            section_index=section_index, section_name=macroarea,
            success=False, error="parsed_data_not_dict",
        )

    docs = _docs_in_macroarea(parsed_data, macroarea)
    if not docs:
        # Nessun documento per questa macroarea: success ma niente file
        return SectionBuildResult(
            section_index=section_index,
            section_name=macroarea,
            success=True,
            output_path=None,
            documents_count=0,
            duration_seconds=round(time.monotonic() - t0, 3),
        )

    section_dir = _section_dir(session_id, base_dir)
    section_dir.mkdir(parents=True, exist_ok=True)
    slug = _slugify(macroarea)
    filename = f"{section_index:02d}_{slug}.docx"
    output_path = section_dir / filename

    try:
        doc = Document()
        _setup_margins(doc)

        # Heading 1 = nome macroarea
        doc.add_heading(macroarea, level=1)

        for d in docs:
            if isinstance(d, dict):
                _render_doc_card(doc, d)

        doc.save(str(output_path))
        size = output_path.stat().st_size

        return SectionBuildResult(
            section_index=section_index,
            section_name=macroarea,
            success=True,
            output_path=output_path,
            file_size_bytes=size,
            documents_count=len(docs),
            duration_seconds=round(time.monotonic() - t0, 3),
        )
    except Exception as e:
        return SectionBuildResult(
            section_index=section_index,
            section_name=macroarea,
            success=False,
            error=f"build_failed: {e}",
            duration_seconds=round(time.monotonic() - t0, 3),
        )


# ──────────────────────────────────────────────────────────────────────────────
# Helper: build all sections for a session
# ──────────────────────────────────────────────────────────────────────────────

def build_all_sections(
    parsed_data: Dict[str, Any],
    session_id: str,
    docs_estratti: int = 0,
    docs_vuoti: int = 0,
    base_dir: Optional[Path] = None,
) -> List[SectionBuildResult]:
    """
    Genera header + tutte le 10 macroaree SOP.

    Returns:
        Lista di SectionBuildResult, ordinati per section_index.
        Le macroaree senza documenti hanno output_path=None.
    """
    results: List[SectionBuildResult] = []
    results.append(build_header_section(
        parsed_data, session_id, docs_estratti, docs_vuoti, base_dir,
    ))

    for i, macroarea in enumerate(MACROAREA_ORDER, start=1):
        results.append(build_macroarea_section(
            parsed_data, macroarea, i, session_id, base_dir,
        ))

    return results


# ──────────────────────────────────────────────────────────────────────────────
# Cleanup
# ──────────────────────────────────────────────────────────────────────────────

def cleanup_session_sections(
    session_id: str,
    base_dir: Optional[Path] = None,
) -> int:
    """
    Cancella la directory di tutte le sezioni parziali per session_id.

    Returns:
        Numero file rimossi (0 se directory non esisteva).
    """
    if os.environ.get("V2_KEEP_PARTIAL_DOCX", "false").lower() == "true":
        return 0

    try:
        _validate_session_id(session_id)
    except ValueError:
        return 0

    section_dir = _section_dir(session_id, base_dir)
    if not section_dir.exists():
        return 0

    count = sum(1 for _ in section_dir.glob("*.docx"))
    try:
        shutil.rmtree(section_dir)
    except OSError as e:
        print(f"[V2 BUILDER] Cleanup fallito {section_dir}: {e}")
        return 0
    return count


# ──────────────────────────────────────────────────────────────────────────────
# Summary
# ──────────────────────────────────────────────────────────────────────────────

def builder_summary(results: List[SectionBuildResult]) -> Dict[str, Any]:
    if not results:
        return {"total": 0, "success": 0, "failed": 0, "files_created": 0}
    success = sum(1 for r in results if r.success)
    files = sum(1 for r in results if r.output_path is not None)
    total_docs = sum(r.documents_count for r in results)
    total_bytes = sum(r.file_size_bytes for r in results)
    return {
        "total_sections": len(results),
        "success": success,
        "failed": len(results) - success,
        "files_created": files,
        "total_documents": total_docs,
        "total_bytes": total_bytes,
        "errors": [
            {"section": r.section_name, "error": r.error}
            for r in results if not r.success
        ],
    }
