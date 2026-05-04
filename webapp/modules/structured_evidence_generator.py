# ==============================================================================
# STRUCTURED EVIDENCE GENERATOR — Generatore .docx da Evidence Card YAML
# ==============================================================================
# Genera il documento Word strutturato a partire dal dict del parser.
#
# Vincoli inderogabili:
#   - generate_report_word() in report_generator.py: INTOCCABILE
#   - Prima riga del docx: deve contenere il nome azienda (compatibilità Tab 2)
#   - parsed_data None/malformato → ValueError con messaggio leggibile
#   - Tab 2 legge il testo del docx tramite extract_text_from_word():
#     tutto il contenuto deve essere in paragrafi/tabelle leggibili
# ==============================================================================

import re
from datetime import datetime
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ==============================================================================
# HELPERS INTERNI
# ==============================================================================


def _compute_deterministic_meta(parsed_data: Dict, docs_estratti: int, docs_vuoti: int) -> None:
    """
    Sovrascrive in-place i campi META calcolabili deterministicamente da Python,
    neutralizzando eventuali allucinazioni del modello.

    Campi determinati:
      - audit.data_estrazione   = oggi
      - audit.docs_estratti     = conteggio reale di file con testo
      - audit.docs_analizzati   = numero di schede documento effettivamente prodotte
      - audit.docs_vuoti        = conteggio reale di file senza testo
      - audit.periodo_copertura = min/max di data_doc reali (se ≥1 data valida)
      - meta.indice             = ricostruito da sezioni reali (N, tipo, titolo, categoria)
    """
    if not isinstance(parsed_data, dict):
        return
    meta = parsed_data.setdefault("meta", {})
    audit = meta.setdefault("audit", {}) or {}
    meta["audit"] = audit
    sezioni = parsed_data.get("sezioni", []) or []

    audit["data_estrazione"] = datetime.now().strftime("%d/%m/%Y")
    total_docs = sum(len(s.get("documenti", []) or []) for s in sezioni)
    audit["docs_estratti"] = docs_estratti
    audit["docs_analizzati"] = total_docs
    audit["docs_vuoti"] = docs_vuoti

    # periodo_copertura: min/max di data_doc in formato DD/MM/YYYY
    date_pattern = re.compile(r'^(\d{2})/(\d{2})/(\d{4})$')
    ordered = []
    for s in sezioni:
        for d in s.get("documenti", []) or []:
            dd = str(d.get("data_doc", "") or "").strip()
            m = date_pattern.match(dd)
            if m:
                ordered.append((f"{m.group(3)}{m.group(2)}{m.group(1)}", dd))
    if ordered:
        ordered.sort()
        audit["periodo_copertura"] = f"{ordered[0][1]} / {ordered[-1][1]}"

    # Indice ricostruito da sezioni reali
    indice_reale = []
    n = 1
    for s in sezioni:
        sid = str(s.get("id", "18"))
        sname = str(s.get("nome", "ALTRI"))
        categoria = f"{sid} · {sname}" if sid and not sname.startswith(sid) else sname
        for d in s.get("documenti", []) or []:
            indice_reale.append({
                "n": n,
                "tipo": str(d.get("tipo", "") or ""),
                "titolo": str(d.get("titolo", "") or ""),
                "categoria": categoria,
            })
            n += 1
    meta["indice"] = indice_reale

def _set_cell_bg(cell, hex_color: str):
    """Imposta colore sfondo cella tabella (hex senza #, es. '2E4057')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_para_bg(para, hex_color: str):
    """Imposta colore sfondo paragrafo (per note_audit)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _add_header_row(table, headers: list, bg_hex: str = "2E4057"):
    """Aggiunge riga header a una tabella con sfondo scuro e testo bianco."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        if i < len(row.cells):
            cell = row.cells[i]
            cell.text = str(header)
            _set_cell_bg(cell, bg_hex)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(header))
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)


def _dict_to_table(doc: "Document", data: dict, max_rows: int = 80):
    """
    Converte un dict in una tabella a due colonne (Campo | Valore).
    Gestisce valori annidati (dict/list) convertendoli in stringa.
    """
    if not data:
        return

    items = list(data.items())[:max_rows]
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = "Table Grid"

    # Header
    _add_header_row(table, ["Campo", "Valore"])

    # Dati
    for i, (k, v) in enumerate(items):
        row = table.rows[i + 1]
        row.cells[0].text = str(k).replace("_", " ").title()
        # Serializza valori complessi
        if isinstance(v, dict):
            row.cells[1].text = ", ".join(f"{kk}: {vv}" for kk, vv in v.items())
        elif isinstance(v, list):
            row.cells[1].text = " | ".join(str(x) for x in v) if v else "—"
        elif v is None or str(v).strip() in ("", "null", "None"):
            row.cells[1].text = "n.d."
        else:
            row.cells[1].text = str(v)

        # Stile celle dati
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


def _value_to_str(v) -> str:
    """Serializza un valore YAML in stringa leggibile."""
    if v is None or str(v).strip() in ("", "null", "None"):
        return "n.d."
    if isinstance(v, dict):
        return " | ".join(f"{kk}: {vv}" for kk, vv in v.items())
    if isinstance(v, list):
        return " | ".join(str(x) for x in v) if v else "—"
    return str(v)


# ==============================================================================
# FUNZIONE PRINCIPALE
# ==============================================================================

def generate_structured_evidence_docx(
    parsed_data: Dict,
    output_path: str,
    docs_estratti: int = 0,
    docs_vuoti: int = 0,
    failed_files: list = None
) -> str:
    """
    Genera il file .docx strutturato a partire dal dict del parser YAML.

    Struttura del documento:
        1. Titolo: "RELAZIONE DI EVIDENZE STRUTTURATE - AUDIT-OS"
        2. Subtitolo: "Audit - {company_name}"  ← Tab 2 legge questa riga
        3. Intestazione: Data | Estratti | Vuoti | Analizzati (deterministici)
        4. Heading 1 "META — AUDIT E AZIENDA" + tabelle meta
        5. Per ogni sezione: Heading 1 "XX · NOME" →
               per ogni documento: Heading 2 "{Tipo} — {titolo}" +
               tabella header doc + tabelle cluster

    Args:
        parsed_data:   dict da parse_structured_response() — NON None
        output_path:   percorso file .docx di output
        docs_estratti: numero documenti con testo estratto
        docs_vuoti:    numero documenti senza testo

    Returns:
        output_path confermato

    Raises:
        ValueError:    se parsed_data è None o manca la chiave "meta"
        ImportError:   se python-docx non è installato
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx non installato. Eseguire: pip install python-docx"
        )

    if parsed_data is None:
        raise ValueError(
            "generate_structured_evidence_docx: parsed_data è None. "
            "Il PROMPT UNIVERSALE non ha prodotto output parsabile."
        )

    if not isinstance(parsed_data, dict) or "meta" not in parsed_data:
        raise ValueError(
            "generate_structured_evidence_docx: parsed_data malformato "
            "(manca la chiave 'meta'). Verificare l'output del parser."
        )

    _compute_deterministic_meta(parsed_data, docs_estratti, docs_vuoti)

    meta = parsed_data.get("meta", {})
    sezioni = parsed_data.get("sezioni", [])

    # Estrai nome azienda
    azienda_dict = meta.get("azienda", {}) or {}
    company_name = str(azienda_dict.get("nome", "") or "").strip().upper()
    if not company_name or company_name in ("N.D.", "N/A"):
        company_name = "AZIENDA NON IDENTIFICATA"

    audit_dict = meta.get("audit", {}) or {}
    data_estrazione = str(audit_dict.get("data_estrazione", "") or "").strip()
    if not data_estrazione:
        data_estrazione = datetime.now().strftime("%d/%m/%Y")

    docs_analizzati = docs_estratti - docs_vuoti
    total_docs = sum(len(s.get("documenti", [])) for s in sezioni)

    # ==================================================================
    # COSTRUZIONE DOCUMENTO
    # ==================================================================
    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ------------------------------------------------------------------
    # 1. TITOLO (prima riga del doc)
    # ------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("RELAZIONE DI EVIDENZE STRUTTURATE - AUDIT-OS")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # ------------------------------------------------------------------
    # 2. SUBTITOLO — OBBLIGATORIO per Tab 2
    # Deve contenere il nome azienda come prima riga significativa
    # Tab 2 estrae questo per identificare l'azienda
    # ------------------------------------------------------------------
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"Audit - {company_name}")
    sub_run.bold = True
    sub_run.font.size = Pt(13)

    # ------------------------------------------------------------------
    # 3. (Intestazione di sistema rimossa: le statistiche sono ora nel
    #     blocco STATISTICHE DI ELABORAZIONE sotto per evitare duplicazioni.)
    # ------------------------------------------------------------------
    doc.add_paragraph()  # spazio prima del META

    # ------------------------------------------------------------------
    # 4. SEZIONE META — INTESTAZIONE MANUALE + DATI AZIENDA + STATISTICHE
    # ------------------------------------------------------------------
    doc.add_heading("INTESTAZIONE AUDIT — A COMPILAZIONE MANUALE", level=1)
    p = doc.add_paragraph()
    note_run = p.add_run(
        "I campi sotto sono lasciati volutamente vuoti: l'auditor lead li compila "
        "a mano sul documento prima della consegna."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    intestazione_manuale = {
        "Tipo audit":            "__________________________________",
        "Auditor lead":          "__________________________________",
        "Data audit":            "__/__/____",
        "Commessa / riferimento": "__________________________________",
        "Note intestazione":     "__________________________________",
    }
    _dict_to_table(doc, intestazione_manuale)
    doc.add_paragraph()

    # Tabella azienda
    if azienda_dict:
        doc.add_heading("DATI AZIENDA", level=1)
        _dict_to_table(doc, azienda_dict)
        doc.add_paragraph()

    # Statistiche deterministiche di elaborazione
    stats_fields = {
        "Data estrazione":   str(audit_dict.get("data_estrazione", data_estrazione)),
        "Documenti estratti": str(audit_dict.get("docs_estratti", docs_estratti)),
        "Documenti vuoti":    str(audit_dict.get("docs_vuoti", docs_vuoti)),
        "Schede generate":    str(audit_dict.get("docs_analizzati", total_docs)),
    }
    periodo = audit_dict.get("periodo_copertura")
    if periodo:
        stats_fields["Periodo copertura"] = str(periodo)
    doc.add_heading("STATISTICHE DI ELABORAZIONE", level=1)
    _dict_to_table(doc, stats_fields)
    doc.add_paragraph()

    # Indice documenti
    indice = meta.get("indice", []) or []
    if indice:
        p = doc.add_paragraph("Indice Documenti")
        p.runs[0].bold = True
        idx_table = doc.add_table(rows=1 + len(indice), cols=4)
        idx_table.style = "Table Grid"
        _add_header_row(idx_table, ["N", "Tipo", "Titolo", "Categoria"])
        for i, entry in enumerate(indice):
            if not isinstance(entry, dict):
                continue
            row = idx_table.rows[i + 1]
            row.cells[0].text = str(entry.get("n", i + 1))
            row.cells[1].text = str(entry.get("tipo", ""))
            row.cells[2].text = str(entry.get("titolo", ""))
            row.cells[3].text = str(entry.get("categoria", ""))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # 5. SEZIONI DOCUMENTI
    # ------------------------------------------------------------------
    import re as _re
    from collections import OrderedDict as _OD

    def _render_doc_detail(doc_obj, doc_entry):
        """Renderizza la scheda completa di un singolo documento."""
        tipo = str(doc_entry.get("tipo", "N/A"))
        titolo = str(doc_entry.get("titolo", "Documento senza titolo"))

        doc_obj.add_heading(f"{tipo} — {titolo}", level=2)

        # Tabella header documento (campi fissi)
        header_fields = {
            "Tipo": tipo,
            "Categoria": doc_entry.get("categoria", ""),
            "Riferimento": doc_entry.get("riferimento", "n.d."),
            "Data documento": doc_entry.get("data_doc", "n.d."),
            "Data scadenza": doc_entry.get("data_scadenza", "n.d."),
            "Emesso da": doc_entry.get("emesso_da", "n.d."),
            "Soggetto": doc_entry.get("soggetto", "n.d."),
        }

        # FIX #4: Commessa eliminata dai campi di header. Se appare in un
        # documento, il prompt la instrada in un cluster tematico libero.

        cat_sec = doc_entry.get("categorie_secondarie", []) or []
        if cat_sec:
            header_fields["Categorie secondarie"] = " | ".join(str(c) for c in cat_sec)
        _dict_to_table(doc_obj, header_fields)
        doc_obj.add_paragraph()

        # Cluster di dati liberi
        clusters = doc_entry.get("cluster", {}) or {}
        for cluster_name, cluster_data in clusters.items():
            if not cluster_data:
                continue
            p = doc_obj.add_paragraph(f"▸ {cluster_name}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            if isinstance(cluster_data, dict):
                _dict_to_table(doc_obj, cluster_data)
            elif isinstance(cluster_data, list):
                if cluster_data and isinstance(cluster_data[0], dict):
                    cols = list(cluster_data[0].keys())
                    tbl = doc_obj.add_table(rows=1 + len(cluster_data), cols=len(cols))
                    tbl.style = "Table Grid"
                    _add_header_row(tbl, cols)
                    for ri, item in enumerate(cluster_data):
                        if not isinstance(item, dict):
                            continue
                        row = tbl.rows[ri + 1]
                        for ci, col in enumerate(cols):
                            if ci < len(row.cells):
                                row.cells[ci].text = _value_to_str(item.get(col))
                                for para in row.cells[ci].paragraphs:
                                    for run in para.runs:
                                        run.font.size = Pt(9)
                else:
                    for item in cluster_data:
                        doc_obj.add_paragraph(str(item), style="List Bullet")
            else:
                doc_obj.add_paragraph(str(cluster_data))
            doc_obj.add_paragraph()

        # Firme
        firme = doc_entry.get("firme", {})
        if isinstance(firme, dict) and firme:
            p = doc_obj.add_paragraph("▸ Firme")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
            _dict_to_table(doc_obj, firme)
            doc_obj.add_paragraph()

    def _render_omogenei_table(doc_obj, group, tipo_label):
        """
        Regola 2.6 — Tabella riepilogativa per gruppi di 3+ documenti omogenei,
        seguita da scheda di dettaglio per ciascun documento del gruppo.
        """
        p = doc_obj.add_paragraph(f"▸ {tipo_label} — Documenti omogenei ({len(group)} documenti)")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

        cols = ["N", "Tipo", "Titolo", "Riferimento", "Data Doc", "Emesso da"]
        tbl = doc_obj.add_table(rows=1 + len(group), cols=len(cols))
        tbl.style = "Table Grid"
        _add_header_row(tbl, cols)
        for ri, entry in enumerate(group):
            row = tbl.rows[ri + 1]
            row.cells[0].text = str(ri + 1)
            row.cells[1].text = str(entry.get("tipo", ""))
            row.cells[2].text = str(entry.get("titolo", ""))[:80]
            row.cells[3].text = _value_to_str(entry.get("riferimento"))
            row.cells[4].text = _value_to_str(entry.get("data_doc"))
            row.cells[5].text = _value_to_str(entry.get("emesso_da"))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc_obj.add_paragraph()

        p2 = doc_obj.add_paragraph("Schede di dettaglio:")
        p2.runs[0].italic = True
        p2.runs[0].font.size = Pt(9)
        doc_obj.add_paragraph()
        for entry in group:
            _render_doc_detail(doc_obj, entry)

    for sezione in sezioni:
        if not isinstance(sezione, dict):
            continue

        sid = sezione.get("id", "18")
        snome = sezione.get("nome", "ALTRI")
        documenti = sezione.get("documenti", []) or []

        if not documenti:
            continue

        snome_clean = _re.sub(r'^\d+\s*[·•\-]\s*', '', snome).strip() or snome
        doc.add_heading(f"{sid} · {snome_clean}", level=1)

        # Raggruppa doc per tipo, preservando l'ordine di prima apparizione
        tipo_order = []
        tipo_groups = _OD()
        for doc_entry in documenti:
            if not isinstance(doc_entry, dict):
                continue
            t = str(doc_entry.get("tipo", "N/A"))
            if t not in tipo_groups:
                tipo_groups[t] = []
                tipo_order.append(t)
            tipo_groups[t].append(doc_entry)

        for t in tipo_order:
            group = tipo_groups[t]
            if len(group) >= 3:
                # Regola 2.7: tabella riepilogativa + dettaglio solo per anomalie
                _render_omogenei_table(doc, group, t)
            else:
                # Comportamento standard
                for doc_entry in group:
                    _render_doc_detail(doc, doc_entry)

    # ==================================================================
    # SEZIONE: DOCUMENTI NON ELABORATI (se presenti)
    # ==================================================================
    if failed_files:
        # Mappa categoria errore → descrizione user-friendly
        REASON_LABELS = {
            "extraction_failed": "Estrazione testo fallita",
            "ocr_failed":         "OCR fallito (testo insufficiente)",
            "ocr_recovery_failed":"OCR fallito anche dopo recovery",
            "ocr_unavailable":    "OCR non disponibile (servizio non raggiungibile)",
            "ocr_timeout":        "OCR oltre il tempo limite (file probabilmente troppo pesante)",
            "unsupported_format": "Formato non supportato",
            "ai_error":           "Errore AI durante l'analisi",
        }

        doc.add_heading("DOCUMENTI NON ELABORATI", level=1)

        intro = doc.add_paragraph(
            "I file elencati di seguito non sono stati inclusi nel report. "
            "Verificare manualmente il loro contenuto se necessario per l'audit."
        )
        intro.runs[0].italic = True

        # Raggruppa per categoria di errore
        from collections import OrderedDict as _OD2
        groups = _OD2()
        for ff in failed_files:
            t = ff.get("type", "unknown")
            groups.setdefault(t, []).append(ff)

        for ftype, items in groups.items():
            label = REASON_LABELS.get(ftype, ftype)
            doc.add_heading(f"{label} ({len(items)})", level=2)

            ftbl = doc.add_table(rows=1 + len(items), cols=2)
            ftbl.style = "Table Grid"
            hdr = ftbl.rows[0].cells
            hdr[0].text = "File"
            hdr[1].text = "Motivo"
            for i, ff in enumerate(items, start=1):
                row = ftbl.rows[i].cells
                row[0].text = ff.get("relative_path") or ff.get("filename", "?")
                row[1].text = str(ff.get("reason", ""))[:200]
            doc.add_paragraph()

    # ==================================================================
    # SALVA FILE
    # ==================================================================
    doc.save(output_path)
    return output_path


# ==============================================================================
# TEST STANDALONE
# ==============================================================================

if __name__ == "__main__":
    import tempfile
    import os

    print("=== Test structured_evidence_generator ===")
    print(f"python-docx: {'OK' if HAS_DOCX else 'NON INSTALLATO — eseguire: pip install python-docx'}")

    if not HAS_DOCX:
        print("Test saltato — python-docx mancante.")
    else:
        # Test con parsed_data=None → ValueError
        try:
            generate_structured_evidence_docx(None, "/tmp/test.docx")
            assert False, "Avrebbe dovuto sollevare ValueError"
        except ValueError as e:
            print(f"✅ ValueError su None: {e}")

        # Test con fixture minima
        fixture = {
            "meta": {
                "audit": {},  # blocco compilato deterministicamente da _compute_deterministic_meta
                "azienda": {
                    "nome": "TEST S.R.L.",
                    "piva": "12345678901",
                    "sede": "Via Roma 1 — 20100 Milano (MI)",
                },
                "indice": [
                    {"n": 1, "tipo": "VIS", "titolo": "Visura Camerale",
                     "categoria": "08 · LEGALE/SOCIETARIA"}
                ],
                "abbrev_aggiunte": []
            },
            "sezioni": [
                {
                    "id": "08",
                    "nome": "LEGALE/SOCIETARIA",
                    "documenti": [
                        {
                            "tipo": "Visura Camerale",
                            "categoria": "08 · LEGALE/SOCIETARIA",
                            "titolo": "Visura Camerale TEST S.R.L.",
                            "riferimento": "n.d.",
                            "data_doc": "15/01/2026",
                            "data_scadenza": "non applicabile",
                            "emesso_da": "CCIAA Milano",
                            "soggetto": "TEST S.R.L.",
                            "norme_pertinenti": ["ISO 9001:2015"],
                            "cluster": {
                                "Identificativi": {
                                    "ragione_sociale": "TEST S.R.L.",
                                    "piva": "12345678901",
                                    "forma_giuridica": "S.r.l.",
                                }
                            },
                            "firme": {
                                "emittente": "Presente",
                                "ricevente": "n.d.",
                                "data_firma": "n.d."
                            },
                            "note_audit": "Rif. ISO 9001 cl. 7.5.2"
                        }
                    ]
                }
            ]
        }

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        result_path = generate_structured_evidence_docx(fixture, tmp_path, docs_estratti=1, docs_vuoti=0)
        assert os.path.exists(result_path), "File .docx non creato"
        size = os.path.getsize(result_path)
        assert size > 1000, f"File .docx troppo piccolo ({size} bytes)"
        print(f"✅ generate_structured_evidence_docx: OK ({size} bytes)")
        print(f"   File: {result_path}")

        # Verifica che la seconda riga contenga il nome azienda (compatibilità Tab 2)
        from docx import Document as _Doc
        check_doc = _Doc(result_path)
        paras = [p.text.strip() for p in check_doc.paragraphs if p.text.strip()]
        assert any("TEST S.R.L." in p for p in paras[:3]), \
            f"Nome azienda non trovato nelle prime 3 righe: {paras[:3]}"
        print(f"✅ Compatibilità Tab 2: nome azienda in riga {next(i for i,p in enumerate(paras[:3]) if 'TEST S.R.L.' in p)+1}")

        os.unlink(tmp_path)
        print("\n✅ Tutti i test passati.")
