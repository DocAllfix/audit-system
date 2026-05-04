# ==============================================================================
# REPORT GENERATOR - Modulo generazione report Word
# ==============================================================================

import os
import sys
import json
import tempfile
import zipfile
import contextlib
import io
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import streamlit as st

# Aggiungi percorsi
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from config import TEMP_DIR, EXECUTION_DIR, OUTPUT_DIR, BATCH_SIZE, MAX_BATCH_CHARS, MAX_API_WORKERS, MAX_OCR_WORKERS
from modules.gemini_client import GeminiClient
from modules.logger import AuditLogger

MACROAREA_ORDER = [
    "DOCUMENTAZIONE LEGALE E SOCIETARIA",
    "REGOLARITÀ CONTRIBUTIVA E FISCALE",
    "SICUREZZA SUL LAVORO",
    "SORVEGLIANZA SANITARIA",
    "FORMAZIONE E ADDESTRAMENTO",
    "GESTIONE RISORSE UMANE",
    "GESTIONE MEZZI E ATTREZZATURE",
    "GESTIONE FORNITORI E APPALTI",
    "GESTIONE AMBIENTALE E RIFIUTI",
    "ALTRO"
]

# Nomi ZIP generici che non devono essere usati come nome azienda (R-3)
NOMI_GENERICI_ZIP = frozenset({
    'allegati', 'documenti', 'files', 'doc', 'docs', 'archive',
    'archivio', 'cartella', 'folder', 'zip', 'upload', 'pratica',
    'materiale', 'materiali', 'tutto', 'vari', 'misc', 'allegato',
    'folder upload', 'cartella selezionata', 'new folder', 'documenti audit',
    'audit', 'report', 'relazione', 'evidenze', 'pratiche'
})


def _categorize_file(filename: str) -> str:
    """Categorizza un file in base all'estensione."""
    ext = os.path.splitext(filename)[1].lower()
    
    # File da saltare (non sono documenti audit)
    SKIP_EXTENSIONS = {
        '.vcf',      # vCard (contatti)
        '.exe', '.dll', '.bat', '.cmd', '.msi', '.com',  # Eseguibili
        '.sys', '.ini', '.log', '.tmp', '.bak',  # File di sistema
        '.lnk', '.url',  # Collegamenti
        '.db', '.sqlite', '.mdb',  # Database
        '.mp3', '.mp4', '.avi', '.mov', '.wav', '.flac',  # Media
        '.psd', '.ai', '.sketch',  # Design
        '.dmg', '.iso', '.tar', '.gz', '.rar', '.7z',  # Archivi (tranne zip)
        # --- Aggiunte 2026-04-19 (da analisi 417 log produzione) ---
        # .dwg/.dwl/.dwl2/.css/.svg/.download/.crdownload/.dsf2/.cam/.mpp
        # → rimossi da SKIP: ora hanno categoria "unsupported" o handler dedicato
    }
    
    if ext in SKIP_EXTENSIONS:
        return "skip"
    elif ext == ".pdf":
        return "pdf"
    elif ext in [".doc", ".docx"]:
        return "word"
    elif ext in [".xls", ".xlsx", ".xlsm"]:
        return "excel"
    elif ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]:
        return "image"
    elif ext in [".txt", ".csv"]:
        return "text"
    elif ext == ".rtf":
        return "rtf"
    elif ext == ".p7m":
        return "p7m"
    elif ext == ".msg":
        return "email"
    elif ext == ".zip":
        return "zip"
    # --- Aggiunte 2026-04-19 ---
    elif ext == ".eml":
        return "eml"
    elif ext in (".html", ".htm"):
        return "html"
    elif ext == ".xml":
        return "xml"
    elif ext == ".odt":
        return "odt"
    elif ext == ".ods":
        return "ods"
    elif ext in (".pptx", ".ppt"):
        return "pptx"
    elif ext in (".heic", ".heif"):
        return "heic"
    elif ext == ".mpp":
        return "mpp"
    elif ext == ".cam":
        return "cam"
    # Formati noti ma senza testo audit estraibile → visibili all'utente come "non supportati"
    elif ext in (".dwg", ".dwl", ".dwl2", ".css", ".svg",
                 ".download", ".crdownload", ".dsf2"):
        return "unsupported"
    return "other"


def _extract_nested_zip(zip_path: str, extract_dir: str, depth: int = 0, max_depth: int = 10) -> List[Dict]:
    """
    Estrae ricorsivamente un file ZIP, includendo ZIP nidificati.
    
    Args:
        zip_path: Percorso del file ZIP da estrarre
        extract_dir: Directory di destinazione
        depth: Livello di nidificazione corrente
        max_depth: Massima profondità di nidificazione (default 10)
    
    Returns:
        Lista di file estratti con metadati
    """
    if depth > max_depth:
        print(f"[WARNING] Raggiunta profondità massima ({max_depth}), stop estrazione ricorsiva")
        return []
    
    extracted_files = []
    nested_zips = []  # ZIP da estrarre successivamente
    
    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            for file_info in zf.infolist():
                if file_info.is_dir():
                    continue
                
                # Ignora file di sistema
                filename = os.path.basename(file_info.filename)
                if filename.startswith('.') or filename.startswith('__') or filename == '':
                    continue
                
                # Estrai file con gestione Long Paths (Fix Windows limit 260 chars)
                import shutil
                # Normalizza path interno allo zip
                normalized_zip_path = os.path.normpath(file_info.filename)
                target_path = os.path.join(extract_dir, normalized_zip_path)
                target_path = os.path.normpath(target_path)
                
                # Applica prefisso Long Path per Windows se necessario e se path assoluto
                if os.name == 'nt' and len(os.path.abspath(target_path)) >= 250:
                    abs_path = os.path.abspath(target_path)
                    if not abs_path.startswith('\\\\?\\'):
                        target_path = '\\\\?\\' + abs_path
                
                # Assicura directory esistente
                os.makedirs(os.path.dirname(target_path), exist_ok=True)
                
                # Scrivi file bypassando limiti OS
                with zf.open(file_info) as source, open(target_path, "wb") as target:
                    shutil.copyfileobj(source, target)
                
                extracted_path = target_path
                category = _categorize_file(filename)
                
                # Salta file non-documento
                if category == "skip":
                    print(f"[INFO] Saltato file non-documento: {filename}")
                    continue
                
                # Se è un ZIP, lo segniamo per estrazione ricorsiva
                if category == "zip":
                    nested_zips.append(extracted_path)
                else:
                    extracted_files.append({
                        "filename": filename,
                        "path": extracted_path,
                        "size": file_info.file_size,
                        "category": category
                    })
    except zipfile.BadZipFile:
        print(f"[ERROR] File ZIP corrotto: {zip_path}")
        return []
    
    # Estrai ricorsivamente i ZIP nidificati
    for nested_zip_path in nested_zips:
        print(f"[INFO] Estrazione ZIP nidificato (livello {depth + 1}): {os.path.basename(nested_zip_path)}")
        nested_extract_dir = os.path.splitext(nested_zip_path)[0]  # Cartella con nome del ZIP
        os.makedirs(nested_extract_dir, exist_ok=True)
        nested_files = _extract_nested_zip(nested_zip_path, nested_extract_dir, depth + 1, max_depth)
        extracted_files.extend(nested_files)
        # Rimuovi il file ZIP dopo l'estrazione per non confonderlo con documenti
        try:
            os.remove(nested_zip_path)
        except:
            pass
    
    return extracted_files


def extract_from_folder(folder_path: str) -> Tuple[List[Dict], str]:
    """
    Legge file direttamente da una cartella locale (alternativa a ZIP).
    Scansiona ricorsivamente tutte le sottocartelle.
    
    Args:
        folder_path: Percorso assoluto della cartella
    
    Returns:
        Tuple (lista file con metadati, percorso cartella)
    """
    if not os.path.isdir(folder_path):
        raise ValueError(f"Il percorso non è una cartella valida: {folder_path}")
    
    extracted_files = []
    
    for root, dirs, files in os.walk(folder_path):
        # Ignora cartelle nascoste
        dirs[:] = [d for d in dirs if not d.startswith('.') and not d.startswith('__')]
        
        for filename in files:
            # Ignora file nascosti
            if filename.startswith('.') or filename.startswith('__'):
                continue
            
            filepath = os.path.join(root, filename)
            category = _categorize_file(filename)
            
            # Salta file non-documento
            if category == "skip":
                print(f"[INFO] Saltato file non-documento: {filename}")
                continue
            
            # Salta ZIP (non estraiamo ZIP annidati da cartelle)
            if category == "zip":
                print(f"[INFO] Saltato ZIP in cartella: {filename} (usa upload ZIP per matrioska)")
                continue
            
            try:
                file_size = os.path.getsize(filepath)
                extracted_files.append({
                    "filename": filename,
                    "path": filepath,
                    "size": file_size,
                    "category": category
                })
            except Exception as e:
                print(f"[WARNING] Impossibile leggere {filename}: {e}")
    
    print(f"[INFO] Scansione cartella completata: {len(extracted_files)} file trovati")
    
    return extracted_files, folder_path


def extract_from_uploaded_files(uploaded_files: List) -> Tuple[List[Dict], str]:
    """
    Gestisce i file caricati via folder picker nativo di Streamlit.
    Salva i file in una cartella temporanea e restituisce metadati.
    
    Args:
        uploaded_files: Lista di UploadedFile objects da st.file_uploader(accept_multiple_files="directory")
    
    Returns:
        Tuple (lista file con metadati, percorso cartella temp)
    """
    import tempfile
    
    # Crea cartella temporanea
    temp_dir = tempfile.mkdtemp(prefix="audit_folder_")
    print(f"[INFO] Cartella temporanea creata: {temp_dir}")
    
    extracted_files = []
    
    for uploaded_file in uploaded_files:
        filename = uploaded_file.name
        category = _categorize_file(filename)
        
        # Salta file non-documento
        if category == "skip":
            print(f"[INFO] Saltato file non-documento: {filename}")
            continue
        
        # Salta ZIP (non estraiamo ZIP annidati da cartelle)
        if category == "zip":
            print(f"[INFO] Saltato ZIP in cartella: {filename}")
            continue
        
        try:
            # Salva file nella cartella temporanea
            # USA IL PATH RELATIVO COMPLETO per preservare la struttura cartelle
            file_path = os.path.join(temp_dir, filename)
            
            # CREA LE SOTTOCARTELLE se il filename contiene path relativi (es. "subfolder/file.pdf")
            file_dir = os.path.dirname(file_path)
            if file_dir and file_dir != temp_dir:
                os.makedirs(file_dir, exist_ok=True)
            
            with open(file_path, "wb") as f:
                data = uploaded_file.read() if hasattr(uploaded_file, 'read') else bytes(uploaded_file)
                f.write(data)
            
            file_size = os.path.getsize(file_path)
            extracted_files.append({
                "filename": filename,
                "path": file_path,
                "size": file_size,
                "category": category
            })
        except Exception as e:
            print(f"[WARNING] Impossibile salvare {filename}: {e}")
    
    print(f"[INFO] File caricati salvati: {len(extracted_files)} file")
    
    return extracted_files, temp_dir


def extract_zip(zip_data) -> Tuple[List[Dict], str]:
    """
    Estrae un file ZIP e restituisce lista file con metadati.
    Supporta ZIP nidificati (matrioska) fino a 10 livelli di profondità.

    Args:
        zip_data: bytes del file ZIP (da FastAPI UploadFile.read())

    Returns:
        Tuple (lista file estratti, percorso temp directory)
    """
    # Crea directory temporanea
    temp_extract_dir = os.path.join(TEMP_DIR, f"extract_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    os.makedirs(temp_extract_dir, exist_ok=True)

    # Normalizza input → bytes
    if isinstance(zip_data, (bytes, bytearray)):
        raw = zip_data
    elif hasattr(zip_data, 'getvalue'):
        raw = zip_data.getvalue()
    elif hasattr(zip_data, 'read'):
        raw = zip_data.read()
    else:
        raw = bytes(zip_data)

    # Salva ZIP su disco per l'estrazione ricorsiva
    temp_zip_path = os.path.join(TEMP_DIR, "uploaded.zip")
    with open(temp_zip_path, "wb") as f:
        f.write(raw)
    
    # Estrazione ricorsiva
    extracted_files = _extract_nested_zip(temp_zip_path, temp_extract_dir, depth=0)
    
    print(f"[INFO] Estrazione completata: {len(extracted_files)} file totali")
    
    return extracted_files, temp_extract_dir


def extract_text_from_files(files: List[Dict], progress_callback=None) -> List[Dict]:
    """
    Estrae testo dai file usando gli script esistenti.
    
    Args:
        files: Lista file estratti
        progress_callback: Funzione per aggiornare progress bar
    
    Returns:
        Lista di dict con filename e content
    """
    # Importa moduli di estrazione esistenti
    sys.path.insert(0, EXECUTION_DIR)
    try:
        from extract_content import (
            extract_text_from_pdf,
            extract_text_from_docx,
            extract_text_from_txt,
            extract_text_from_doc_legacy,
            extract_text_from_excel,
            extract_text_from_msg,
        )
    except ImportError as e:
        st.error(f"Errore import moduli estrazione: {e}")
        return []
    
    documents = []
    total = len(files)
    
    for i, file_info in enumerate(files):
        if progress_callback:
            progress_callback(i / total, f"Estrazione {file_info['filename']}...")
        
        filepath = file_info["path"]
        category = file_info["category"]
        content = ""
        
        try:
            if category == "pdf":
                with contextlib.redirect_stderr(io.StringIO()):
                    result = extract_text_from_pdf(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "word":
                if filepath.endswith(".docx"):
                    result = extract_text_from_docx(filepath)
                else:
                    result = extract_text_from_doc_legacy(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "excel":
                result = extract_text_from_excel(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "text":
                result = extract_text_from_txt(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "image":
                # Per immagini, segna come da OCR
                content = "[IMMAGINE - OCR richiesto]"
            elif category == "email":
                result = extract_text_from_msg(filepath)
                content = result[0] if isinstance(result, tuple) else result
        except Exception as e:
            content = f"[Errore estrazione: {str(e)}]"
        
        # Assicura che content sia una stringa
        if content and isinstance(content, str) and len(content.strip()) > 10:
            documents.append({
                "filename": file_info["filename"],
                "content": content
            })
    
    if progress_callback:
        progress_callback(1.0, "Estrazione completata!")
    
    return documents


def generate_report_word(paragraphs: List[Dict], stats: Dict, company_name: str = "") -> bytes:
    """
    Genera il file Word del report.
    
    Args:
        paragraphs: Lista paragrafi generati
        stats: Statistiche qualità
        company_name: Nome azienda per filename
    
    Returns:
        Bytes del file Word generato
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Imposta margini
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # TITOLO
    title = doc.add_heading("RELAZIONE DI EVIDENZE DI AUDIT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sottotitolo con nome azienda
    if company_name:
        subtitle = doc.add_paragraph(f"Audit - {company_name}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # INTESTAZIONE OBBLIGATORIA
    doc.add_paragraph()
    header_para = doc.add_paragraph()
    header_para.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    header_para.add_run(f"Documenti estratti: {stats.get('total_docs', len(paragraphs))}\n")
    header_para.add_run(f"Documenti analizzati: {stats.get('total_paragraphs', len(paragraphs))}\n")
    header_para.add_run(f"Media parole per paragrafo: {stats.get('avg_words', 0):.0f}\n")
    
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    doc.add_paragraph()
    
    # PARAGRAFI
    current_category = None
    for para in paragraphs:
        # Categoria (se presente)
        category = para.get("categoria", "")
        if category and category != current_category:
            doc.add_heading(category, 1)
            current_category = category
        
        # Sottotitolo numerato
        numero = para.get("numero", "")
        sottotitolo = para.get("sottotitolo", "")
        heading_text = f"[{numero}] {sottotitolo}" if numero else sottotitolo
        doc.add_heading(heading_text, 2)
        
        # Contenuto
        contenuto = para.get("contenuto", "")
        content_para = doc.add_paragraph(contenuto)
        content_para.paragraph_format.line_spacing = 1.5
        content_para.paragraph_format.space_after = Pt(12)
    
    # Salva in bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def optimize_text_content(text: str) -> str:
    """
    Ottimizza il contenuto testuale rimuovendo whitespace superfluo.
    Riduce il consumo di token e velocizza l'elaborazione LLM senza perdere info.
    """
    if not text:
        return ""
    import re
    # Sostituisce tab e spazi multipli con spazio singolo
    text = re.sub(r'[ \t]+', ' ', text)
    # Riduce sequenze di newline (max 2)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def create_smart_batches(documents: List[Dict], max_files: int = 8, max_chars: int = 100000) -> List[List[Dict]]:
    """
    Raggruppa i documenti in batch intelligenti basati sulla lunghezza del testo.
    Usa algoritmo First Fit Decreasing per ottimizzare il riempimento ma isolare i file enormi.
    """
    # Ordina per lunghezza decrescente
    sorted_docs = sorted(documents, key=lambda d: len(d.get("content", "")), reverse=True)
    
    batches = []
    batch_chars = [] # Traccia caratteri per ogni batch
    
    for doc in sorted_docs:
        content_len = len(doc.get("content", ""))
        filename = doc.get("filename", "unknown")
        
        # Cerca il primo batch dove il file "entra"
        assigned = False
        for i, batch in enumerate(batches):
            # Se il batch è vuoto (nuovo), entra sempre (anche se > max_chars, sarà isolato)
            # Se ha già file, deve rispettare entrambi i limiti
            current_len = batch_chars[i]
            
            # Condizioni per aggiungere:
            # 1. C'è spazio negli slot file
            # 2. C'è spazio nel budget caratteri
            if len(batch) < max_files and (current_len + content_len <= max_chars):
                batch.append(doc)
                batch_chars[i] += content_len
                assigned = True
                break
        
        # Se non entra in nessuno, crea nuovo batch
        if not assigned:
            batches.append([doc])
            batch_chars.append(content_len)
            
    return batches


# ==============================================================================
# PIPELINE STRUTTURATA — PROMPT UNIVERSALE ADATTIVO v2.1
# Sostituisce la pipeline narrativa come default di Tab 1 (decisione 2026-04-18)
# La pipeline narrativa (analyze_batch + generate_report_word) rimane nel file
# come dead code — NON cancellare: è il fallback di sicurezza.
# ==============================================================================

def _pipeline_strutturata(
    documents: List[Dict],
    failed_files: List[Dict],
    api_key: str,
    input_name: str,
    total_files: int,
    progress_callback=None,
    status_callback=None,
    logger=None
) -> Tuple[bytes, Dict, str]:
    """
    Pipeline modalità strutturata (PROMPT UNIVERSALE ADATTIVO v2.1).
    Ritorna (docx_bytes, stats_dict, filename) — stesso contratto di
    process_zip_and_generate_report(). Non toccare analyze_batch().
    """
    from modules.structured_evidence_parser import (
        parse_structured_response,
        extract_company_name_from_meta
    )
    from modules.structured_evidence_generator import generate_structured_evidence_docx

    # Carica universal prompt (letto da file, non cache — sicuro in thread)
    client = GeminiClient(api_key)
    prompt_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'prompts', 'universal_evidence_prompt.md'
    )
    universal_prompt = client._load_universal_prompt(prompt_path)

    # ── B: Smart batching adattivo ─────────────────────────────────────────
    # create_smart_batches() raggruppa per peso reale in caratteri (First Fit
    # Decreasing). max_files=4 (FIX #2): coerente con cap dinamico fino a 30k
    # per doc chiave (Visura/DVR/Statuto). max_chars=50000 per assorbire
    # batch con 1 doc pesante (es. visura 30k) + 2-3 doc leggeri.
    batches = create_smart_batches(documents, max_files=4, max_chars=50000)
    num_batches = len(batches)
    total_docs_count = len(documents)

    # ── A: Esecuzione parallela batch ──────────────────────────────────────
    # 7 worker allineati al semaforo gemini_structured_slot (cap reale 7).
    # min(7, num_batches) evita thread inutili su ZIP piccoli.
    MAX_STRUCTURED_WORKERS = min(7, num_batches)

    # Preallochiamo la lista con None per preservare l'ordine dei batch
    # indipendentemente dall'ordine di completamento dei thread.
    raw_outputs_ordered = [None] * num_batches
    completed_count = [0]  # lista per mutabilità nel closure

    if status_callback:
        status_callback(f"🤖 Analisi strutturata: {num_batches} batch, {MAX_STRUCTURED_WORKERS} worker paralleli...")

    from concurrent.futures import ThreadPoolExecutor, as_completed as _as_completed

    with ThreadPoolExecutor(max_workers=MAX_STRUCTURED_WORKERS) as executor:
        # Mappa future → indice batch (per preservare l'ordine al merge)
        future_to_idx = {
            executor.submit(
                client.analyze_batch_structured,
                batch, i, total_docs_count, universal_prompt
            ): i
            for i, batch in enumerate(batches)
        }

        # as_completed() viene chiamato dal thread principale → i callback
        # di Streamlit sono sempre chiamati dal thread principale. ✅
        for future in _as_completed(future_to_idx):
            batch_idx = future_to_idx[future]
            completed_count[0] += 1
            try:
                raw = future.result()
                if raw:
                    raw_outputs_ordered[batch_idx] = raw
            except Exception as e:
                # Un batch fallito non blocca gli altri — il parser
                # salterà semplicemente i documenti mancanti.
                print(f"[STRUCTURED] Batch {batch_idx} fallito in parallelo: "
                      f"{str(e)[:200]}")

            if status_callback:
                status_callback(
                    f"🤖 Batch strutturati completati: {completed_count[0]}/{num_batches}..."
                )
            if progress_callback:
                progress_callback(
                    0.30 + (completed_count[0] / num_batches) * 0.55,
                    f"Analisi strutturata: {completed_count[0]}/{num_batches} batch completati..."
                )

    # ── RECOVERY: recupero batch falliti ───────────────────────────────────
    # Identifica i batch che non hanno prodotto output nel loop parallelo.
    failed_batch_indices = [i for i, r in enumerate(raw_outputs_ordered) if r is None]

    if failed_batch_indices:
        recovery_outputs = []  # output recuperati (non devono rispettare ordine)
        if status_callback:
            status_callback(f"⚠️ {len(failed_batch_indices)} batch falliti — tentativo recupero...")

        for batch_idx in failed_batch_indices:
            batch_docs = batches[batch_idx]

            # Livello 1: retry dell'intero batch (può risolvere errori transienti)
            print(f"[RECOVERY L1] Retry batch {batch_idx} ({len(batch_docs)} doc)...")
            raw = client.analyze_batch_structured(
                batch_docs, batch_idx, total_docs_count, universal_prompt
            )
            if raw:
                raw_outputs_ordered[batch_idx] = raw
                print(f"[RECOVERY L1] Batch {batch_idx} recuperato")
                continue

            # Livello 2: splitta in 1-doc-per-volta e riprova ciascuno
            print(f"[RECOVERY L2] Batch {batch_idx} ancora fallito — retry 1-doc-per-volta...")
            for single_doc in batch_docs:
                fname = single_doc.get("filename", "sconosciuto")
                raw_single = client.analyze_batch_structured(
                    [single_doc], batch_idx, total_docs_count, universal_prompt
                )
                if raw_single:
                    recovery_outputs.append(raw_single)
                    print(f"[RECOVERY L2] Recuperato doc singolo: {fname}")
                else:
                    # Livello 3: documento irrecuperabile → failed_files
                    print(f"[RECOVERY L3] Doc irrecuperabile: {fname}")
                    failed_files.append({
                        "filename": fname,
                        "type": "ai_error",
                        "error": "Documento non elaborabile dal modello AI dopo tentativi multipli"
                    })

        # Appendi gli output di recovery alla lista principale (il parser li integra)
        raw_outputs_ordered_final = [r for r in raw_outputs_ordered if r] + recovery_outputs
    else:
        raw_outputs_ordered_final = [r for r in raw_outputs_ordered if r]

    raw_outputs = raw_outputs_ordered_final

    if not raw_outputs:
        raise ValueError(
            "Il modello non ha prodotto output. "
            "Verifica la chiave API e riprova."
        )

    # Parse YAML accumulato da tutti i batch
    full_yaml = "\n\n---\n\n".join(raw_outputs)
    parsed_data = parse_structured_response(full_yaml)

    if not parsed_data:
        raise ValueError(
            "L'analisi non ha prodotto un output interpretabile. "
            "Riprova o contatta il supporto se il problema persiste."
        )

    # Estrai nome azienda: META YAML → fallback visura/ZIP
    company_name = extract_company_name_from_meta(parsed_data)
    if company_name == "AZIENDA NON IDENTIFICATA":
        company_name = extract_company_name(documents, logger, zip_filename=input_name)

    docs_estratti = len(documents)
    docs_vuoti = total_files - docs_estratti

    # Genera docx strutturato
    if status_callback:
        status_callback("📄 Generazione documento strutturato Word...")
    if progress_callback:
        progress_callback(0.95, "Generazione Word strutturato...")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        generate_structured_evidence_docx(
            parsed_data, tmp_path, docs_estratti, docs_vuoti,
            failed_files=failed_files
        )
        with open(tmp_path, 'rb') as f:
            word_bytes = f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    # Stats compatibili con app.py — stesso contratto del narrativo
    stats = {
        "company_name": company_name,
        "failed_files": failed_files,
        "total_docs": total_files,
        "docs_with_text": docs_estratti,
        "total_paragraphs": docs_estratti,
        "paragraphs_generated": docs_estratti,
        "avg_words": 0,          # N/A per modalità strutturata
        "unique_verbs": 0,       # N/A
        "privacy_violations": 0,
        "log_summary": logger.get_summary() if logger else {},
        "output_mode": "strutturato",
    }

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_company = "".join(c for c in company_name if c.isalnum() or c == " ")[:30]
    safe_company = safe_company.replace(" ", "_") if safe_company else "AUDIT"
    filename = f"Evidenze_Strutturate_{safe_company}_{timestamp}.docx"

    if progress_callback:
        progress_callback(1.0, "Completato!")
    if status_callback:
        status_callback("✅ Evidenze strutturate generate con successo!")

    return word_bytes, stats, filename


def process_zip_and_generate_report(
    input_source,
    api_key: str,
    progress_callback=None,
    status_callback=None,
    input_type: str = "zip",
    filename: str = "upload.zip",
) -> Tuple[bytes, Dict, str]:
    """
    Pipeline completa: ZIP → estrazione → analisi AI → Word.

    Args:
        input_source: bytes del file ZIP (da FastAPI), oppure path stringa per cartella locale
        api_key:      Chiave API Gemini
        progress_callback: callable(float, str) opzionale per progress bar
        status_callback:   callable(str) opzionale per messaggi di stato
        input_type:   "zip" (default) o "folder" per cartella locale su disco
        filename:     Nome originale del file ZIP (usato per estrarre nome azienda)

    Returns:
        Tuple (bytes del documento Word, dict stats, nome file suggerito)
    """
    # Inizializza logger
    logger = AuditLogger()
    all_paragraphs = []
    
    # ═══ BENCHMARK TIMING ═══
    import time
    benchmark_start = time.time()
    phase_times = {
        "extraction": 0,
        "text_extraction": 0,
        "ai_analysis": 0,
        "report_generation": 0
    }
    input_name = ""
    total_size_bytes = 0
    
    try:
        # STEP 1: Estrazione file
        phase_start = time.time()
        if status_callback:
            status_callback("📦 Estrazione file ZIP..." if input_type == "zip" else "📂 Scansione cartella locale...")
        if progress_callback:
            progress_callback(0.05, "Estrazione documenti...")

        if input_type == "folder":
            files, temp_dir = extract_from_folder(input_source)
            input_name = os.path.basename(input_source) if input_source else "folder"
        else:
            # ZIP: input_source può essere bytes, BytesIO, o qualsiasi file-like
            files, temp_dir = extract_zip(input_source)
            # Nome: usa il parametro esplicito, poi prova .name sull'oggetto, poi fallback
            input_name = filename or getattr(input_source, 'name', 'upload.zip')

        total_files = len(files)
        
        logger.log_extraction_start(total_files)
        for f in files:
            logger.log_file_extracted(f["filename"], f["category"], f["size"])
        
        if status_callback:
            status_callback(f"📄 {total_files} file trovati")
        
        phase_times["extraction"] = time.time() - phase_start
        
        # STEP 2: Estrazione testo
        phase_start = time.time()
        if status_callback:
            status_callback("📝 Estrazione testo dai documenti...")
        
        def extraction_progress(pct, msg):
            if progress_callback:
                progress_callback(0.05 + pct * 0.25, msg)
        
        # MODIFICATO: ora restituisce tupla (documents, failed_files)
        documents, failed_files = extract_text_from_files_with_logging(
            files, extraction_progress, logger, api_key, base_extract_dir=temp_dir
        )
        
        if status_callback:
            failed_count = len(failed_files)
            if failed_count > 0:
                status_callback(f"✅ {len(documents)} documenti estratti | ⚠️ {failed_count} file non elaborati")
            else:
                status_callback(f"✅ {len(documents)} documenti con testo estratto")
        
        phase_times["text_extraction"] = time.time() - phase_start
        
        if len(documents) == 0:
            logger.log_error("Nessun documento con testo estratto", "text_extraction")
            raise ValueError("Nessun documento valido trovato nel file ZIP")

        # ── STEP 3: pipeline strutturata (PROMPT UNIVERSALE ADATTIVO v2.1) ──
        # Il codice sottostante (STEP 3 narrativo + STEP 4) è DEAD CODE.
        # Mantenuto come fallback di sicurezza — NON cancellare (2026-04-18).
        return _pipeline_strutturata(
            documents, failed_files, api_key, input_name,
            total_files, progress_callback, status_callback, logger
        )

        # ════ DEAD CODE — pipeline narrativa (non attiva dal 2026-04-18) ════
        # STEP 3: Analisi AI in batch PARALLELI
        phase_start = time.time()
        if status_callback:
            status_callback("🤖 Analisi AI in corso (batch paralleli)...")
        
        client = GeminiClient(api_key)
        
        # SMART BATCHING 2.0
        # Invece di raggruppare alla cieca, usa l'algoritmo intelligente
        smart_batches_list = create_smart_batches(documents, max_files=BATCH_SIZE, max_chars=MAX_BATCH_CHARS)
        num_batches = len(smart_batches_list)
        MAX_BATCH_RETRIES = 2
        
        # Prepara tutti i batch con indici
        batches = []
        for i, batch_docs in enumerate(smart_batches_list):
            batches.append((i, batch_docs))
            # Log composizione batch per debug
            total_chars = sum(len(d.get("content", "")) for d in batch_docs)
            # Log composizione batch per debug
            total_chars = sum(len(d.get("content", "")) for d in batch_docs)
            # logger.log_info non esiste, uso print per debug console
            print(f"DEBUG - Batch {i}: {len(batch_docs)} docs, {total_chars} chars")
        
        def process_single_batch(batch_data):
            """Processa un singolo batch - usato per parallelizzazione."""
            batch_idx, batch_docs = batch_data
            doc_names = [d["filename"] for d in batch_docs]
            logger.log_batch_start(batch_idx, doc_names, num_batches)
            
            paragraphs = []
            for retry in range(MAX_BATCH_RETRIES + 1):
                try:
                    paragraphs = client.analyze_batch(batch_docs, batch_idx, len(documents))
                    
                    if paragraphs and len(paragraphs) > 0:
                        logger.log_batch_complete(batch_idx, len(paragraphs))
                        return batch_idx, paragraphs
                    else:
                        if retry < MAX_BATCH_RETRIES:
                            import time
                            time.sleep(2)
                
                except Exception as e:
                    error_msg = f"Errore batch {batch_idx}: {str(e)}"
                    logger.log_batch_error(batch_idx, error_msg)
                    if retry < MAX_BATCH_RETRIES:
                        import time
                        time.sleep(3)
                    else:
                        logger.log_error(f"Batch {batch_idx} fallito dopo {MAX_BATCH_RETRIES} retry", "ai_analysis")
            
            # Se siamo qui, il batch è fallito definitivamente
            error_msg = f"Batch {batch_idx} completato senza risultati (lista vuota)"
            logger.log_batch_error(batch_idx, error_msg)
            return batch_idx, []
        
        # ESECUZIONE PARALLELA con max worker ottimizzati (da config)
        from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
        
        if status_callback:
            status_callback(f"🚀 Esecuzione {num_batches} batch in parallelo (max {MAX_API_WORKERS} thread)...")
        
        batch_results = {}
        # Usa MAX_API_WORKERS per massimizzare il throughput su Flash
        with ThreadPoolExecutor(max_workers=MAX_API_WORKERS) as executor:
            futures = {executor.submit(process_single_batch, b): b[0] for b in batches}
            
            completed = 0
            try:
                # Timeout DINAMICO: 120s base + 30s per batch, cap 45 min
                dynamic_timeout = min(120 + (num_batches * 30), 2700)
                print(f"[TIMEOUT] Dinamico: {dynamic_timeout}s per {num_batches} batch")
                for future in as_completed(futures, timeout=dynamic_timeout):
                    batch_idx = futures[future]
                    try:
                        idx, paragraphs = future.result(timeout=180)  # 3 min max per batch result
                        batch_results[idx] = paragraphs if paragraphs else []
                        completed += 1
                        
                        if progress_callback:
                            batch_progress = 0.30 + (completed / num_batches) * 0.60
                            progress_callback(batch_progress, f"Completati {completed}/{num_batches} batch...")
                    except TimeoutError:
                        logger.log_error(f"Timeout result batch {batch_idx}", "ai_analysis")
                        batch_results[batch_idx] = []
                    except Exception as e:
                        logger.log_error(f"Errore batch {batch_idx}: {e}", "ai_analysis")
                        batch_results[batch_idx] = []
            except TimeoutError:
                logger.log_error(f"Timeout globale esecuzione parallela ({dynamic_timeout}s dinamico)", "ai_analysis")
                # Recupera eventuali risultati parziali o segna falliti quelli mancanti
                for f in futures:
                    idx = futures[f]
                    if idx not in batch_results:
                        batch_results[idx] = []
            except Exception as e:
                logger.log_error(f"Errore critico loop parallelo: {e}", "ai_analysis")
        
        # Unisci risultati nell'ordine corretto
        for batch_idx in sorted(batch_results.keys()):
            if batch_results[batch_idx]:
                all_paragraphs.extend(batch_results[batch_idx])
        
        # ORDINAMENTO PER MACROAREE LOGICHE
        # Mappa priorità macroaree
        macroarea_map = {name: i for i, name in enumerate(MACROAREA_ORDER)}
        
        def sort_key(p):
            cat = p.get("categoria", "ALTRO").upper().strip()
            # Cerca match esatto o parziale
            priority = macroarea_map.get(cat, 999)
            if priority == 999:
                 for k, v in macroarea_map.items():
                      if k in cat:
                           priority = v
                           break
            # Secondaria: Sottotitolo per ordine alfabetico dentro la categoria
            return (priority, p.get("sottotitolo", ""))

        all_paragraphs.sort(key=sort_key)
        
        # RINUMERAZIONE SEQUENZIALE (1..N)
        for i, p in enumerate(all_paragraphs):
            p["numero"] = i + 1  # Sovrascrive il numero temporaneo dei batch
        
        phase_times["ai_analysis"] = time.time() - phase_start
        
        # Verifica copertura
        if len(all_paragraphs) < len(documents):
            logger.log_warning(
                f"Copertura incompleta: {len(all_paragraphs)}/{len(documents)} documenti",
                "ai_analysis"
            )
        
        # STEP 4: Generazione Word
        phase_start = time.time()
        if status_callback:
            status_callback("📄 Generazione documento Word...")
        if progress_callback:
            progress_callback(0.95, "Generazione Word...")
        
        stats = client.get_quality_stats(all_paragraphs)
        stats["total_docs"] = total_files
        stats["docs_with_text"] = len(documents)
        stats["paragraphs_generated"] = len(all_paragraphs)
        stats["log_file"] = logger.log_file
        stats["failed_files"] = failed_files  # NUOVO: lista file non elaborati
        
        # Estrai nome azienda con logica migliorata
        # Estrai nome azienda con logica migliorata (usa anche output AI e nome ZIP)
        company_name = extract_company_name(documents, logger, all_paragraphs, zip_filename=input_name)
        
        word_bytes = generate_report_word(all_paragraphs, stats, company_name)
        
        phase_times["report_generation"] = time.time() - phase_start
        total_processing_time = time.time() - benchmark_start
        
        logger.log_word_generation(True, len(all_paragraphs), company_name, "extracted")
        
        # Log tempo elaborazione (benchmark semplificato)
        print(f"[BENCHMARK] Elaborazione completata: {total_processing_time:.1f}s per {len(documents)} file")
        
        # Genera filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company = "".join(c for c in company_name if c.isalnum() or c == " ")[:30]
        safe_company = safe_company.replace(" ", "_") if safe_company else "AUDIT"
        filename = f"Relazione_Evidenze_{safe_company}_{timestamp}.docx"
        
        if progress_callback:
            progress_callback(1.0, "Completato!")
        if status_callback:
            status_callback("✅ Report generato con successo!")
        
        # Aggiungi summary del logger alle stats
        stats["log_summary"] = logger.get_summary()
        
        # NUOVO: Salva company_name in stats per passarlo a Tab 2
        stats["company_name"] = company_name
        
        logger.complete("success")
        
        return word_bytes, stats, filename
    
    except Exception as e:
        logger.log_error(str(e), "pipeline")
        logger.complete("error")
        raise


def extract_text_from_files_with_logging(files: List[Dict], progress_callback=None, logger=None, api_key: str = None, base_extract_dir: str = None) -> Tuple[List[Dict], List[Dict]]:
    """
    Estrae testo dai file con logging dettagliato.
    Usa PyMuPDF per velocità (5-10x più veloce di PyPDF2).
    
    Args:
        files: Lista file da processare
        progress_callback: Callback per progress bar
        logger: Logger per diagnostica
        api_key: API key per Gemini OCR (fallback per PDF scansionati)
        base_extract_dir: Directory base per calcolare percorsi relativi
    
    Returns:
        Tuple (documenti_estratti, file_falliti)
        - documenti_estratti: lista dict con filename e content
        - file_falliti: lista dict con relative_path, filename, reason, category
    """
    # Importa moduli di estrazione esistenti
    sys.path.insert(0, EXECUTION_DIR)
    try:
        from extract_content import (
            extract_text_from_pdf,
            extract_text_from_docx,
            extract_text_from_txt,
            extract_text_from_doc_legacy,
            extract_text_from_excel,
            extract_text_from_rtf,
            extract_text_from_p7m,
            extract_text_from_msg,
            # Aggiunte 2026-04-19
            extract_text_from_eml,
            extract_text_from_html,
            extract_text_from_xml,
            extract_text_from_odt,
            extract_text_from_ods,
            extract_text_from_pptx,
            # Aggiunte 2026-04-20
            extract_text_from_mpp,
            extract_text_from_cam,
        )
    except ImportError as e:
        st.error(f"Errore import moduli estrazione: {e}")
        if logger:
            logger.log_error(f"Import fallito: {e}", "text_extraction")
        return [], []
    
    # Inizializza Gemini OCR per fallback
    gemini_ocr = None
    if api_key:
        try:
            from modules.gemini_ocr import GeminiOCR
            gemini_ocr = GeminiOCR(api_key)
        except ImportError:
            if logger:
                logger.log_warning("GeminiOCR non disponibile", "ocr")
    
    # Configurazione soglie (hardcoded per ora, ma centralizzabili)
    MIN_CHARS_FOR_OCR_TRIGGER = 15  # Se < 15, prova OCR (era 50)
    MIN_CHARS_VALID_DOC = 3        # Se > 3, accetta (era 10/20)

    documents = []
    failed_files = [] 
    files_needing_ocr = []
    total = len(files) # FIX: Variabile necessaria per progress_callback

    # Estrazione sequenziale (PyMuPDF è già veloce)
    for i, file_info in enumerate(files):
        if progress_callback:
            progress_callback(i / total * 0.7, f"Estrazione {file_info['filename']}...")
        
        filepath = file_info["path"]
        category = file_info["category"]
        content = ""
        error = None
        needs_ocr = False
        original_content = "" # Salva contenuto originale per fallback
        
        try:
            if category == "pdf":
                with contextlib.redirect_stderr(io.StringIO()):
                    result = extract_text_from_pdf(filepath)
                content = result[0] if isinstance(result, tuple) else result
                original_content = content # Backup
                
                # Verifica se il testo è insufficiente (probabilmente scansionato)
                # SOGLIA RILASSATA: Solo se veramente vuoto o brevissimo
                if not content or len(str(content).strip()) < MIN_CHARS_FOR_OCR_TRIGGER:
                    needs_ocr = True
                    # Salva original_content nel file_info per recupero successivo
                    file_info["original_text"] = str(content) if content else ""
                    files_needing_ocr.append(file_info)
                    
            elif category == "word":
                if filepath.endswith(".docx"):
                    result = extract_text_from_docx(filepath)
                else:
                    result = extract_text_from_doc_legacy(filepath)
                content = result[0] if isinstance(result, tuple) else result
                # Fix: Se DOCX vuoto ma file >1KB, probabilmente ha solo immagini → OCR
                if not content or len(str(content).strip()) < MIN_CHARS_FOR_OCR_TRIGGER:
                    file_size = file_info.get("size", 0)
                    if file_size > 1024:  # File >1KB ma senza testo = probabilmente immagini
                        needs_ocr = True
                        file_info["original_text"] = str(content) if content else ""
                        files_needing_ocr.append(file_info)
            elif category == "excel":
                result = extract_text_from_excel(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "text":
                result = extract_text_from_txt(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "image":
                needs_ocr = True
                files_needing_ocr.append(file_info)
            elif category == "rtf":
                result = extract_text_from_rtf(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "p7m":
                result = extract_text_from_p7m(filepath)
                content = result[0] if isinstance(result, tuple) else result
                # Se P7M contiene un PDF scansionato, potrebbe servire OCR
                if not content or len(str(content).strip()) < MIN_CHARS_FOR_OCR_TRIGGER:
                    needs_ocr = True
                    file_info["original_text"] = str(content) if content else ""
                    files_needing_ocr.append(file_info)
            elif category == "email":
                result = extract_text_from_msg(filepath)
                content = result[0] if isinstance(result, tuple) else result
            # --- Handler aggiunti 2026-04-19 ---
            elif category == "eml":
                result = extract_text_from_eml(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "html":
                result = extract_text_from_html(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "xml":
                result = extract_text_from_xml(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "odt":
                result = extract_text_from_odt(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "ods":
                result = extract_text_from_ods(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "pptx":
                result = extract_text_from_pptx(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "heic":
                # Gemini API supporta image/heic come inline data — invia bytes diretti
                if gemini_ocr:
                    try:
                        with open(filepath, "rb") as _hf:
                            heic_bytes = _hf.read()
                        content, _ = gemini_ocr.extract_text_from_image_raw(
                            heic_bytes, "image/heic"
                        )
                    except Exception:
                        content = ""
            elif category == "mpp":
                result = extract_text_from_mpp(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "cam":
                result = extract_text_from_cam(filepath)
                content = result[0] if isinstance(result, tuple) else result
            elif category == "unsupported":
                # Formato noto ma senza testo audit estraibile — segnala all'utente
                relative_path = filepath
                if base_extract_dir and filepath.startswith(base_extract_dir):
                    relative_path = filepath[len(base_extract_dir):].lstrip(os.sep)
                ext_display = os.path.splitext(file_info["filename"])[1].upper()
                failed_files.append({
                    "relative_path": relative_path,
                    "filename": file_info["filename"],
                    "reason": f"Formato {ext_display} non supportato (file tecnico/grafico)",
                    "category": category,
                    "type": "unsupported_format"
                })
                continue  # Non processare oltre
        except Exception as e:
            content = ""
            error = str(e)

        # Aggiungi documento se ha contenuto (e non serve OCR)
        if not needs_ocr and content and isinstance(content, str) and len(content.strip()) >= MIN_CHARS_VALID_DOC:
            # OTTIMIZZAZIONE: Pulisci whitespace per risparmiare token e tempo
            optimized_content = optimize_text_content(content)
            
            documents.append({
                "filename": file_info["filename"],
                "content": optimized_content
            })
            if logger:
                logger.log_text_extraction(file_info["filename"], "success", len(content))
                
        elif not needs_ocr:
            # NUOVO: Traccia file falliti con percorso relativo
            relative_path = filepath
            if base_extract_dir and filepath.startswith(base_extract_dir):
                relative_path = filepath[len(base_extract_dir):].lstrip(os.sep)
            
            reason = f"Testo insufficiente (< {MIN_CHARS_VALID_DOC} chars)"
            if error:
                reason = f"Errore estrazione: {error}"
            
            failed_files.append({
                "relative_path": relative_path,
                "filename": file_info["filename"],
                "reason": reason,
                "category": category,
                "type": "extraction_failed"
            })
            
            if logger:
                logger.log_text_extraction(
                    file_info["filename"],
                    "empty" if not error else "failed",
                    len(content) if content else 0,
                    error
                )
    
    # Seconda passata: OCR PARALLELO per PDF scansionati
    if files_needing_ocr and gemini_ocr:
        if progress_callback:
            progress_callback(0.70, f"OCR parallelo per {len(files_needing_ocr)} PDF scansionati...")
        
        from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
        import threading
        
        ocr_results = {}
        ocr_lock = threading.Lock()
        completed_count = [0]
        OCR_TIMEOUT_PER_FILE = 60  # 60 secondi per gestire file pesanti senza abortire troppo presto
        OCR_GLOBAL_TIMEOUT = 300   # (Non usato nel recovery, solo legacy)
        
        def process_ocr_file(file_info):
            """Processa un singolo file con OCR."""
            try:
                filepath = file_info["path"]
                file_size = file_info.get("size", 0)
                # max_pages adattivo: file grandi ottengono più pagine per coprire contenuto critico
                if file_size > 5_000_000:
                    max_p = 12
                elif file_size > 2_000_000:
                    max_p = 10
                else:
                    max_p = 8
                content, method = gemini_ocr.extract_text_from_pdf(filepath, max_pages=max_p)
                return file_info["filename"], content, None
            except Exception as e:
                return file_info["filename"], "", str(e)
        
        # OCR PARALLELO con MAX_OCR_WORKERS worker e TIMEOUT
        OCR_TIMEOUT_INACTIVITY = 180  # Timeout inattività: 3min per tollerare burst di file grandi senza zombie cascade
        
        with ThreadPoolExecutor(max_workers=MAX_OCR_WORKERS) as executor:
            # Mappa future -> file_info
            future_to_file = {executor.submit(process_ocr_file, f): f for f in files_needing_ocr}
            processed_futures = set()
            
            try:
                # as_completed con timeout: se passa OCR_TIMEOUT_INACTIVITY senza che NESSUNO finisca, lancia TimeoutError
                for future in as_completed(future_to_file, timeout=OCR_TIMEOUT_INACTIVITY):
                    processed_futures.add(future)
                    file_info = future_to_file[future]
                    
                    try:
                        filename, content, error = future.result() # Ritorna subito perché future è completed
                        
                        # === LOGICA STANDARD ===
                        with ocr_lock:
                            completed_count[0] += 1
                            if progress_callback:
                                pct = 0.70 + (completed_count[0] / len(files_needing_ocr)) * 0.25
                                progress_callback(pct, f"OCR completati {completed_count[0]}/{len(files_needing_ocr)}...")
                        
                        original_text = file_info.get("original_text", "")
                        final_content = content
                        used_method = "ocr"
                        
                        if not content or len(content.strip()) < MIN_CHARS_VALID_DOC:
                            if original_text and len(original_text.strip()) >= MIN_CHARS_VALID_DOC:
                                final_content = original_text
                                used_method = "fallback_original"
                                error = None 
                        
                        if final_content and len(final_content.strip()) >= MIN_CHARS_VALID_DOC:
                            optimized_content = optimize_text_content(final_content)
                            documents.append({
                                "filename": filename,
                                "content": optimized_content
                            })
                            if logger:
                                logger.log_text_extraction(filename, "success", len(final_content))
                                if used_method == "ocr":
                                    logger.log_warning(f"Usato Gemini OCR per {filename}", "ocr")
                                else:
                                    logger.log_warning(f"OCR fallito/scarso per {filename}, usato testo originale (fallback)", "ocr")
                        else:
                            filepath = file_info["path"]
                            relative_path = filepath
                            if base_extract_dir and filepath.startswith(base_extract_dir):
                                relative_path = filepath[len(base_extract_dir):].lstrip(os.sep)
                            
                            fail_reason = error or f"Testo insufficiente"
                            failed_files.append({
                                "relative_path": relative_path,
                                "filename": filename,
                                "reason": fail_reason,
                                "category": file_info.get("category", "pdf"),
                                "type": "ocr_failed"
                            })
                            if logger:
                                logger.log_text_extraction(filename, "empty", 0, fail_reason)
                    
                    except Exception as e:
                        if logger: logger.log_error(f"Errore interno loop OCR per {file_info['filename']}: {e}", "ocr")
            
            except TimeoutError:
                if logger:
                    logger.log_warning(f"Rilevato STALLO OCR (> {OCR_TIMEOUT_INACTIVITY}s inattività). Attivo recupero file bloccati...", "ocr_monitor")
            
            # === RECUPERO FILE BLOCCATI (ZOMBIE) - PARALLELO ===
            zombies = []
            for future, file_info in future_to_file.items():
                if future not in processed_futures:
                    zombies.append(file_info)
            
            if zombies:
                if logger: logger.log_warning(f"Rilevati {len(zombies)} file bloccati. Avvio recupero PARALLELO Fast Mode...", "ocr_monitor")
                
                def recover_single_zombie(z_info):
                    # Helper per esecuzione parallela
                    fname = z_info["filename"]
                    fpath = z_info["path"]
                    try:
                        # Fallback Fast Mode
                        c, _ = gemini_ocr.extract_text_from_pdf(fpath, max_pages=3, fast_mode=True)
                        if c and len(c.strip()) >= MIN_CHARS_VALID_DOC:
                            return True, c, None
                        return False, None, "Testo insufficiente in Fast Mode"
                    except Exception as e:
                        return False, None, str(e)

                # Esegui recupero PARALLELO POTENZIATO (Max 4 Worker)
                # "Quad Turbo": smaltisce la coda velocemente ora che non c'è più il rischio di timeout globale.
                with ThreadPoolExecutor(max_workers=4) as recover_executor:
                    future_to_zombie = {
                        recover_executor.submit(recover_single_zombie, z): z
                        for z in zombies
                    }
                    
                    try:
                        # RIMOSSO TIMEOUT GLOBALE: Se ci sono 70+ file, 300s non bastano.
                        # Lasciamo che finisca. Il timeout per singolo file è gestito internamente o non necessario (Fast Mode è veloce).
                        for r_future in as_completed(future_to_zombie):
                            z_info = future_to_zombie[r_future]
                            z_filename = z_info["filename"]
                            
                            try:
                                success, content, err_msg = r_future.result()
                                
                                if success:
                                    optimized_content = optimize_text_content(content)
                                    documents.append({
                                        "filename": z_filename,
                                        "content": optimized_content
                                    })
                                    if logger:
                                        logger.log_text_extraction(z_filename, "success", len(content))
                                        logger.log_warning(f"RECUPERATO {z_filename} con OCR FAST MODE", "ocr_fallback")
                                else:
                                    raise Exception(err_msg)
                                    
                            except Exception as final_err:
                                # Fallimento definitivo recovery per questo file
                                filepath = z_info["path"]
                                relative_path = filepath
                                if base_extract_dir and filepath.startswith(base_extract_dir):
                                    relative_path = filepath[len(base_extract_dir):].lstrip(os.sep)
                                
                                failed_files.append({
                                    "relative_path": relative_path,
                                    "filename": z_filename,
                                    "reason": f"Recovery Fail: {final_err}",
                                    "category": z_info.get("category", "pdf"),
                                    "type": "ocr_recovery_failed"
                                })
                                if logger:
                                    logger.log_error(f"Errore Recovery OCR {z_filename}: {final_err}", "ocr")
                            
                            # Aggiorna contatore globale e UI
                            with ocr_lock:
                                completed_count[0] += 1
                            
                            # Calcolo Float Progress Corretto
                            # Usa completed_count[0] aggiornato
                            final_pct = 0.70 + (completed_count[0] / len(files_needing_ocr)) * 0.25
                            
                            if progress_callback:
                                msg = f"Recuperato: {z_filename[:20]}... ({completed_count[0]}/{len(files_needing_ocr)})"
                                progress_callback(final_pct, msg)
                                
                    except KeyboardInterrupt:
                        # Gestione interruzione manuale sicura
                        if logger: 
                            logger.log_error("Interruzione utente durante Recovery.", "ocr_interrupt")
                        raise
                            

    elif files_needing_ocr and not gemini_ocr:
        # OCR non disponibile - PROVA FALLBACK SU ORIGINALE prima di fallire
        for file_info in files_needing_ocr:
            original_text = file_info.get("original_text", "")
            
            if original_text and len(original_text.strip()) >= MIN_CHARS_VALID_DOC:
                 documents.append({
                    "filename": file_info["filename"],
                    "content": optimize_text_content(original_text)
                })
                 if logger:
                    logger.log_text_extraction(file_info["filename"], "success", len(original_text))
                    logger.log_warning(f"OCR non disp., usato testo originale (limitato) per {file_info['filename']}", "fallback")
            else:
                filepath = file_info["path"]
                relative_path = filepath
                if base_extract_dir and filepath.startswith(base_extract_dir):
                    relative_path = filepath[len(base_extract_dir):].lstrip(os.sep)
                
                failed_files.append({
                    "relative_path": relative_path,
                    "filename": file_info["filename"],
                    "reason": "OCR non disponibile e testo originale insufficiente",
                    "category": file_info.get("category", "pdf"),
                    "type": "ocr_unavailable"
                })
                if logger:
                    logger.log_text_extraction(
                        file_info["filename"], 
                        "empty", 
                        0, 
                        "PDF scansionato - OCR non disponibile"
                    )
    
    if progress_callback:
        progress_callback(1.0, "Estrazione completata!")
    
    return documents, failed_files


def extract_company_name(documents: List[Dict], logger=None, paragraphs: List[Dict] = None, zip_filename: str = "") -> str:
    """
    Estrae il nome dell'azienda con la seguente priorità:
    1. VISURA (documento ufficiale Camera di Commercio) - MASSIMA AFFIDABILITÀ
    2. Campo ente_auditato dall'AI, con filtri per escludere subappaltatori
    3. Regex su contenuti AI
    
    Supporta TUTTE le forme giuridiche italiane:
    - S.R.L., S.P.A., S.N.C., S.A.S.
    - COOPERATIVA SOCIALE, SOC. COOP., CONSORZIO
    - ONLUS, IMPRESA SOCIALE
    """
    company_name = "AZIENDA NON IDENTIFICATA"
    visura_name = None
    
    # Regex universale per forme giuridiche italiane
    FORME_GIURIDICHE = (
        r"(?:"
        r"S\.?R\.?L\.?|S\.?P\.?A\.?|S\.?N\.?C\.?|S\.?A\.?S\.?"  # Società di capitali/persone
        r"|SOCIETA'?\s+(?:PER\s+AZIONI|A\s+RESPONSABILITA'?\s+LIMITATA)"  # Forme estese
        r"|(?:SOC\.?\s*)?COOP(?:ERATIVA)?(?:\s+SOCIALE)?"  # Cooperative
        r"|COOPERATIVA\s+SOCIALE"  # Cooperative sociali
        r"|CONSORZIO"  # Consorzi
        r"|O\.?N\.?L\.?U\.?S\.?"  # ONLUS
        r"|IMPRESA\s+SOCIALE"  # Imprese sociali
        r"|SOC\.?\s*CONSORTILE"  # Società consortili
        r")"
    )

    # ═══════════════════════════════════════════════════════════════════
    # PRIORITÀ 0: NOME FILE ZIP (Candidato Utente)
    # ═══════════════════════════════════════════════════════════════════
    zip_candidate = None
    if zip_filename:
        try:
            import re
            base = os.path.splitext(os.path.basename(zip_filename))[0]
            clean = base.replace('_', ' ').replace('-', ' ')
            clean = re.sub(r'(?i)\b(audit|report|relazione|doc|files|documenti|\d{8}|\d{6})\b', '', clean).strip()
            clean = re.sub(r'\s+', ' ', clean)
            
            if len(clean) > 3 and "UPLOAD" not in clean.upper() and "CARTELLA" not in clean.upper():
                if clean.lower() not in NOMI_GENERICI_ZIP:
                    zip_candidate = clean.upper()
                    print(f"DEBUG: Candidato nome da ZIP: {zip_candidate}")
                else:
                    print(f"DEBUG: Nome ZIP generico ignorato: '{clean}' — attesa visura")
        except Exception as z_err:
             print(f"Errore parsing nome zip: {z_err}")
    
    try:
        import re
        from collections import Counter
        
        # ═══════════════════════════════════════════════════════════════════
        # PRIORITÀ 1: VISURA CAMERALE (Documento Ufficiale - Max Affidabilità)
        # ═══════════════════════════════════════════════════════════════════
        for doc in documents:
            filename_lower = doc["filename"].lower()
            if "visura" in filename_lower:
                content = doc.get("content", "")[:8000]  # Primi 8000 caratteri
                
                # Pattern specifici per visura camerale italiana (ordine di priorità)
                visura_patterns = [
                    # Pattern 1: "VISURA ORDINARIA SOCIETA' DI CAPITALE relativa a NOME"
                    r"VISURA\s+ORDINARIA\s+SOCIET[A']'?\s+DI\s+CAPITALE\s+(?:relativa?\s+a|ente)\s+([A-Z][A-Z0-9\s\.'\-&]+?" + FORME_GIURIDICHE + r")",
                    
                    # Pattern 2: Denominazione seguita da nome
                    r"(?:Denominazione|Ragione\s+sociale)[:\s]+([A-Z][A-Z0-9\s\.'\-&]+?" + FORME_GIURIDICHE + r")",
                    
                    # Pattern 3: Nome con COOPERATIVA SOCIALE (cattura tutto il blocco)
                    r"([A-Z][A-Z0-9\s\.'\-&]*?COOPERATIVA\s+SOCIALE[A-Z0-9\s\.'\-&]*)",
                    
                    # Pattern 4: Nome con CONSORZIO (cattura tutto il blocco)
                    r"(CONSORZIO\s+[A-Z][A-Z0-9\s\.'\-&]+)",
                    
                    # Pattern 5: Nome con SOC. COOP.
                    r"([A-Z][A-Z0-9\s\.'\-&]+SOC\.?\s*COOP[A-Z0-9\s\.'\-&]*)",
                    
                    # Pattern 6: "denominata anche" per acronimi ufficiali
                    r"denominata?\s+anche\s+([A-Z][A-Z0-9\.\s\-&]+(?:S\.?P\.?A\.?|S\.?R\.?L\.?))",
                    
                    # Pattern 7: Fallback generico per S.R.L., S.P.A., etc.
                    r"([A-Z][A-Z0-9\s\.'\-&]+(?:S\.?R\.?L\.?|S\.?P\.?A\.?|S\.?N\.?C\.?|S\.?A\.?S\.?))",
                ]
                
                for pattern in visura_patterns:
                    match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
                    if match:
                        candidate = match.group(1).strip().upper()
                        # Pulizia aggressiva
                        candidate = re.sub(r'\s+', ' ', candidate)  # Rimuovi spazi multipli
                        candidate = re.sub(r'^[^A-Z]+', '', candidate)  # Rimuovi prefissi non-alfa
                        candidate = candidate.strip()
                        
                        # Validazione: almeno 5 caratteri, max 100, deve contenere lettere
                        if len(candidate) > 5 and len(candidate) < 100 and re.search(r'[A-Z]', candidate):
                            visura_name = candidate
                            print(f"DEBUG: Nome azienda da VISURA (pattern matched): {visura_name}")
                            break
                
                if visura_name:
                    break
        
        # Se trovato dalla visura, usa quello!
        # ═══════════════════════════════════════════════════════════════════
        # CROSS-VALIDATION: ZIP vs VISURA
        # ═══════════════════════════════════════════════════════════════════
        if visura_name:
            if zip_candidate:
                # Clean strings for comparison
                v_clean = re.sub(r'[^\w\s]', '', visura_name.upper())
                z_clean = re.sub(r'[^\w\s]', '', zip_candidate.upper())
                
                v_tokens = set(v_clean.split())
                z_tokens = set(z_clean.split())
                
                # Intersezione token significativi (>3 chars per evitare 'SPA', 'SRL', 'DEL', 'SOC')
                common = {w for w in v_tokens.intersection(z_tokens) if len(w) > 3}
                
                # Se c'è overlap, assumiamo sia la stessa azienda -> Usa Visura (Formale)
                if len(common) > 0:
                     print(f"DEBUG: Match ZIP/Visura su '{common}'. Uso Visura: {visura_name}")
                     return visura_name
                else:
                     # Mismatch Totale! (Es. "Impianti AR" vs "Consorzio Jeocon")
                     # Regola Utente: "Confronta... validazione".
                     # Se c'è mismatch, l'OCR Visura potrebbe aver preso il file sbagliato.
                     # Diamo fiducia all'Utente (ZIP) ma logghiamo warning.
                     
                     # Eccezione: se ZIP è molto generico ("AUDIT 2026"), usa Visura
                     if len(zip_candidate) < 5:
                         return visura_name
                         
                     print(f"WARNING: Conflitto Azienda! ZIP='{zip_candidate}' vs VISURA='{visura_name}'. Vincono i token ZIP.")
                     return zip_candidate
            else:
                return visura_name

        # Se non c'è Visura ma c'è ZIP, usa ZIP
        if zip_candidate:
             print(f"DEBUG: Usa nome da ZIP (Nessuna Visura trovata): {zip_candidate}")
             return zip_candidate
        
        # ═══════════════════════════════════════════════════════════════════
        # PRIORITÀ 2: Campo ente_auditato dall'AI (con filtri estesi)
        # ═══════════════════════════════════════════════════════════════════
        if paragraphs:
            candidates = []
            for p in paragraphs:
                if "ente_auditato" in p and p["ente_auditato"]:
                    name = p["ente_auditato"].strip().upper()
                    # Filtri estesi per tutte le forme giuridiche
                    if len(name) > 3 and re.search(FORME_GIURIDICHE, name, re.IGNORECASE):
                        candidates.append(name)
            
            if candidates:
                # Voto a maggioranza
                most_common = Counter(candidates).most_common(3)
                
                if most_common:
                    company_name = most_common[0][0]
                    print(f"DEBUG: Azienda da AI metadata: {company_name} ({most_common[0][1]} voti)")
                    return company_name

        # ═══════════════════════════════════════════════════════════════════
        # PRIORITÀ 3: Regex su contenuti AI (cercando Visura Camerale nei testi)
        # ═══════════════════════════════════════════════════════════════════
        if paragraphs:
            ai_patterns = [
                r"VISURA\s+(?:ORDINARIA|CAMERALE)[^:]*(?:relativa?\s+a|ente|di)\s+([A-Z][A-Z0-9\s\.'\-&]+?" + FORME_GIURIDICHE + r")",
                r"denominata?\s+anche\s+([A-Z][A-Z0-9\.\s]+(?:S\.?P\.?A\.?|S\.?R\.?L\.?))",
                r"(CONSORZIO\s+[A-Z][A-Z0-9\s\.'\-&]+)",
            ]
            candidates = []
            for p in paragraphs:
                content = p.get("contenuto", "")[:800]
                sottotitolo = p.get("sottotitolo", "")
                text_to_search = f"{sottotitolo} {content}"
                
                for pattern in ai_patterns:
                    match = re.search(pattern, text_to_search, re.IGNORECASE)
                    if match:
                        candidate = match.group(1).strip().upper()
                        candidate = re.sub(r'\s+', ' ', candidate)
                        if len(candidate) > 3:
                            candidates.append(candidate)
            
            if candidates:
                most_common = Counter(candidates).most_common(1)
                if most_common:
                    return most_common[0][0]

        # Log finale se nulla trovato
        if logger:
            logger.log_warning("Nome azienda non rilevato", "company_extraction")

        return company_name

    except Exception as e:
        if logger:
            logger.log_error(f"Errore estrazione nome azienda: {e}", "company_extraction")
        print(f"Errore estrazione nome azienda: {e}")
        return "AZIENDA NON IDENTIFICATA"


