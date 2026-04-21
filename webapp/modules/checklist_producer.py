# ==============================================================================
# CHECKLIST PRODUCER - Generazione JSON evidenze da Report Word
# ==============================================================================
# Tab 2: Riceve Report Word, chiama Gemini 3 Pro, produce JSON strutturato

import os
import sys
import json
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from io import BytesIO

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (
    CHECKLIST_PROMPTS,
    PROMPTS_CHECKLIST_DIR,
    GEMINI_MODEL_CHECKLIST,
    SUPPORTED_NORMS
)
from modules.genai_factory import create_genai_client

# ═══ SCHEMI CLAUSOLE PER NORMA ═══
# Definisce le chiavi delle clausole per ogni norma (usato per response_schema)
CHECKLIST_SCHEMAS_DIR = os.path.join(os.path.dirname(PROMPTS_CHECKLIST_DIR), "prompts", "checklist_schemas")

# Mappatura statica delle clausole per ogni norma (estratte dai prompt)
NORM_CLAUSES = {
    "ISO 9001": [
        "iso9001_4_1", "iso9001_4_2", "iso9001_4_3", "iso9001_4_4",
        "iso9001_5_1", "iso9001_5_1_1", "iso9001_5_1_2", "iso9001_5_2", "iso9001_5_2_1", "iso9001_5_2_2", "iso9001_5_3",
        "iso9001_6_1", "iso9001_6_2", "iso9001_6_2_1", "iso9001_6_2_2", "iso9001_6_3",
        "iso9001_7_1", "iso9001_7_1_1", "iso9001_7_1_2", "iso9001_7_1_3", "iso9001_7_1_4", "iso9001_7_1_5", "iso9001_7_1_5_1", "iso9001_7_1_5_2", "iso9001_7_1_6",
        "iso9001_7_2", "iso9001_7_3", "iso9001_7_4", "iso9001_7_5", "iso9001_7_5_1", "iso9001_7_5_2", "iso9001_7_5_3",
        "iso9001_8_1", "iso9001_8_2", "iso9001_8_2_1", "iso9001_8_2_2", "iso9001_8_2_3", "iso9001_8_2_3_1", "iso9001_8_2_3_2", "iso9001_8_2_4",
        "iso9001_8_3", "iso9001_8_3_1", "iso9001_8_3_2", "iso9001_8_3_3", "iso9001_8_3_4", "iso9001_8_3_5", "iso9001_8_3_6",
        "iso9001_8_4", "iso9001_8_4_1", "iso9001_8_4_2", "iso9001_8_4_3",
        "iso9001_8_5", "iso9001_8_5_1", "iso9001_8_5_2", "iso9001_8_5_3", "iso9001_8_5_4", "iso9001_8_5_5", "iso9001_8_5_6",
        "iso9001_8_6", "iso9001_8_7", "iso9001_8_7_1", "iso9001_8_7_2",
        "iso9001_9_1", "iso9001_9_1_1", "iso9001_9_1_2", "iso9001_9_1_3", "iso9001_9_2", "iso9001_9_2_1", "iso9001_9_2_2", "iso9001_9_3", "iso9001_9_3_1", "iso9001_9_3_2", "iso9001_9_3_3",
        "iso9001_10_1", "iso9001_10_2", "iso9001_10_3"
    ],
    "ISO 14001": [
        "iso14001_4_1", "iso14001_4_2", "iso14001_4_3", "iso14001_4_4",
        "iso14001_5_1", "iso14001_5_2", "iso14001_5_3",
        "iso14001_6_1", "iso14001_6_1_1", "iso14001_6_1_2", "iso14001_6_1_3", "iso14001_6_1_4", "iso14001_6_2", "iso14001_6_2_1", "iso14001_6_2_2",
        "iso14001_7_1", "iso14001_7_2", "iso14001_7_3", "iso14001_7_4", "iso14001_7_4_1", "iso14001_7_4_2", "iso14001_7_4_3", "iso14001_7_5", "iso14001_7_5_1", "iso14001_7_5_2", "iso14001_7_5_3",
        "iso14001_8_1", "iso14001_8_2",
        "iso14001_9_1", "iso14001_9_1_1", "iso14001_9_1_2", "iso14001_9_2", "iso14001_9_2_1", "iso14001_9_2_2", "iso14001_9_3",
        "iso14001_10_1", "iso14001_10_2", "iso14001_10_3"
    ],
    "ISO 45001": [
        "iso45001_4_1", "iso45001_4_2", "iso45001_4_3", "iso45001_4_4",
        "iso45001_5_1", "iso45001_5_2", "iso45001_5_3", "iso45001_5_4",
        "iso45001_6_1_1", "iso45001_6_1_2", "iso45001_6_1_2_1", "iso45001_6_1_2_2", "iso45001_6_1_2_3", "iso45001_6_1_3", "iso45001_6_1_4", "iso45001_6_2", "iso45001_6_2_1", "iso45001_6_2_2",
        "iso45001_7_1", "iso45001_7_2", "iso45001_7_3", "iso45001_7_4_1", "iso45001_7_4_2", "iso45001_7_4_3", "iso45001_7_5_1", "iso45001_7_5_2", "iso45001_7_5_3",
        "iso45001_8_1_1", "iso45001_8_1_2", "iso45001_8_1_3", "iso45001_8_1_4", "iso45001_8_2",
        "iso45001_9_1_1", "iso45001_9_1_2", "iso45001_9_2", "iso45001_9_3",
        "iso45001_10_1", "iso45001_10_2", "iso45001_10_3"
    ],
    "ISO 14064": [
        "iso14064_5_1", "iso14064_5_2", "iso14064_5_2_2", "iso14064_5_2_3", "iso14064_5_2_4",
        "iso14064_6", "iso14064_6_1", "iso14064_6_2", "iso14064_6_3", "iso14064_6_4", "iso14064_6_4_1", "iso14064_6_4_2",
        "iso14064_7_1",
        "iso14064_8_1_1", "iso14064_8_1_2", "iso14064_8_2",
        "iso14064_9_1", "iso14064_9_2", "iso14064_9_3_1"
    ],
    "ESG": [
        "easi_1_1", "easi_1_2", "easi_1_3", "easi_1_4", "easi_1_5", "easi_1_6", "easi_1_7", "easi_1_8",
        "easi_2_1", "easi_2_2", "easi_2_3", "easi_2_4", "easi_2_5", "easi_2_6", "easi_2_7", "easi_2_8",
        "easi_3_1", "easi_3_2", "easi_3_3", "easi_3_4", "easi_3_5", "easi_3_6",
        "easi_4_1", "easi_4_2", "easi_4_3", "easi_4_4", "easi_4_5", "easi_4_6", "easi_4_7", "easi_4_8", "easi_4_9", "easi_4_10"
    ],
    "PAS 24000": [
        "pas24000_4_1", "pas24000_4_2", "pas24000_4_3", "pas24000_4_4",
        "pas24000_5_1", "pas24000_5_2", "pas24000_5_3", "pas24000_5_4",
        "pas24000_6_1_1", "pas24000_6_1_2", "pas24000_6_2_1", "pas24000_6_2_2", "pas24000_6_3",
        "pas24000_7_1", "pas24000_7_2", "pas24000_7_3", "pas24000_7_4_1", "pas24000_7_4_2", "pas24000_7_4_3", "pas24000_7_5_2", "pas24000_7_5_3",
        "pas24000_8_1_1", "pas24000_8_1_2", "pas24000_8_1_3", "pas24000_8_2", "pas24000_8_5_6", "pas24000_8_7",
        "pas24000_9_1_1", "pas24000_9_1_2", "pas24000_9_2_1", "pas24000_9_2_2", "pas24000_9_3_1", "pas24000_9_3_2", "pas24000_9_3_3",
        "pas24000_10_1", "pas24000_10_2"
    ],
    "ISO 27001": [
        "iso27001_4_1", "iso27001_4_2", "iso27001_4_3", "iso27001_4_4",
        "iso27001_5_1", "iso27001_5_2", "iso27001_5_3",
        "iso27001_6_1", "iso27001_6_1_1", "iso27001_6_1_2", "iso27001_6_1_3", "iso27001_6_2", "iso27001_6_3",
        "iso27001_7_1", "iso27001_7_2", "iso27001_7_3", "iso27001_7_4", "iso27001_7_5", "iso27001_7_5_1", "iso27001_7_5_2", "iso27001_7_5_3",
        "iso27001_8_1", "iso27001_8_2", "iso27001_8_3",
        "iso27001_9_1", "iso27001_9_2", "iso27001_9_3",
        "iso27001_10_1", "iso27001_10_2"
    ],
    # Alias per varianti ISO/IEC 27001 (stesso contenuto, formato diverso nel JSON)
    "ISO/IEC 27001": [
        "iso27001_4_1", "iso27001_4_2", "iso27001_4_3", "iso27001_4_4",
        "iso27001_5_1", "iso27001_5_2", "iso27001_5_3",
        "iso27001_6_1", "iso27001_6_1_1", "iso27001_6_1_2", "iso27001_6_1_3", "iso27001_6_2", "iso27001_6_3",
        "iso27001_7_1", "iso27001_7_2", "iso27001_7_3", "iso27001_7_4", "iso27001_7_5", "iso27001_7_5_1", "iso27001_7_5_2", "iso27001_7_5_3",
        "iso27001_8_1", "iso27001_8_2", "iso27001_8_3",
        "iso27001_9_1", "iso27001_9_2", "iso27001_9_3",
        "iso27001_10_1", "iso27001_10_2"
    ],
    "ISO/IEC 27001:2022": [
        "iso27001_4_1", "iso27001_4_2", "iso27001_4_3", "iso27001_4_4",
        "iso27001_5_1", "iso27001_5_2", "iso27001_5_3",
        "iso27001_6_1", "iso27001_6_1_1", "iso27001_6_1_2", "iso27001_6_1_3", "iso27001_6_2", "iso27001_6_3",
        "iso27001_7_1", "iso27001_7_2", "iso27001_7_3", "iso27001_7_4", "iso27001_7_5", "iso27001_7_5_1", "iso27001_7_5_2", "iso27001_7_5_3",
        "iso27001_8_1", "iso27001_8_2", "iso27001_8_3",
        "iso27001_9_1", "iso27001_9_2", "iso27001_9_3",
        "iso27001_10_1", "iso27001_10_2"
    ],
    "ISO 37001": [
        "iso37001_4_1", "iso37001_4_2", "iso37001_4_3", "iso37001_4_4", "iso37001_4_5",
        "iso37001_5_1", "iso37001_5_2", "iso37001_5_3",
        "iso37001_6_1", "iso37001_6_2",
        "iso37001_7_1", "iso37001_7_2", "iso37001_7_3", "iso37001_7_4", "iso37001_7_5",
        "iso37001_8_1", "iso37001_8_2", "iso37001_8_3", "iso37001_8_4", "iso37001_8_5", "iso37001_8_6", "iso37001_8_7", "iso37001_8_8", "iso37001_8_9", "iso37001_8_10",
        "iso37001_9_1", "iso37001_9_2", "iso37001_9_3", "iso37001_9_4",
        "iso37001_10_1", "iso37001_10_2"
    ],
    "ISO 39001": [
        "iso39001_4_1", "iso39001_4_2", "iso39001_4_3", "iso39001_4_4",
        "iso39001_5_1", "iso39001_5_2", "iso39001_5_3",
        "iso39001_6_1", "iso39001_6_2", "iso39001_6_3", "iso39001_6_4",
        "iso39001_7_1", "iso39001_7_2", "iso39001_7_3", "iso39001_7_4", "iso39001_7_5",
        "iso39001_7_6_1", "iso39001_7_6_2", "iso39001_7_6_3",
        "iso39001_8_1", "iso39001_8_2",
        "iso39001_9_1", "iso39001_9_2", "iso39001_9_3", "iso39001_9_4",
        "iso39001_10_1", "iso39001_10_2"
    ],
    "ISO 50001": [
        "iso50001_4_1", "iso50001_4_2", "iso50001_4_3", "iso50001_4_4",
        "iso50001_5_1", "iso50001_5_2", "iso50001_5_3",
        "iso50001_6_1", "iso50001_6_2", "iso50001_6_3", "iso50001_6_4", "iso50001_6_5", "iso50001_6_6",
        "iso50001_7_1", "iso50001_7_2", "iso50001_7_3", "iso50001_7_4", "iso50001_7_5",
        "iso50001_8_1", "iso50001_8_2", "iso50001_8_3",
        "iso50001_9_1", "iso50001_9_2", "iso50001_9_3",
        "iso50001_10_1", "iso50001_10_2"
    ],
    # Alias per varianti ISO 50001 (stesso prompt, template diverso)
    "ISO 50001 ESQ": [
        "iso50001_4_1", "iso50001_4_2", "iso50001_4_3", "iso50001_4_4",
        "iso50001_5_1", "iso50001_5_2", "iso50001_5_3",
        "iso50001_6_1", "iso50001_6_2", "iso50001_6_3", "iso50001_6_4", "iso50001_6_5", "iso50001_6_6",
        "iso50001_7_1", "iso50001_7_2", "iso50001_7_3", "iso50001_7_4", "iso50001_7_5",
        "iso50001_8_1", "iso50001_8_2", "iso50001_8_3",
        "iso50001_9_1", "iso50001_9_2", "iso50001_9_3",
        "iso50001_10_1", "iso50001_10_2"
    ],
    "ISO 50001 CERTIS": [
        "iso50001_4_1", "iso50001_4_2", "iso50001_4_3", "iso50001_4_4",
        "iso50001_5_1", "iso50001_5_2", "iso50001_5_3",
        "iso50001_6_1", "iso50001_6_2", "iso50001_6_3", "iso50001_6_4", "iso50001_6_5", "iso50001_6_6",
        "iso50001_7_1", "iso50001_7_2", "iso50001_7_3", "iso50001_7_4", "iso50001_7_5",
        "iso50001_8_1", "iso50001_8_2", "iso50001_8_3",
        "iso50001_9_1", "iso50001_9_2", "iso50001_9_3",
        "iso50001_10_1", "iso50001_10_2"
    ]
}

# Soglia minima parole per clausola (fallback globale)
MIN_CLAUSE_WORDS = 80

# Soglia minima parole PER NORMA - sotto questa soglia = ERRORE CRITICO (richiede recovery)
# Soglia = 100 parole per tutte le norme (clausole < 100 = troppo corte)
MIN_CLAUSE_WORDS_PER_NORM = {
    "ISO 9001": 100,
    "ISO 14001": 100,
    "ISO 45001": 100,
    "ISO 14064": 100,
    "ESG": 100,
    "PAS 24000": 100,
    "ISO 27001": 100,
    "ISO 37001": 100,
    "ISO 39001": 100,
    "ISO 50001": 100,
    "ISO 50001 ESQ": 100,
    "ISO 50001 CERTIS": 100,
    "ISO/IEC 27001": 100,
    "ISO/IEC 27001:2022": 100,
}

# Soglia WARNING parole PER NORMA - tra MIN e WARNING = warning giallo (integrate ma segnalate)
# Clausole >= WARNING = complete
WARNING_CLAUSE_WORDS_PER_NORM = {
    "ISO 9001": 150,
    "ISO 14001": 150,
    "ISO 45001": 150,
    "ISO 14064": 150,
    "ESG": 150,
    "PAS 24000": 150,
    "ISO 27001": 175,     # Prompt: 250-500 → warning sotto 175
    "ISO 37001": 175,
    "ISO 39001": 175,
    "ISO 50001": 175,
    "ISO 50001 ESQ": 175,
    "ISO 50001 CERTIS": 175,
    "ISO/IEC 27001": 175,
    "ISO/IEC 27001:2022": 175,
}


def clean_company_name(name: str) -> str:
    """
    Rimuove la forma giuridica dal nome azienda.
    
    Args:
        name: Nome azienda con possibile forma giuridica
    
    Returns:
        Nome azienda pulito senza forma giuridica
    """
    if not name:
        return name
    
    import re
    
    # Pattern per forme giuridiche (case insensitive, alla fine del nome)
    patterns = [
        r'\s+S\.?r\.?l\.?\s*$',           # S.r.l., Srl, S.R.L.
        r'\s+S\.?p\.?A\.?\s*$',           # S.p.A., SPA, Spa
        r'\s+S\.?a\.?s\.?\s*$',           # S.a.s., Sas
        r'\s+S\.?n\.?c\.?\s*$',           # S.n.c., Snc
        r'\s+S\.?c\.?a\.?r\.?l\.?\s*$',   # S.c.a.r.l., Scarl
        r'\s+S\.?c\.?r\.?l\.?\s*$',       # S.c.r.l., Scrl
        r'\s+Scarl\s*$',
        r'\s+società\s+cooperativa\s*$',
        r'\s+cooperativa\s*$',
        r'\s+Ltd\.?\s*$',
        r'\s+LLC\s*$',
        r'\s+Inc\.?\s*$',
        r'\s+GmbH\s*$',
        r'\s+S\.?r\.?l\.?s\.?\s*$',       # Srls
    ]
    
    cleaned = name
    for pattern in patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    
    return cleaned.strip()


def clean_legal_suffix_from_text(text: str, company_name: str) -> str:
    """
    Rimuove la forma giuridica dal nome azienda ovunque appaia nel testo.
    
    Cerca il nome azienda con la forma giuridica e lo sostituisce con il nome pulito.
    
    Args:
        text: Testo da pulire (es. contenuto clausola)
        company_name: Nome azienda originale (con possibile forma giuridica)
    
    Returns:
        Testo con nome azienda senza forma giuridica
    """
    if not text or not company_name:
        return text
    
    import re
    
    # Estrai nome pulito
    clean_name = clean_company_name(company_name)
    
    # Se il nome è già pulito, niente da fare
    if clean_name == company_name:
        return text
    
    # Sostituisci tutte le occorrenze del nome completo con quello pulito
    # Usa word boundary per evitare sostituzioni parziali
    pattern = re.escape(company_name)
    text = re.sub(pattern, clean_name, text, flags=re.IGNORECASE)
    
    return text


def clean_legal_suffix_from_json(json_data: Dict) -> Dict:
    """
    Pulisce la forma giuridica dal campo 'azienda' e da tutte le clausole.
    
    Args:
        json_data: JSON con campo 'azienda' e 'clausole'
    
    Returns:
        JSON con nomi puliti
    """
    if not json_data:
        return json_data
    
    # Pulisci campo azienda
    original_name = json_data.get("azienda", "")
    clean_name = clean_company_name(original_name)
    json_data["azienda"] = clean_name
    
    # Pulisci tutte le clausole (sostituisci nome completo con nome pulito)
    if original_name and original_name != clean_name:
        clausole = json_data.get("clausole", {})
        for key, value in clausole.items():
            if value and isinstance(value, str):
                clausole[key] = clean_legal_suffix_from_text(value, original_name)
    
    return json_data


def extract_company_from_clauses(clausole: Dict, api_key: str) -> str:
    """
    Estrae il nome azienda COMPLETO dal contenuto delle clausole generate.
    
    Usa SOLO Gemini per garantire estrazione accurata del nome completo.
    Chiamata come FALLBACK quando il nome dal report è vuoto/invalido.
    
    Args:
        clausole: Dizionario delle clausole generate
        api_key: Chiave API Gemini
    
    Returns:
        Nome azienda estratto (pulito, senza forma giuridica) o stringa vuota
    """
    from google import genai
    from google.genai import types
    
    if not clausole:
        return ""
    
    # Prendi il testo delle prime clausole (4.x e 5.x hanno il nome azienda)
    priority_keys = [k for k in clausole.keys() if '_4_' in k or '_5_' in k]
    priority_keys = sorted(priority_keys)[:4]
    
    if not priority_keys:
        priority_keys = list(clausole.keys())[:4]
    
    combined_text = ""
    for key in priority_keys:
        if clausole.get(key):
            combined_text += f"\n{clausole[key][:2500]}"
        if len(combined_text) > 8000:
            break
    
    if not combined_text.strip():
        return ""
    
    try:
        client = create_genai_client(api_key)
        
        prompt = f"""COMPITO: Estrai il NOME COMPLETO dell'azienda auditata dal testo seguente.

TESTO DELLE CLAUSOLE DI AUDIT:
{combined_text[:8000]}

ISTRUZIONI FONDAMENTALI:
1. L'azienda auditata è il SOGGETTO principale del testo (l'organizzazione valutata)
2. Restituisci il NOME COMMERCIALE COMPLETO con TUTTE LE PAROLE
3. NON troncare il nome - se è "Emmeci Appalti" restituisci "Emmeci Appalti", NON "Emmeci" o "Emm"
4. ESCLUDI la forma giuridica (S.r.l., S.p.A., Srl, ecc.) ma MANTIENI tutte le altre parole
5. NON restituire nomi di enti certificatori, auditor, consulenti

ESEMPI DI OUTPUT CORRETTO:
- "Emmeci Appalti S.r.l." → "Emmeci Appalti"
- "Fratelli Bianchi Trasporti S.p.A." → "Fratelli Bianchi Trasporti"  
- "Consorzio Edilizio Romano" → "Consorzio Edilizio Romano"
- "La Fenice Costruzioni" → "La Fenice Costruzioni"

OUTPUT: Scrivi SOLO il nome completo dell'azienda, nient'altro. Se non identificabile: "UNKNOWN"."""
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.0, max_output_tokens=100)
        )
        
        result = response.text.strip().replace('"', '').replace('\n', '').replace('*', '')
        
        # Rimuovi eventuali prefissi tipo "Nome:" o "Azienda:"
        import re
        result = re.sub(r'^(?:Nome|Azienda|Organizzazione|Output|Risposta)\s*:\s*', '', result, flags=re.IGNORECASE).strip()
        
        if "UNKNOWN" in result.upper() or len(result) < 3 or len(result) > 100:
            return ""
        
        # Pulisci forma giuridica come sicurezza aggiuntiva
        result = clean_company_name(result)
        print(f"DEBUG: Nome azienda estratto da clausole: '{result}'")
        return result
        
    except Exception as e:
        print(f"Errore estrazione nome azienda da clausole: {e}")
        return ""


def get_min_clause_words(norma: str) -> int:
    """
    Restituisce la soglia minima parole per clausola in base alla norma.
    Sotto questa soglia = ERRORE CRITICO (richiede recovery).
    
    Args:
        norma: Nome della norma (es. "ISO 50001")
    
    Returns:
        Soglia minima parole (fallback a MIN_CLAUSE_WORDS se norma non trovata)
    """
    return MIN_CLAUSE_WORDS_PER_NORM.get(norma, MIN_CLAUSE_WORDS)


def get_warning_clause_words(norma: str) -> int:
    """
    Restituisce la soglia warning parole per clausola in base alla norma.
    Tra MIN e WARNING = warning giallo (integrate ma segnalate come potenzialmente corte).
    
    Args:
        norma: Nome della norma (es. "ISO 50001")
    
    Returns:
        Soglia warning parole (fallback a 150 se norma non trovata)
    """
    return WARNING_CLAUSE_WORDS_PER_NORM.get(norma, 150)


# Numero di gruppi paralleli per elaborazione clausole
PARALLEL_GROUPS = 6


def split_clauses_into_groups(norma: str, num_groups: int = PARALLEL_GROUPS) -> List[List[str]]:
    """
    Divide le clausole di una norma in gruppi bilanciati per elaborazione parallela.
    
    Args:
        norma: Nome della norma (es. "ISO 9001")
        num_groups: Numero di gruppi da creare
    
    Returns:
        Lista di liste, ogni sotto-lista contiene le chiavi clausole per quel gruppo
    """
    clauses = NORM_CLAUSES.get(norma, [])
    if not clauses:
        return [[]]
    
    # Dividi in gruppi bilanciati
    groups = [[] for _ in range(num_groups)]
    for i, clause in enumerate(clauses):
        groups[i % num_groups].append(clause)
    
    # Rimuovi gruppi vuoti
    return [g for g in groups if g]


def build_group_schema(clauses: List[str]) -> Dict:
    """
    Costruisce schema JSON per un subset di clausole (usato per gruppi paralleli).
    
    Args:
        clauses: Lista chiavi clausole per questo gruppo
    
    Returns:
        Schema JSON compatibile con response_schema
    """
    clausole_properties = {
        clause: {"type": "string", "description": f"Evidenza per clausola {clause}"}
        for clause in clauses
    }
    
    return {
        "type": "object",
        "properties": {
            "clausole": {
                "type": "object",
                "properties": clausole_properties,
                "required": clauses
            }
        },
        "required": ["clausole"]
    }


def process_clause_group(
    group_idx: int,
    clauses: List[str],
    prompt_content: str,
    report_text: str,
    norma: str,
    api_key: str
) -> Tuple[int, Optional[Dict], Optional[str]]:
    """
    Elabora un singolo gruppo di clausole con chiamata API.
    
    Args:
        group_idx: Indice del gruppo (per ordinamento)
        clauses: Lista chiavi clausole da generare
        prompt_content: Prompt di sistema
        report_text: Testo del report Word
        norma: Norma di riferimento
        api_key: Chiave API Gemini
    
    Returns:
        Tuple (indice_gruppo, dict_clausole, errore)
    """
    from google import genai
    from google.genai import types
    
    MAX_RETRIES = 3
    import time
    
    last_error = None
    
    for attempt in range(MAX_RETRIES):
        try:
            client = create_genai_client(api_key)
            
            # Costruisci prompt specifico per questo gruppo
            clause_list = ", ".join(clauses)
            group_prompt = f"""{prompt_content}
    
    ---
    
    # DOCUMENTO WORD DA ANALIZZARE
    
    {report_text}
    
    ---
    
    # ISTRUZIONI GRUPPO {group_idx + 1}
    
    Genera SOLO le seguenti clausole: {clause_list}
    
    Output: JSON con chiave "clausole" contenente SOLO le chiavi richieste.
    """
            
            # Schema per questo gruppo specifico
            group_schema = build_group_schema(clauses)
            
            response = client.models.generate_content(
                model=GEMINI_MODEL_CHECKLIST,
                contents=group_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                    max_output_tokens=32000,
                    response_mime_type="application/json",
                    response_schema=group_schema
                )
            )
            
            import json
            response_text = response.text.strip()
            
            # Rimuovi markdown se presente
            if response_text.startswith("```"):
                lines = response_text.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines[-1].strip() == "```":
                    lines = lines[:-1]
                response_text = "\n".join(lines)
            
            result = json.loads(response_text)
            return group_idx, result.get("clausole", {}), None
            
        except Exception as e:
            last_error = str(e)
            # Gestione Rate Limit (429) o Resource Exhausted
            if "429" in last_error or "quota" in last_error.lower() or "resource" in last_error.lower():
                wait_time = 2 * (attempt + 1)  # 2s, 4s, 6s
                time.sleep(wait_time)
                continue
            
            # Se errore diverso, riprova comunque (es. overload server)
            if attempt < MAX_RETRIES - 1:
                time.sleep(1)
                continue
                
    return group_idx, None, f"Fallito dopo {MAX_RETRIES} tentativi: {last_error}"


def build_response_schema(norma: str) -> Dict:
    """
    Costruisce lo schema JSON per response_schema di Gemini.
    Forza l'API a produrre JSON valido con le chiavi esatte richieste.
    
    Args:
        norma: Nome della norma (es. "ISO 9001")
    
    Returns:
        Schema JSON compatibile con Gemini response_schema
    """
    clauses = NORM_CLAUSES.get(norma, [])
    
    # Costruisci proprietà clausole
    clausole_properties = {
        clause: {"type": "string", "description": f"Evidenza per clausola {clause}"}
        for clause in clauses
    }
    
    schema = {
        "type": "object",
        "properties": {
            "norma": {"type": "string", "description": "Nome della norma ISO"},
            "azienda": {"type": "string", "description": "Nome dell'azienda"},
            "data_elaborazione": {"type": "string", "description": "Data elaborazione YYYY-MM-DD"},
            "clausole": {
                "type": "object",
                "properties": clausole_properties,
                "required": clauses
            }
        },
        "required": ["norma", "azienda", "data_elaborazione", "clausole"]
    }
    
    return schema


def get_incomplete_clauses(json_data: Dict, norma: str, min_words: int = None) -> Dict:
    """
    Identifica TUTTE le clausole incomplete per una norma.
    
    Trova:
    - Clausole MANCANTI (chiave non presente)
    - Clausole VUOTE (chiave presente ma testo vuoto/None)
    - Clausole CORTE CRITICHE (< min_words, errore rosso, richiede recovery)
    - Clausole CORTE WARNING (tra min_words e warning_words, integrate ma segnalate)
    
    Args:
        json_data: JSON con clausole generate
        norma: Norma di riferimento (es. "ISO 14001")
        min_words: Soglia minima parole (default: soglia dinamica per norma)
    
    Returns:
        Dict con:
        - total_expected: Numero totale clausole attese per la norma
        - total_generated: Numero clausole valide (>= warning_words)
        - missing: Lista chiavi clausole mancanti
        - empty: Lista chiavi clausole con testo vuoto
        - short: Lista di tuple (chiave, conteggio_parole) per clausole < min_words (critiche)
        - short_warning: Lista di tuple (chiave, conteggio_parole) per clausole tra min e warning
        - all_incomplete: Lista clausole CRITICHE (missing + empty + short) per recovery
    
    Note:
        - Clausole < min_words (100): ERRORE CRITICO, richiedono recovery
        - Clausole tra min_words e warning_words: WARNING, integrate ma segnalate
        - Clausole >= warning_words: COMPLETE
    """
    # Usa soglia dinamica per norma se non specificata
    if min_words is None:
        min_words = get_min_clause_words(norma)
    
    # Soglia warning per questa norma
    warning_words = get_warning_clause_words(norma)
    
    # Clausole attese per questa norma
    expected_keys = NORM_CLAUSES.get(norma, [])
    total_expected = len(expected_keys)
    
    if not expected_keys:
        return {
            "total_expected": 0,
            "total_generated": 0,
            "missing": [],
            "empty": [],
            "short": [],
            "short_warning": [],
            "all_incomplete": []
        }
    
    clausole = json_data.get("clausole", {})
    generated_keys = set(clausole.keys())
    
    # Clausole mancanti (chiave non presente)
    missing = [k for k in expected_keys if k not in generated_keys]
    
    # Clausole vuote, corte critiche e corte warning
    empty = []
    short = []  # < min_words (100) = errore critico
    short_warning = []  # tra min_words e warning_words = warning giallo
    valid_count = 0
    
    for key in expected_keys:
        if key in clausole:
            text = clausole[key]
            if not text or not isinstance(text, str) or len(text.strip()) == 0:
                # Testo vuoto o None
                empty.append(key)
            else:
                word_count = len(text.split())
                if word_count < min_words:
                    # Critico: < 100 parole, richiede recovery
                    short.append((key, word_count))
                elif word_count < warning_words:
                    # Warning: tra 100 e 175, integrata ma segnalata
                    short_warning.append((key, word_count))
                    valid_count += 1  # Comunque contata come valida
                else:
                    # Completa: >= warning_words
                    valid_count += 1
    
    # Solo clausole CRITICHE per pulsante "Recupera" (non le warning)
    all_incomplete = missing + empty + [k for k, _ in short]
    
    return {
        "total_expected": total_expected,
        "total_generated": valid_count,
        "missing": missing,
        "empty": empty,
        "short": short,
        "short_warning": short_warning,
        "all_incomplete": all_incomplete
    }


def find_short_clauses(json_data: Dict, norma: str = None, min_words: int = None) -> List[Tuple[str, int]]:
    """
    Identifica clausole con numero di parole insufficiente.
    
    Args:
        json_data: JSON con clausole
        norma: Norma di riferimento (per soglia dinamica)
        min_words: Soglia minima parole (override manuale)
    
    Returns:
        Lista di tuple (chiave_clausola, conteggio_parole) per clausole troppo corte
    """
    # Usa soglia dinamica se norma specificata e min_words non override
    if min_words is None:
        min_words = get_min_clause_words(norma) if norma else MIN_CLAUSE_WORDS
    short = []
    clausole = json_data.get("clausole", {})
    
    for key, value in clausole.items():
        if value:
            word_count = len(value.split())
            if word_count < min_words:
                short.append((key, word_count))
    
    return short


def regenerate_short_clauses(
    json_data: Dict,
    short_clauses: List[Tuple[str, int]],
    report_text: str,
    norma: str,
    api_key: str,
    progress_callback=None
) -> Dict:
    """
    Rigenera solo le clausole troppo corte con chiamate API mirate.
    
    Args:
        json_data: JSON originale
        short_clauses: Lista clausole da rigenerare
        report_text: Testo del report Word
        norma: Norma di riferimento
        api_key: Chiave API
        progress_callback: Callback per progresso
    
    Returns:
        JSON aggiornato con clausole rigenerate
    """
    from google import genai
    from google.genai import types
    
    if not short_clauses:
        return json_data
    
    client = create_genai_client(api_key)
    
    # Rigenera ogni clausola corta
    total = len(short_clauses)
    for i, (clause_key, current_words) in enumerate(short_clauses):
        if progress_callback:
            progress_callback(
                85 + int((i / total) * 10),
                f"Rielaborazione clausola {clause_key} ({current_words} parole)..."
            )
        
        # Prompt specifico per rigenerare UNA clausola
        regen_prompt = f"""
Sei un auditor ISO. Devi espandere la seguente clausola con più dettagli.

CLAUSOLA DA ESPANDERE: {clause_key}
TESTO ATTUALE ({current_words} parole): {json_data['clausole'].get(clause_key, '')}

NORMA: {norma}
DOCUMENTI DI RIFERIMENTO:
{report_text}

ISTRUZIONI:
- Riscrivi SOLO questa clausola con minimo 150 parole
- Includi evidenze oggettive, numeri, date, riferimenti documentali
- Stile: prosa discorsiva formale, no elenchi puntati
- Output: SOLO il testo della clausola, niente JSON

CLAUSOLA ESPANSA:
"""
        
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL_CHECKLIST,
                contents=regen_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2,
                    max_output_tokens=2000
                )
            )
            
            new_text = response.text.strip()
            # Usa soglia dinamica per la norma specifica
            min_words_for_norm = get_min_clause_words(norma)
            if len(new_text.split()) >= min_words_for_norm:
                json_data["clausole"][clause_key] = new_text
                
        except Exception as e:
            # Se fallisce, mantieni la versione originale
            pass
    
    return json_data


def recover_missing_clauses(
    json_data: Dict,
    report_text: str,
    norma: str,
    api_key: str,
    progress_callback=None
) -> Tuple[Dict, List[str]]:
    """
    Identifica e rigenera le clausole MANCANTI dal JSON.
    
    Confronta le clausole generate con quelle attese da NORM_CLAUSES.
    Le clausole mancanti vengono rigenerate singolarmente con retry e cooldown
    per evitare rate limits.
    
    Args:
        json_data: JSON con clausole generate
        report_text: Testo del report Word
        norma: Norma di riferimento (es. "ISO 14001")
        api_key: Chiave API Gemini
        progress_callback: Callback per aggiornare progresso
    
    Returns:
        Tuple (JSON aggiornato, lista clausole ancora mancanti dopo recovery)
    """
    import time
    from google import genai
    from google.genai import types
    
    # Trova clausole attese per questa norma
    expected_keys = NORM_CLAUSES.get(norma, [])
    if not expected_keys:
        return json_data, []
    
    # Identifica clausole mancanti
    generated_keys = set(json_data.get("clausole", {}).keys())
    missing_keys = [k for k in expected_keys if k not in generated_keys]
    
    if not missing_keys:
        return json_data, []
    
    print(f"RECOVERY: {len(missing_keys)} clausole mancanti rilevate: {missing_keys}")
    
    client = create_genai_client(api_key)
    recovered = []
    still_missing = []
    
    total = len(missing_keys)
    for i, clause_key in enumerate(missing_keys):
        if progress_callback:
            progress_callback(
                88 + int((i / total) * 8),
                f"🔄 Recovery clausola {clause_key} ({i+1}/{total})..."
            )
        
        # Cooldown tra le chiamate per evitare rate limit (aumentato per pratiche grandi)
        if i > 0:
            time.sleep(5)  # 5 secondi tra ogni chiamata (aumentato da 3)
        
        # Prompt per generare la clausola mancante
        recovery_prompt = f"""
Sei un auditor ISO esperto. Devi generare l'evidenza per una clausola specifica.

CLAUSOLA DA GENERARE: {clause_key}
NORMA: {norma}

DOCUMENTI DI RIFERIMENTO (Report Audit):
{report_text[:80000]}

ISTRUZIONI TASSATIVE:
- Genera SOLO il testo per la clausola {clause_key}
- Minimo 150 parole, massimo 400 parole
- Includi evidenze oggettive, numeri, date, riferimenti documentali dal report
- Stile: prosa discorsiva formale, professionale
- NO elenchi puntati, NO JSON, SOLO testo
- Se non trovi evidenze specifiche, dichiara che l'informazione è stata visionata in corso di audit e risulta conforme

TESTO CLAUSOLA:
"""
        
        # Retry con backoff esponenziale (aumentato per pratiche grandi)
        max_retries = 5  # Aumentato da 3 per maggiore resilienza
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=GEMINI_MODEL_CHECKLIST,
                    contents=recovery_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.3,
                        max_output_tokens=4000  # Aumentato da 2500 per clausole complete
                    )
                )
                
                # NUOVO: Verifica finish_reason per rilevare troncature
                finish_reason = getattr(response.candidates[0], 'finish_reason', None) if response.candidates else None
                if finish_reason and str(finish_reason) == "MAX_TOKENS":
                    print(f"[WARNING] Troncatura rilevata su {clause_key}, tentativo recupero...")
                    # Riprova con prompt più corto
                    continue
                
                new_text = response.text.strip()
                
                # Pulisci eventuali wrapper markdown
                if new_text.startswith("```"):
                    lines = new_text.split("\n")
                    new_text = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
                
                if len(new_text.split()) >= 50:  # Soglia minima per accettare
                    json_data["clausole"][clause_key] = new_text
                    recovered.append(clause_key)
                    print(f"RECOVERY OK: {clause_key} ({len(new_text.split())} parole)")
                    break
                    
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    # Rate limit - attendi più a lungo (aumentato)
                    wait_time = (attempt + 1) * 15  # 15s, 30s, 45s, 60s, 75s
                    print(f"Rate limit su {clause_key}, attendo {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    print(f"Errore recovery {clause_key}: {e}")
                    if attempt == max_retries - 1:
                        break
                    time.sleep(5)  # Breve pausa anche per altri errori
        else:
            still_missing.append(clause_key)
            print(f"RECOVERY FALLITO: {clause_key} dopo {max_retries} tentativi")
    
    if recovered:
        print(f"RECOVERY COMPLETATO: {len(recovered)} clausole recuperate")
    
    return json_data, still_missing


def recover_incomplete_clauses(
    json_data: Dict,
    incomplete_keys: List[str],
    report_text: str,
    norma: str,
    api_key: str,
    progress_callback=None
) -> Tuple[Dict, List[str]]:
    """
    Recupera clausole incomplete specifiche (mancanti, vuote o corte).
    
    Differenza da recover_missing_clauses: accetta una lista esplicita di chiavi
    da recuperare, permettendo all'utente di scegliere quali recuperare.
    
    Args:
        json_data: JSON con clausole esistenti
        incomplete_keys: Lista di chiavi clausole da recuperare
        report_text: Testo del report Word
        norma: Norma di riferimento
        api_key: Chiave API Gemini
        progress_callback: Callback per aggiornare progresso
    
    Returns:
        Tuple (JSON aggiornato, lista clausole ancora mancanti dopo recovery)
    """
    import time
    from google import genai
    from google.genai import types
    
    if not incomplete_keys:
        return json_data, []
    
    print(f"RECOVERY INCOMPLETE: {len(incomplete_keys)} clausole da recuperare: {incomplete_keys}")
    
    client = create_genai_client(api_key)
    recovered = []
    still_missing = []
    
    # Assicura che esista la struttura clausole
    if "clausole" not in json_data:
        json_data["clausole"] = {}
    
    total = len(incomplete_keys)
    for i, clause_key in enumerate(incomplete_keys):
        if progress_callback:
            pct = int((i / total) * 100)
            progress_callback(pct, f"🔄 Recupero {clause_key} ({i+1}/{total})...")
        
        # Cooldown tra le chiamate
        if i > 0:
            time.sleep(5)
        
        # Calcola range parole in base alla norma
        min_words_norm = get_min_clause_words(norma)
        target_min = max(250, min_words_norm)  # Minimo 250 per nuove norme
        target_max = target_min + 250  # Range di 250 parole
        
        # Prompt per generare la clausola
        recovery_prompt = f"""
Sei un auditor ISO esperto. Devi generare l'evidenza per una clausola specifica.

CLAUSOLA DA GENERARE: {clause_key}
NORMA: {norma}

DOCUMENTI DI RIFERIMENTO (Report Audit):
{report_text[:80000]}

ISTRUZIONI TASSATIVE:
- Genera SOLO il testo per la clausola {clause_key}
- Minimo {target_min} parole, massimo {target_max} parole
- Includi evidenze oggettive, numeri, date, riferimenti documentali dal report
- Stile: prosa discorsiva formale, professionale
- NO elenchi puntati, NO JSON, SOLO testo
- NON usare citazioni numeriche [N], (documento X), (allegato N) o riferimenti a immagini/file
- NON usare virgolette per titoli documenti
- Se non trovi evidenze specifiche, dichiara che l'informazione è stata visionata in corso di audit e risulta conforme

TESTO CLAUSOLA:
"""
        
        # Retry con backoff esponenziale
        max_retries = 5
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=GEMINI_MODEL_CHECKLIST,
                    contents=recovery_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.0,
                        max_output_tokens=4000
                    )
                )
                
                new_text = response.text.strip()
                
                # Pulisci eventuali wrapper markdown
                if new_text.startswith("```"):
                    lines = new_text.split("\n")
                    new_text = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
                
                if len(new_text.split()) >= 120:  # Soglia minima recovery: 120 parole
                    json_data["clausole"][clause_key] = new_text
                    recovered.append(clause_key)
                    print(f"RECOVERY OK: {clause_key} ({len(new_text.split())} parole)")
                    break
                    
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    wait_time = (attempt + 1) * 15
                    print(f"Rate limit su {clause_key}, attendo {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    print(f"Errore recovery {clause_key}: {e}")
                    if attempt == max_retries - 1:
                        break
                    time.sleep(5)
        else:
            still_missing.append(clause_key)
            print(f"RECOVERY FALLITO: {clause_key} dopo {max_retries} tentativi")
    
    if recovered:
        print(f"RECOVERY COMPLETATO: {len(recovered)} clausole recuperate")
    
    return json_data, still_missing


# ═══ CACHE PROMPT IN MEMORIA ═══
# Evita reletture disco ad ogni chiamata
_prompt_cache: Dict[str, str] = {}
_prompts_prefetched: bool = False


def prefetch_all_prompts() -> None:
    """
    Precarica tutti i prompt in memoria all'avvio dell'app.
    Chiamare questa funzione una volta all'avvio per eliminare latenza I/O.
    """
    global _prompts_prefetched
    if _prompts_prefetched:
        return
    
    for norma in SUPPORTED_NORMS:
        load_prompt(norma)  # Popola la cache
    
    _prompts_prefetched = True


def compress_text(text: str) -> str:
    """
    Comprime il testo rimuovendo whitespace eccessivo per ridurre token input.
    Mantiene la leggibilità preservando singoli newline e spazi.
    """
    import re
    # Rimuovi spazi multipli
    text = re.sub(r' +', ' ', text)
    # Riduci newline multipli a massimo 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Rimuovi spazi a inizio/fine riga
    text = '\n'.join(line.strip() for line in text.split('\n'))
    # Rimuovi righe completamente vuote consecutive
    text = re.sub(r'\n\n\n+', '\n\n', text)
    return text.strip()


def load_prompt(norma: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Carica il prompt specifico per la norma selezionata.
    
    Args:
        norma: Nome della norma (es. "ISO 9001")
    
    Returns:
        Tuple (contenuto prompt, errore se presente)
    """
    # Controlla cache prima
    if norma in _prompt_cache:
        return _prompt_cache[norma], None
    
    prompt_file = CHECKLIST_PROMPTS.get(norma)
    if not prompt_file:
        return None, f"Norma non supportata: {norma}. Norme valide: {', '.join(SUPPORTED_NORMS)}"
    
    prompt_path = os.path.join(PROMPTS_CHECKLIST_DIR, prompt_file)
    
    if not os.path.exists(prompt_path):
        return None, f"File prompt non trovato: {prompt_path}"
    
    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        # Salva in cache
        _prompt_cache[norma] = content
        return content, None
    except Exception as e:
        return None, f"Errore lettura prompt: {str(e)}"


def extract_text_from_word(word_input) -> Tuple[Optional[str], Optional[str]]:
    """
    Estrae il testo completo da un documento Word.
    
    Args:
        word_input: Bytes, BytesIO, o percorso file Word
    
    Returns:
        Tuple (testo estratto, errore se presente)
    """
    from docx import Document
    
    try:
        # Gestione diversi tipi di input
        if isinstance(word_input, bytes):
            doc = Document(BytesIO(word_input))
        elif isinstance(word_input, BytesIO):
            doc = Document(word_input)
        elif hasattr(word_input, 'read'):
            # File-like object (es. UploadedFile di Streamlit)
            content = word_input.read()
            doc = Document(BytesIO(content))
        else:
            # Percorso file
            doc = Document(word_input)
        
        # 1. Estrai testo dal CORPO
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        
        # 2. Estrai testo da INTESTAZIONI e PIÈ DI PAGINA (Cruciale per Nome Azienda)
        headers_footers = []
        for section in doc.sections:
            for part in [section.header, section.footer]:
                if not part: continue
                try:
                    # Estrai dai paragrafi dell'header/footer
                    for para in part.paragraphs:
                        if para.text.strip():
                            headers_footers.append(para.text.strip())
                    # Estrai dalle tabelle dell'header/footer
                    for table in part.tables:
                        for row in table.rows:
                            row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_vals:
                                headers_footers.append(" | ".join(row_vals))
                except Exception:
                    continue # Ignora errori su singole sezioni
        
        # Aggiungi headers all'inizio del testo per dare contesto immediato
        if headers_footers:
            paragraphs = list(set(headers_footers)) + ["--- INIZIO REPORT ---"] + paragraphs

        # 3. Estrai testo dalle TABELLE del corpo (ottimizzato: set per dedup celle ripetute)
        seen_cells = set()
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Skip celle già viste (Word ripete celle merged)
                    cell_key = (id(table), cell_text[:50] if cell_text else "")
                    if cell_text and cell_key not in seen_cells:
                        seen_cells.add(cell_key)
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n\n".join(paragraphs)
        
        # Applica compressione per ridurre token
        full_text = compress_text(full_text)
        
        if not full_text.strip():
            return None, "Il documento Word è vuoto o non contiene testo estraibile"
        
        return full_text, None
        
    except Exception as e:
        return None, f"Errore estrazione testo da Word: {str(e)}"


def produce_checklist_json(
    report_input,
    norma: str,
    api_key: str,
    progress_callback=None
) -> Tuple[Optional[Dict], Optional[str]]:
    """
    Genera il JSON strutturato delle evidenze dal Report Word.
    
    Args:
        report_input: Bytes, BytesIO, o percorso file Word del report
        norma: Nome della norma (es. "ISO 9001")
        api_key: Chiave API Gemini
        progress_callback: Funzione opzionale per aggiornare progresso (0-100)
    
    Returns:
        Tuple (dict JSON, errore se presente)
    """
    from google import genai
    from google.genai import types
    
    # Step 1: Carica prompt
    if progress_callback:
        progress_callback(10, "Caricamento prompt...")
    
    prompt_content, error = load_prompt(norma)
    if error:
        return None, error
    
    # Step 2: Estrai testo dal Word
    if progress_callback:
        progress_callback(20, "Estrazione testo dal Report...")
    
    report_text, error = extract_text_from_word(report_input)
    if error:
        return None, error
    
    # Info sul testo estratto
    word_count = len(report_text.split())
    char_count = len(report_text)
    
    if progress_callback:
        progress_callback(30, f"Testo estratto: {word_count} parole, {char_count} caratteri")
    
    # Step 3: Costruisci il messaggio completo
    full_message = f"""{prompt_content}

---

# DOCUMENTO WORD DA ANALIZZARE

Di seguito il contenuto completo del Report Word con le evidenze oggettive già estratte.
Analizza attentamente tutto il contenuto e produci il JSON richiesto.

---

{report_text}
"""
    
    # Step 4: Chiama Gemini con progress simulato
    if progress_callback:
        progress_callback(40, "Avvio elaborazione AI...")
    
    try:
        import threading
        import time
        
        client = create_genai_client(api_key)
        
        # Messaggi di progresso contestuali (simulati durante attesa API)
        num_clausole = len(NORM_CLAUSES.get(norma, []))
        progress_messages = [
            (42, f"🔄 Elaborazione AI in corso... ({num_clausole} clausole da generare)"),
            (48, "📋 Analisi sezioni 4-5 (Contesto e Leadership)..."),
            (54, "📋 Analisi sezioni 6-7 (Pianificazione e Supporto)..."),
            (60, "📋 Analisi sezione 8 (Attività operative)..."),
            (66, "📋 Analisi sezione 9 (Valutazione prestazioni)..."),
            (72, "📋 Analisi sezione 10 (Miglioramento)..."),
            (78, "🔨 Composizione JSON strutturato..."),
        ]
        
        # Flag per fermare il thread quando API risponde
        api_complete = threading.Event()
        api_response = [None]  # Contenitore per risposta (mutabile)
        api_error = [None]
        
        def call_api():
            """Thread che esegue la chiamata API."""
            try:
                # Costruisci schema per la norma specifica
                response_schema = build_response_schema(norma)
                
                api_response[0] = client.models.generate_content(
                    model=GEMINI_MODEL_CHECKLIST,
                    contents=full_message,
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        max_output_tokens=65536,
                        response_mime_type="application/json",
                        response_schema=response_schema  # Schema stretto per JSON valido
                    )
                )
            except Exception as e:
                api_error[0] = e
            finally:
                api_complete.set()
        
        def simulate_progress():
            """Thread che simula progresso durante attesa API."""
            for pct, msg in progress_messages:
                if api_complete.is_set():
                    break
                if progress_callback:
                    progress_callback(pct, msg)
                # Attesa variabile per sembrare naturale
                wait_time = 2.0 + (pct % 3) * 0.5  # 2-3.5 sec tra step
                api_complete.wait(timeout=wait_time)
        
        # Avvia entrambi i thread
        api_thread = threading.Thread(target=call_api)
        progress_thread = threading.Thread(target=simulate_progress)
        
        api_thread.start()
        progress_thread.start()
        
        # Attendi completamento API
        api_thread.join()
        progress_thread.join()  # Attendi che finisca il ciclo progress
        
        # Verifica errori
        if api_error[0]:
            raise api_error[0]
        
        response = api_response[0]
        
        if progress_callback:
            progress_callback(80, "Parsing risposta JSON...")
        
        # Step 5: Parsa la risposta JSON
        response_text = response.text.strip()
        
        # Rimuovi eventuali markdown code blocks
        if response_text.startswith("```"):
            # Rimuovi ```json all'inizio e ``` alla fine
            lines = response_text.split("\n")
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines[-1].strip() == "```":
                lines = lines[:-1]
            response_text = "\n".join(lines)
        
        json_data = json.loads(response_text)
        
        # Step 6: Valida clausole e richiedi rielaborazione se troppo corte
        if progress_callback:
            total_clauses = len(json_data.get("clausole", {}))
            progress_callback(82, f"Validazione {total_clauses} clausole...")
        
        short_clauses = find_short_clauses(json_data)
        
        if short_clauses:
            if progress_callback:
                progress_callback(84, f"Trovate {len(short_clauses)} clausole da espandere...")
            
            # Rigenera clausole troppo corte
            json_data = regenerate_short_clauses(
                json_data=json_data,
                short_clauses=short_clauses,
                report_text=report_text,
                norma=norma,
                api_key=api_key,
                progress_callback=progress_callback
            )
        
        # Aggiungi metadati
        json_data["_metadata"] = {
            "norma": norma,
            "generato_il": datetime.now().isoformat(),
            "modello": GEMINI_MODEL_CHECKLIST,
            "parole_input": word_count,
            "caratteri_input": char_count,
            "clausole_rigenerate": len(short_clauses) if short_clauses else 0
        }
        
        # Riepilogo finale
        final_short = find_short_clauses(json_data)
        total_clauses = len(json_data.get("clausole", {}))
        
        if progress_callback:
            if final_short:
                progress_callback(98, f"Completato: {total_clauses} clausole ({len(final_short)} ancora brevi)")
            else:
                progress_callback(100, f"✓ {total_clauses} clausole generate correttamente!")
        
        # Pulisci forma giuridica dal nome azienda e dalle clausole
        json_data = clean_legal_suffix_from_json(json_data)
        
        return json_data, None
        
    except json.JSONDecodeError as e:
        return None, f"Errore parsing JSON dalla risposta AI: {str(e)}\n\nRisposta ricevuta:\n{response_text[:500]}..."
    except Exception as e:
        return None, f"Errore chiamata API Gemini: {str(e)}"


def extract_company_from_docx_structure(word_input) -> str:
    """
    Estrae il nome azienda in modo DETERMINISTICO dal docx generato dalla Fase 1.

    Il report di Fase 1 (structured_evidence_generator) inserisce come secondo
    paragrafo un subtitle nel formato "Audit - <NOME AZIENDA>", pensato
    esplicitamente per permettere a Tab 2 di identificare l'azienda senza
    ricorrere ad estrazione AI (che può troncare o sbagliare nomi ambigui).

    Args:
        word_input: Bytes, BytesIO, percorso file, o file-like object

    Returns:
        Nome azienda se identificato in modo attendibile, altrimenti stringa vuota.
    """
    from docx import Document
    import re

    try:
        # Gestione input: evita di consumare file-like non riavvolgibili
        if isinstance(word_input, bytes):
            doc = Document(BytesIO(word_input))
        elif isinstance(word_input, BytesIO):
            pos = word_input.tell()
            word_input.seek(0)
            doc = Document(word_input)
            word_input.seek(pos)
        elif hasattr(word_input, "read"):
            # File-like non riavvolgibile: skip sicuro, il chiamante userà il fallback AI
            return ""
        else:
            doc = Document(word_input)

        subtitle_pattern = re.compile(r"^\s*Audit\s*[-–—]\s*(.+?)\s*$", re.IGNORECASE)
        invalid_values = {
            "AZIENDA NON IDENTIFICATA", "N.D.", "N/A", "UNKNOWN", "NONE", "NULL", ""
        }

        # Il subtitle è al paragrafo 1 (dopo il titolo). Scansione di 10 per tolleranza.
        for para in doc.paragraphs[:10]:
            m = subtitle_pattern.match(para.text)
            if not m:
                continue
            candidate = m.group(1).strip()
            if candidate and candidate.upper() not in invalid_values and len(candidate) >= 3:
                return candidate

        return ""
    except Exception as e:
        print(f"DEBUG: extract_company_from_docx_structure fallita: {e}")
        return ""


def extract_company_metadata(report_text: str, api_key: str, filename: str = "") -> str:
    """
    Estrae il nome COMPLETO dell'azienda dal testo del report Word.

    Strategy:
    1. Cerca nel testo con AI (inizio + fine report per trovare intestazioni/firme).
    2. Se fallisce, usa il nome del file pulito come fallback.
    """
    from google import genai
    from google.genai import types
    import re
    
    # Tentativo 1: Nome dal file (Fallback solido)
    clean_filename = ""
    if filename:
        # Rimuovi estensione e prefissi comuni
        base = os.path.splitext(os.path.basename(filename))[0]
        # Pulisci da underscore e date
        clean_filename = base.replace('_', ' ').replace('-', ' ')
        # Rimuovi parole comuni report
        clean_filename = re.sub(r'(?i)\b(report|relazione|audit|evidenze|checklist|bozza|finale)\b', '', clean_filename).strip()
    
    try:
        client = create_genai_client(api_key)
        
        # Prepara contesto ricco: Inizio (intestazione) + Fine (firme timbri)
        context_text = report_text[:25000]
        if len(report_text) > 30000:
            context_text += "\n\n[...OMISSIS...]\n\n" + report_text[-8000:]
            
        prompt = f"""COMPITO: Estrai il NOME COMPLETO dell'azienda auditata da questo Report.

TESTO DEL REPORT:
{context_text}

NOME FILE (fallback): "{clean_filename}"

ISTRUZIONI FONDAMENTALI:
1. Identifica l'AZIENDA AUDITATA (il soggetto controllato, NON l'ente certificatore o l'auditor)
2. Cerca label: "Organizzazione:", "Ragione Sociale:", "Audit presso:", "Cliente:", "Sito:"
3. Restituisci il NOME COMMERCIALE COMPLETO con TUTTE LE PAROLE
4. NON troncare - "Emmeci Appalti" deve restare "Emmeci Appalti", NON "Emmeci"
5. ESCLUDI la forma giuridica (S.r.l., S.p.A., ecc.) ma MANTIENI tutto il resto

ESEMPI CORRETTI:
- "Emmeci Appalti S.r.l." → "Emmeci Appalti"
- "Fratelli Bianchi Trasporti" → "Fratelli Bianchi Trasporti"
- "Consorzio Edilizio Romano" → "Consorzio Edilizio Romano"

OUTPUT: Scrivi SOLO il nome completo dell'azienda, nient'altro. Se non identificabile: "UNKNOWN"."""
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.0,
                max_output_tokens=100  # Aumentato per nomi lunghi
            )
        )
        
        result = response.text.strip().replace('"', '').replace('\n', '').replace('*', '')
        
        # Rimuovi eventuali prefissi tipo "Nome:" o "Azienda:"
        result = re.sub(r'^(?:Nome|Azienda|Organizzazione|Output|Risposta)\s*:\s*', '', result, flags=re.IGNORECASE).strip()
        
        # Validazione base
        if "UNKNOWN" in result.upper() or len(result) < 3:
            return clean_filename if clean_filename else "AZIENDA_SCONOSCIUTA"
        
        # Pulisci forma giuridica
        result = clean_company_name(result)
        print(f"DEBUG: Nome azienda da report: '{result}'")
        return result
        
    except Exception as e:
        print(f"Errore estrazione nome da report: {e}")
        # In caso di crash AI, usa il file
        return clean_filename if clean_filename else ""


def produce_checklist_json_parallel(
    report_input,
    norma: str,
    api_key: str,
    progress_callback=None,
    company_name_override: str = ""
) -> Tuple[Optional[Dict], Optional[str]]:
    """
    Genera il JSON usando elaborazione PARALLELA delle clausole (3x più veloce).
    
    Divide le clausole in gruppi e li elabora contemporaneamente con ThreadPoolExecutor.
    
    Args:
        report_input: Bytes, BytesIO, o percorso file Word del report
        norma: Nome della norma (es. "ISO 9001")
        api_key: Chiave API Gemini
        progress_callback: Funzione opzionale per aggiornare progresso (0-100)
        company_name_override: Nome azienda già estratto da Tab 1 (priorità massima)
    
    Returns:
        Tuple (dict JSON, errore se presente)
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    
    # Step 1: Carica prompt
    if progress_callback:
        progress_callback(5, "📁 Caricamento prompt...")
    
    prompt_content, error = load_prompt(norma)
    if error:
        return None, error
    
    # Step 2: Estrai testo dal Word
    if progress_callback:
        progress_callback(10, "📄 Estrazione testo dal Report...")
    
    report_text, error = extract_text_from_word(report_input)
    if error:
        return None, error
    
    word_count = len(report_text.split())
    char_count = len(report_text)
    
    # Step 3: Dividi clausole in gruppi
    groups = split_clauses_into_groups(norma)
    num_groups = len(groups)
    total_clauses = sum(len(g) for g in groups)
    
    if progress_callback:
        progress_callback(15, f"🔄 Elaborazione parallela: {num_groups} gruppi, {total_clauses} clausole...")
    
    # Step 4: Lancia elaborazione parallela
    all_clausole = {}
    errors = []
    completed = 0
    company_name = ""
    
    with ThreadPoolExecutor(max_workers=PARALLEL_GROUPS + 1) as executor:
        # Risoluzione nome azienda per priorità:
        #   1) Override esplicito dal frontend (Tab 1)
        #   2) Subtitle deterministico del docx Fase 1 ("Audit - XXX")
        #   3) Estrazione AI in parallelo (fallback)
        # L'AI viene lanciata SOLO se (1) e (2) non risolvono, così si risparmia
        # una chiamata Gemini nel caso comune e si evitano troncamenti AI su nomi ambigui.
        future_company = None
        if company_name_override:
            company_name = company_name_override
            print(f"DEBUG: Uso nome azienda da Tab 1 (override): {company_name}")
        else:
            structural_name = extract_company_from_docx_structure(report_input)
            if structural_name:
                company_name = structural_name
                print(f"DEBUG: Uso nome azienda da subtitle docx (strutturale): {company_name}")
            else:
                filename_str = str(report_input) if isinstance(report_input, str) else getattr(report_input, "name", "")
                future_company = executor.submit(extract_company_metadata, report_text, api_key, filename_str)

        # Task paralleli per gruppi clausole
        futures = {
            executor.submit(
                process_clause_group,
                idx, group, prompt_content, report_text, norma, api_key
            ): idx
            for idx, group in enumerate(groups)
        }

        for future in as_completed(futures):
            group_idx = futures[future]
            try:
                idx, clausole, error = future.result()
                completed += 1

                if error:
                    errors.append(f"Gruppo {idx + 1}: {error}")
                elif clausole:
                    all_clausole.update(clausole)

                if progress_callback:
                    pct = 15 + int((completed / num_groups) * 60)
                    progress_callback(pct, f"✓ Gruppo {completed}/{num_groups} completato ({len(clausole or {})} clausole)")

            except Exception as e:
                errors.append(f"Gruppo {group_idx + 1}: {str(e)}")

        # Recupera risultato AI azienda SOLO se il fallback è stato attivato
        if future_company is not None:
            try:
                company_name = future_company.result()
                print(f"DEBUG: Uso nome azienda da AI: {company_name}")
            except Exception:
                company_name = ""
    
    if not all_clausole:
        return None, f"Nessuna clausola generata. Errori: {'; '.join(errors)}"
    
    # Step 5: Costruisci JSON finale
    if progress_callback:
        progress_callback(78, f"🔧 Composizione JSON: {len(all_clausole)} clausole...")
    
    json_data = {
        "norma": norma,
        "azienda": company_name,
        "data_elaborazione": datetime.now().strftime("%Y-%m-%d"),
        "clausole": all_clausole
    }
    
    # Step 6: Valida e rigenera clausole corte
    if progress_callback:
        progress_callback(80, f"📋 Validazione {len(all_clausole)} clausole...")
    
    short_clauses = find_short_clauses(json_data)
    
    if short_clauses:
        if progress_callback:
            progress_callback(82, f"🔁 Espansione {len(short_clauses)} clausole brevi...")
        
        json_data = regenerate_short_clauses(
            json_data=json_data,
            short_clauses=short_clauses,
            report_text=report_text,
            norma=norma,
            api_key=api_key,
            progress_callback=progress_callback
        )
    
    # Step 7: Recovery clausole mancanti (NEW!)
    # Verifica completezza e rigenera eventuali clausole perse per rate limit
    if progress_callback:
        progress_callback(85, f"🔍 Verifica completezza clausole...")
    
    json_data, still_missing = recover_missing_clauses(
        json_data=json_data,
        report_text=report_text,
        norma=norma,
        api_key=api_key,
        progress_callback=progress_callback
    )
    
    # Calcola clausole recuperate
    expected_count = len(NORM_CLAUSES.get(norma, []))
    generated_count = len(json_data.get("clausole", {}))
    recovered_count = generated_count - len(all_clausole)
    
    # Metadati
    json_data["_metadata"] = {
        "norma": norma,
        "generato_il": datetime.now().isoformat(),
        "modello": GEMINI_MODEL_CHECKLIST,
        "parole_input": word_count,
        "caratteri_input": char_count,
        "modalita": "parallela",
        "gruppi": num_groups,
        "clausole_rigenerate": len(short_clauses) if short_clauses else 0,
        "clausole_recuperate": recovered_count if recovered_count > 0 else 0,
        "clausole_ancora_mancanti": still_missing if still_missing else None,
        "errori_gruppi": errors if errors else None
    }
    
    final_short = find_short_clauses(json_data)
    
    if progress_callback:
        if final_short:
            progress_callback(98, f"⚠️ {len(all_clausole)} clausole ({len(final_short)} brevi)")
        else:
            progress_callback(100, f"✅ {len(all_clausole)} clausole generate!")
    
    # Estrai nome azienda dalle clausole SOLO SE il nome corrente è vuoto o invalido
    # Questo è un FALLBACK - il nome da Tab 1 ha PRIORITÀ
    current_name = json_data.get("azienda", "")
    if not current_name or current_name.upper() in ["AZIENDA_SCONOSCIUTA", "UNKNOWN", ""]:
        if json_data.get("clausole"):
            extracted_name = extract_company_from_clauses(json_data["clausole"], api_key)
            if extracted_name:
                json_data["azienda"] = extracted_name
                print(f"DEBUG: Nome azienda estratto dalle clausole (fallback): {extracted_name}")
    else:
        print(f"DEBUG: Uso nome azienda già presente: {current_name}")
    
    # Pulisci forma giuridica dal nome azienda e dalle clausole
    json_data = clean_legal_suffix_from_json(json_data)
    
    return json_data, None


def validate_json_structure(json_data: Dict, norma: str) -> Tuple[bool, List[str]]:
    """
    Valida che il JSON contenga i campi richiesti.
    
    Args:
        json_data: Dati JSON da validare
        norma: Norma di riferimento
    
    Returns:
        Tuple (is_valid, lista di warning/errori)
    """
    issues = []
    
    # Campi obbligatori top-level
    required_fields = ["norma", "azienda"]
    for field in required_fields:
        if field not in json_data:
            issues.append(f"Campo obbligatorio mancante: '{field}'")
    
    # Verifica presenza clausole
    clausole = json_data.get("clausole", {})
    if not clausole:
        issues.append("Campo 'clausole' mancante o vuoto")
    else:
        # Conta clausole vuote
        empty_count = sum(1 for v in clausole.values() if not v or v.strip() == "")
        if empty_count > 0:
            issues.append(f"{empty_count} clausole con valore vuoto")
        
        # Verifica lunghezza testi
        short_count = 0
        for key, value in clausole.items():
            if value and len(value.split()) < 50:
                short_count += 1
        if short_count > 0:
            issues.append(f"{short_count} clausole con testo troppo breve (<50 parole)")
    
    is_valid = len([i for i in issues if "obbligatorio" in i.lower()]) == 0
    
    return is_valid, issues


def get_json_stats(json_data: Dict) -> Dict:
    """
    Genera statistiche sul JSON prodotto.
    
    Args:
        json_data: Dati JSON
    
    Returns:
        Dict con statistiche
    """
    clausole = json_data.get("clausole", {})
    
    stats = {
        "norma": json_data.get("norma", "Non specificata"),
        "azienda": json_data.get("azienda", "Non specificata"),
        "num_clausole": len(clausole),
        "clausole_compilate": sum(1 for v in clausole.values() if v and v.strip()),
        "clausole_vuote": sum(1 for v in clausole.values() if not v or not v.strip()),
        "parole_totali": sum(len(str(v).split()) for v in clausole.values() if v),
        "caratteri_totali": sum(len(str(v)) for v in clausole.values() if v)
    }
    
    if stats["num_clausole"] > 0:
        stats["percentuale_compilazione"] = round(
            stats["clausole_compilate"] / stats["num_clausole"] * 100, 1
        )
    else:
        stats["percentuale_compilazione"] = 0
    
    return stats


def generate_json_filename(json_data: Dict) -> str:
    """
    Genera nome file per il JSON.
    
    Args:
        json_data: Dati JSON
    
    Returns:
        Nome file suggerito
    """
    norma = json_data.get("norma", "AUDIT")
    azienda = json_data.get("azienda", "")
    
    # Pulisci caratteri speciali
    safe_norma = re.sub(r'[^\w\s-]', '', str(norma)).replace(' ', '_')[:20]
    safe_azienda = re.sub(r'[^\w\s-]', '', str(azienda)).replace(' ', '_')[:30]
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if safe_azienda:
        return f"Checklist_{safe_norma}_{safe_azienda}_{timestamp}.json"
    return f"Checklist_{safe_norma}_{timestamp}.json"
