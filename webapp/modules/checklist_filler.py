# ==============================================================================
# CHECKLIST FILLER - Compilazione automatica checklist Word
# ==============================================================================

import os
import sys
import json
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from io import BytesIO

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import CHECKLIST_TEMPLATES, TEMPLATES_DIR


def parse_json_input(json_input) -> Tuple[Optional[Dict], Optional[str]]:
    """
    Parsa input JSON (file o testo).
    
    Args:
        json_input: File object o stringa JSON
    
    Returns:
        Tuple (dict parsato, messaggio errore se presente)
    """
    try:
        if hasattr(json_input, 'read'):
            # È un file object
            content = json_input.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            data = json.loads(content)
        else:
            # È una stringa
            data = json.loads(json_input)
        
        return data, None
    except json.JSONDecodeError as e:
        return None, f"Errore parsing JSON: {str(e)}"
    except Exception as e:
        return None, f"Errore lettura input: {str(e)}"


def detect_standard(json_data: Dict) -> Tuple[Optional[str], Optional[str]]:
    """
    Rileva automaticamente la norma dal campo 'norma' del JSON.
    
    Args:
        json_data: Dati JSON parsati
    
    Returns:
        Tuple (norma rilevata, percorso template) o (None, errore)
    """
    norma_field = json_data.get("norma", "")
    
    if not norma_field:
        # Cerca in altri campi possibili
        norma_field = json_data.get("standard", "")
        norma_field = norma_field or json_data.get("norm", "")
        norma_field = norma_field or json_data.get("iso", "")
    
    if not norma_field:
        return None, "Campo 'norma' non trovato nel JSON"
    
    # Cerca corrispondenza nella mappatura
    # IMPORTANTE: Ordina le chiavi per lunghezza decrescente per matchare prima le più specifiche
    # Questo evita che "ISO 9001" matchi con "ISO 27001" (9001 è substring di 27001)
    norma_upper = norma_field.upper()
    
    # Prima prova match esatto
    for key, template_file in CHECKLIST_TEMPLATES.items():
        if key.upper() == norma_upper:
            template_path = os.path.join(TEMPLATES_DIR, template_file)
            return norma_field, template_path
    
    # Poi ordina chiavi per lunghezza decrescente e cerca match parziale
    sorted_keys = sorted(CHECKLIST_TEMPLATES.keys(), key=len, reverse=True)
    for key in sorted_keys:
        # Verifica che la chiave sia contenuta nella norma
        # E che non sia un falso positivo (es. "9001" in "27001")
        if key.upper() in norma_upper:
            # Estrai il numero della norma dalla chiave (es. "9001" da "ISO 9001")
            import re
            key_numbers = re.findall(r'\d+', key)
            norma_numbers = re.findall(r'\d+', norma_field)
            
            # Verifica che almeno un numero della chiave corrisponda esattamente a un numero nella norma
            if key_numbers and norma_numbers:
                if any(kn in norma_numbers for kn in key_numbers):
                    template_path = os.path.join(TEMPLATES_DIR, CHECKLIST_TEMPLATES[key])
                    return norma_field, template_path
            else:
                # Se non ci sono numeri, usa il match semplice
                template_path = os.path.join(TEMPLATES_DIR, CHECKLIST_TEMPLATES[key])
                return norma_field, template_path
    
    return norma_field, None  # Norma trovata ma nessun template


def get_json_summary(json_data: Dict) -> Dict:
    """
    Genera un riepilogo del JSON per conferma utente.
    
    Args:
        json_data: Dati JSON parsati
    
    Returns:
        Dict con statistiche e info
    """
    summary = {
        "norma": json_data.get("norma", "Non specificata"),
        "azienda": json_data.get("azienda", json_data.get("organizzazione", "Non specificata")),
        "data_audit": json_data.get("data_audit", "Non specificata"),
        "auditor": json_data.get("auditor", "Non specificato"),
        "num_requisiti": 0,
        "requisiti_conformi": 0,
        "requisiti_non_conformi": 0,
        "requisiti_con_osservazioni": 0
    }
    
    contenuto = json_data.get("contenuto", json_data.get("evidenze", []))
    if isinstance(contenuto, list):
        summary["num_requisiti"] = len(contenuto)
        
        for item in contenuto:
            conformita = str(item.get("conformita", "")).upper()
            if "NON CONFORME" in conformita:
                summary["requisiti_non_conformi"] += 1
            elif "OSSERVAZION" in conformita:
                summary["requisiti_con_osservazioni"] += 1
            elif "CONFORME" in conformita:
                summary["requisiti_conformi"] += 1
    
    return summary


def fill_checklist(json_data: Dict, template_path: str) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Compila la checklist (Word o Excel) con i dati del JSON.
    
    Rileva automaticamente il tipo di file dal percorso template e chiama
    la funzione appropriata (Word o Excel).
    
    Supporta due formati JSON:
    1. Nuovo formato (da Tab 2): { "clausole": { "iso9001_4_1": "testo...", ... } }
    2. Vecchio formato: { "contenuto": [ { "requisito": "4.1", "evidenza": "..." }, ... ] }
    
    Args:
        json_data: Dati JSON parsati
        template_path: Percorso template Word (.docx) o Excel (.xlsx)
    
    Returns:
        Tuple (bytes documento compilato, errore se presente)
    """
    # Verifica template esiste
    if not os.path.exists(template_path):
        return None, f"Template non trovato: {template_path}"
    
    # Rileva tipo template e smista alla funzione corretta
    ext = os.path.splitext(template_path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        return fill_checklist_excel(json_data, template_path)
    elif ext == '.docx':
        return fill_checklist_word(json_data, template_path)
    else:
        return None, f"Formato template non supportato: {ext}. Usa .docx o .xlsx"


def fill_checklist_word(json_data: Dict, template_path: str) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Compila la checklist Word con i dati del JSON.
    
    Args:
        json_data: Dati JSON parsati
        template_path: Percorso template Word (.docx)
    
    Returns:
        Tuple (bytes Word compilato, errore se presente)
    """
    from docx import Document
    from docx.shared import Pt
    
    try:
        doc = Document(template_path)
    except Exception as e:
        return None, f"Errore apertura template Word: {str(e)}"
    
    # Determina formato JSON
    clausole = json_data.get("clausole", {})
    is_new_format = isinstance(clausole, dict) and len(clausole) > 0
    
    # Mappa dei placeholder base da sostituire
    replacements = {
        "{{AZIENDA}}": json_data.get("azienda", ""),
        "{{DATA_AUDIT}}": json_data.get("data_audit", datetime.now().strftime("%d/%m/%Y")),
        "{{AUDITOR}}": json_data.get("auditor", ""),
        "{{NORMA}}": json_data.get("norma", ""),
    }
    
    # Aggiungi clausole come placeholder {{chiave}}
    if is_new_format:
        for key, value in clausole.items():
            val_str = str(value) if value else ""
            # Supporta sia lowercase che uppercase per robustezza
            replacements[f"{{{{{key}}}}}"] = val_str       # es. {{iso9001_4_1}}
            replacements[f"{{{{{key.upper()}}}}}"] = val_str   # es. {{ISO9001_4_1}}
    
    # Contatori per statistiche
    replaced_count = 0
    
    # Funzione helper per sostituire in un run mantenendo formattazione
    def replace_in_paragraph(paragraph):
        nonlocal replaced_count
        full_text = paragraph.text
        modified = False
        
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
                modified = True
                replaced_count += 1
        
        if modified:
            # Sostituisci il testo mantenendo la formattazione del primo run
            if paragraph.runs:
                # Salva formattazione del primo run
                first_run = paragraph.runs[0]
                # Pulisci tutti i run
                for run in paragraph.runs[1:]:
                    run.clear()
                first_run.text = full_text
            else:
                paragraph.text = full_text
    
    # Funzione helper per iterare ricorsivamente le tabelle annidate
    def iter_all_tables(parent):
        """Itera tutte le tabelle, incluse quelle annidate dentro altre tabelle."""
        from docx.table import Table
        from docx.oxml.ns import qn
        
        # Trova tutti gli elementi tabella nel parent
        for tbl in parent._element.iter(qn('w:tbl')):
            yield Table(tbl, parent)
    
    # Sostituisci placeholder nei paragrafi
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Cerca e sostituisci in TUTTE le tabelle (incluse annidate)
    for table in iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Cerca anche nelle intestazioni/piè di pagina
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)
    
    # Fallback per vecchio formato (lista di requisiti)
    if not is_new_format:
        contenuto = json_data.get("contenuto", json_data.get("evidenze", []))
        if isinstance(contenuto, list):
            for table in doc.tables:
                for row in table.rows:
                    cells_text = [cell.text.strip() for cell in row.cells]
                    
                    for item in contenuto:
                        req_id = str(item.get("requisito", ""))
                        
                        if req_id and req_id in cells_text[0]:
                            for i, cell in enumerate(row.cells):
                                cell_text = cell.text.strip().upper()
                                
                                if "EVIDENZA" in cell_text or i >= 2:
                                    evidenza = item.get("evidenza", "")
                                    if evidenza and i + 1 < len(row.cells):
                                        row.cells[i + 1].text = evidenza
                                        replaced_count += 1
                                    break
                                
                                if "CONFORM" in cell_text:
                                    conformita = item.get("conformita", "")
                                    if conformita and i + 1 < len(row.cells):
                                        row.cells[i + 1].text = conformita
    
    # Salva in bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue(), None


def fill_checklist_excel(json_data: Dict, template_path: str) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Compila la checklist Excel con i dati del JSON.
    
    Cerca placeholder {{chiave}} in tutte le celle del foglio e li sostituisce
    con i valori corrispondenti dal JSON.
    
    Args:
        json_data: Dati JSON parsati
        template_path: Percorso template Excel (.xlsx)
    
    Returns:
        Tuple (bytes Excel compilato, errore se presente)
    """
    from openpyxl import load_workbook
    
    try:
        wb = load_workbook(template_path)
    except Exception as e:
        return None, f"Errore apertura template Excel: {str(e)}"
    
    # Determina formato JSON
    clausole = json_data.get("clausole", {})
    is_new_format = isinstance(clausole, dict) and len(clausole) > 0
    
    # Mappa dei placeholder base da sostituire
    replacements = {
        "{{AZIENDA}}": json_data.get("azienda", ""),
        "{{DATA_AUDIT}}": json_data.get("data_audit", datetime.now().strftime("%d/%m/%Y")),
        "{{DATA_ELABORAZIONE}}": json_data.get("data_elaborazione", datetime.now().strftime("%Y-%m-%d")),
        "{{AUDITOR}}": json_data.get("auditor", ""),
        "{{NORMA}}": json_data.get("norma", ""),
    }
    
    # Aggiungi clausole come placeholder {{chiave}}
    if is_new_format:
        for key, value in clausole.items():
            val_str = str(value) if value else ""
            # Supporta sia lowercase che uppercase per robustezza
            replacements[f"{{{{{key}}}}}"] = val_str       # es. {{iso27001_4_1}}
            replacements[f"{{{{{key.upper()}}}}}"] = val_str   # es. {{ISO27001_4_1}}
    
    # Contatori per statistiche
    replaced_count = 0
    
    # Itera su tutti i fogli
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Itera su tutte le celle (incluse quelle in range definiti)
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    original_value = cell.value
                    new_value = original_value
                    
                    # Sostituisci tutti i placeholder trovati nella cella
                    for placeholder, replacement in replacements.items():
                        if placeholder in new_value:
                            new_value = new_value.replace(placeholder, replacement)
                            replaced_count += 1
                    
                    # Aggiorna solo se modificato
                    if new_value != original_value:
                        cell.value = new_value
                        
                        # Se il testo è lungo, abilita wrap text per mostrarlo tutto
                        if len(new_value) > 50:
                            from openpyxl.styles import Alignment
                            cell.alignment = Alignment(wrap_text=True, vertical='top')
    
    # Salva in bytes
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    print(f"[Excel Filler] Sostituiti {replaced_count} placeholder")
    
    return buffer.getvalue(), None


def generate_checklist_filename(json_data: Dict, template_path: str = None) -> str:
    """
    Genera nome file per checklist compilata.
    
    Rileva automaticamente l'estensione dal template se fornito.
    
    Args:
        json_data: Dati JSON parsati
        template_path: Percorso template (opzionale, per determinare estensione)
    
    Returns:
        Nome file con estensione corretta
    """
    norma = json_data.get("norma", "AUDIT")
    azienda = json_data.get("azienda", "")
    
    # Pulisci caratteri speciali
    safe_norma = re.sub(r'[^\w\s-]', '', norma).replace(' ', '_')[:20]
    safe_azienda = re.sub(r'[^\w\s-]', '', azienda).replace(' ', '_')[:30]
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Determina estensione dal template o default a .docx
    ext = ".docx"
    if template_path:
        template_ext = os.path.splitext(template_path)[1].lower()
        if template_ext in ['.xlsx', '.xls']:
            ext = ".xlsx"
    
    if safe_azienda:
        return f"Checklist_{safe_norma}_{safe_azienda}_{timestamp}{ext}"
    return f"Checklist_{safe_norma}_{timestamp}{ext}"

