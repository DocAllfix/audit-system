"""
Test webapp/v2/incremental_docx_builder.build_unprocessed_section.

Verifica:
- Lista vuota → niente file creato (success ma output_path=None)
- Lista popolata → docx valido con tabella di tracking
- Raggruppamento per phase preserva l'ordine di apparizione
- build_all_sections accetta il nuovo kwarg senza rompere il default
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

_REPO_ROOT = Path(__file__).resolve().parent.parent.parent
_WEBAPP_DIR = _REPO_ROOT / "webapp"
sys.path.insert(0, str(_WEBAPP_DIR))

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from v2.incremental_docx_builder import (  # noqa: E402
    _UNPROCESSED_SECTION_INDEX,
    build_all_sections,
    build_unprocessed_section,
)


def test_empty_list_no_file_created(tmp_path):
    res = build_unprocessed_section(
        unprocessed_files=[],
        session_id="empty_test",
        base_dir=tmp_path,
    )
    assert res.success is True
    assert res.output_path is None
    assert res.documents_count == 0


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_populated_list_creates_docx(tmp_path):
    items = [
        {"filename": "scan1.pdf", "reason": "OCR fallito", "phase": "ocr"},
        {"filename": "broken.pdf", "reason": "PDF corrupted", "phase": "triage"},
        {"filename": "x.pdf", "reason": "batch_0 perso", "phase": "analyze"},
    ]
    res = build_unprocessed_section(
        unprocessed_files=items,
        session_id="popul_test",
        base_dir=tmp_path,
    )
    assert res.success is True
    assert res.output_path is not None
    assert res.output_path.exists()
    assert res.documents_count == 3
    assert res.section_index == _UNPROCESSED_SECTION_INDEX
    # Filename con prefix 99 per ordinamento merge
    assert res.output_path.name.startswith("99_")

    # Validazione contenuto
    doc = Document(str(res.output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    assert "DOCUMENTI NON ELABORATI" in full_text
    # 3 phases distinte = 3 sezioni Heading 2 + 3 tabelle
    assert len(doc.tables) == 3
    # Conta righe complessive (header + 1 per file ciascuna)
    total_rows = sum(len(t.rows) for t in doc.tables)
    assert total_rows == 6  # 3 header + 3 data rows


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_phase_grouping_preserves_first_appearance_order(tmp_path):
    items = [
        {"filename": "a.pdf", "reason": "x", "phase": "analyze"},
        {"filename": "b.pdf", "reason": "y", "phase": "triage"},
        {"filename": "c.pdf", "reason": "z", "phase": "analyze"},
    ]
    res = build_unprocessed_section(
        unprocessed_files=items,
        session_id="order_test",
        base_dir=tmp_path,
    )
    assert res.success
    doc = Document(str(res.output_path))
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    # Ordine di prima apparizione: analyze, poi triage
    assert h2[0].startswith("Analyze")
    assert h2[1].startswith("Triage")


def test_invalid_session_id_returns_error():
    res = build_unprocessed_section(
        unprocessed_files=[{"filename": "x.pdf", "reason": "test"}],
        session_id="../invalid/path",
    )
    assert res.success is False
    assert "session_id" in (res.error or "").lower()


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_build_all_sections_includes_unprocessed_when_provided(tmp_path):
    """Assicura zero regressione: senza unprocessed_files il comportamento è identico."""
    parsed_data = {"sezioni": [], "meta": {}}
    # Senza unprocessed
    r1 = build_all_sections(parsed_data, "noreg_test", base_dir=tmp_path)
    n_default = len(r1)

    # Con unprocessed
    r2 = build_all_sections(
        parsed_data, "noreg_test2", base_dir=tmp_path,
        unprocessed_files=[{"filename": "x.pdf", "reason": "y", "phase": "z"}],
    )
    assert len(r2) == n_default + 1
    last = r2[-1]
    assert last.section_index == _UNPROCESSED_SECTION_INDEX
    assert last.success is True
