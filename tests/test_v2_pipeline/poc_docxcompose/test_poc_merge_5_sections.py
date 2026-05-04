"""
PoC Fase 0.5 — Validazione docxcompose su 5 sezioni di test.

Obiettivo: verificare che docxcompose regga gli scenari reali del nostro audit
PRIMA di committarlo in Fase 7. Se uno di questi test fallisce, dobbiamo
ripensare l'approccio (es. pandoc, oppure assemblaggio XML manuale).

Scenari coperti (rappresentativi degli audit reali):
1. Sezione semplice — solo paragrafi (caso base, ATTESTATI)
2. Sezione con tabella complessa — header colorato, celle multi-riga (VISURA, DVR)
3. Sezione con stili eterogenei — heading, sottotitoli, grassetti (BILANCIO)
4. Sezione con loop dinamico — molte righe simili ripetute (ELENCO ATTESTATI)
5. Sezione con caratteri speciali italiani — accenti, simboli € (CCNL)

Verifiche post-merge:
- Il file finale apre senza errori (load via python-docx)
- Tutto il contenuto delle 5 sezioni è presente nel merge finale
- Nessun crash su python-docx 1.x + lxml su Windows
"""
from __future__ import annotations

from pathlib import Path

import pytest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docxcompose.composer import Composer


# ──────────────────────────────────────────────────────────────────────────────
# Generatori di sezioni di test
# ──────────────────────────────────────────────────────────────────────────────

def _new_doc() -> Document:
    """Crea un Document con margini standard del nostro audit."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    return doc


def _section_simple_paragraphs(out_path: Path) -> Path:
    """Sezione 1: paragrafi semplici (caso ATTESTATI)."""
    doc = _new_doc()
    doc.add_heading("01 — DOCUMENTAZIONE LEGALE", level=1)
    for i in range(5):
        p = doc.add_paragraph(
            f"Paragrafo {i+1}. Esaminato il documento numero {i+1}, "
            f"si rileva che il contenuto è conforme alle aspettative. " * 5
        )
        p.paragraph_format.line_spacing = 1.5
    doc.save(str(out_path))
    return out_path


def _section_with_table(out_path: Path) -> Path:
    """Sezione 2: tabella con header colorato, celle multi-riga (VISURA)."""
    doc = _new_doc()
    doc.add_heading("02 — VISURA CAMERALE", level=1)
    doc.add_paragraph("Estrazione dati anagrafici dalla visura.")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid"
    headers = [("Campo", "Valore")]
    rows = [
        ("Ragione sociale", "DEMO CONSULTING S.R.L."),
        ("P. IVA", "02012345678"),
        ("Capitale sociale", "10.000,00 €"),
    ]
    for i, (k, v) in enumerate(headers + rows):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v

    doc.save(str(out_path))
    return out_path


def _section_heterogeneous_styles(out_path: Path) -> Path:
    """Sezione 3: heading multipli, grassetti, allineamenti (BILANCIO)."""
    doc = _new_doc()
    doc.add_heading("03 — BILANCIO 2024", level=1)
    doc.add_heading("3.1 Stato Patrimoniale", level=2)
    p = doc.add_paragraph()
    p.add_run("ATTIVO: ").bold = True
    p.add_run("euro 1.250.000,00")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("3.2 Conto Economico", level=2)
    p2 = doc.add_paragraph()
    run = p2.add_run("Ricavi: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    p2.add_run("euro 980.000,00")

    doc.save(str(out_path))
    return out_path


def _section_dynamic_loop(out_path: Path) -> Path:
    """Sezione 4: molte righe ripetute (ELENCO ATTESTATI)."""
    doc = _new_doc()
    doc.add_heading("04 — ELENCO ATTESTATI", level=1)
    table = doc.add_table(rows=21, cols=3)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "N°"
    table.rows[0].cells[1].text = "Lavoratore"
    table.rows[0].cells[2].text = "Scadenza"
    for i in range(1, 21):
        table.rows[i].cells[0].text = str(i)
        table.rows[i].cells[1].text = f"Cognome{i} Nome{i}"
        table.rows[i].cells[2].text = f"31/12/{2025 + (i % 5)}"

    doc.save(str(out_path))
    return out_path


def _section_special_chars(out_path: Path) -> Path:
    """Sezione 5: accenti, simboli, caratteri speciali italiani (CCNL)."""
    doc = _new_doc()
    doc.add_heading("05 — CCNL e Norme Contrattuali", level=1)
    doc.add_paragraph(
        "L'azienda applica il CCNL Cooperative Sociali. "
        "Importi: 1.500,00 € mensili; rivalutazione semestrale. "
        "Riferimenti: art. 36 della Costituzione; ex art. 2087 c.c. "
        "Lavoratori coperti: ≥ 90% del personale; copertura sanitaria attivata. "
        "Forme giuridiche di riferimento: S.p.A., S.r.l., Coop. Soc."
    )
    doc.add_paragraph("Caratteri speciali: à è ì ò ù — § © ® ™ ½ ¼ ¾")

    doc.save(str(out_path))
    return out_path


SECTION_BUILDERS = [
    _section_simple_paragraphs,
    _section_with_table,
    _section_heterogeneous_styles,
    _section_dynamic_loop,
    _section_special_chars,
]


# ──────────────────────────────────────────────────────────────────────────────
# Test PoC
# ──────────────────────────────────────────────────────────────────────────────

def test_poc_each_section_builds_standalone(tmp_path):
    """Ogni section builder produce un .docx valido individualmente."""
    for i, builder in enumerate(SECTION_BUILDERS):
        path = builder(tmp_path / f"section_{i:02d}.docx")
        assert path.exists()
        # Riapri con python-docx: valida che il file è ben formato
        doc = Document(str(path))
        # Almeno un heading presente
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1


def test_poc_merge_5_sections_into_single_docx(tmp_path):
    """
    Cuore del PoC: docxcompose deve mergiare 5 sezioni eterogenee in 1 docx.
    """
    section_paths = [
        builder(tmp_path / f"section_{i:02d}.docx")
        for i, builder in enumerate(SECTION_BUILDERS)
    ]

    # Apri la prima come base, accodaci le altre 4
    base = Document(str(section_paths[0]))
    composer = Composer(base)
    for path in section_paths[1:]:
        composer.append(Document(str(path)))

    final_path = tmp_path / "merged.docx"
    composer.save(str(final_path))

    assert final_path.exists()
    assert final_path.stat().st_size > 5000, "Merged docx sospetto piccolo"


def test_poc_merged_content_completeness(tmp_path):
    """
    Verifica che TUTTO il contenuto delle 5 sezioni sia nel merge finale.
    Cerca marker univoci di ogni sezione nel testo del docx mergiato.
    """
    section_paths = [
        builder(tmp_path / f"section_{i:02d}.docx")
        for i, builder in enumerate(SECTION_BUILDERS)
    ]
    base = Document(str(section_paths[0]))
    composer = Composer(base)
    for path in section_paths[1:]:
        composer.append(Document(str(path)))

    final_path = tmp_path / "merged.docx"
    composer.save(str(final_path))

    # Estrai tutto il testo dal merged
    final_doc = Document(str(final_path))
    full_text = "\n".join(p.text for p in final_doc.paragraphs)
    # Aggiungi anche testo nelle tabelle
    for tbl in final_doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    # Marker univoci di ogni sezione
    markers = [
        "01 — DOCUMENTAZIONE LEGALE",
        "02 — VISURA CAMERALE",
        "DEMO CONSULTING S.R.L.",
        "03 — BILANCIO 2024",
        "04 — ELENCO ATTESTATI",
        "05 — CCNL e Norme Contrattuali",
        "à è ì ò ù",
    ]
    missing = [m for m in markers if m not in full_text]
    assert not missing, f"Marker mancanti nel merge: {missing}"


def test_poc_merged_table_count(tmp_path):
    """Le tabelle delle sezioni 2 e 4 devono essere preservate nel merge."""
    section_paths = [
        builder(tmp_path / f"section_{i:02d}.docx")
        for i, builder in enumerate(SECTION_BUILDERS)
    ]
    base = Document(str(section_paths[0]))
    composer = Composer(base)
    for path in section_paths[1:]:
        composer.append(Document(str(path)))

    final_path = tmp_path / "merged.docx"
    composer.save(str(final_path))

    final_doc = Document(str(final_path))
    # Almeno 2 tabelle (visura + elenco attestati)
    assert len(final_doc.tables) >= 2


def test_poc_merge_preserves_heading_levels(tmp_path):
    """Le gerarchie heading 1/2 delle sezioni devono restare nel merge."""
    section_paths = [
        builder(tmp_path / f"section_{i:02d}.docx")
        for i, builder in enumerate(SECTION_BUILDERS)
    ]
    base = Document(str(section_paths[0]))
    composer = Composer(base)
    for path in section_paths[1:]:
        composer.append(Document(str(path)))

    final_path = tmp_path / "merged.docx"
    composer.save(str(final_path))

    final_doc = Document(str(final_path))
    h1_count = sum(1 for p in final_doc.paragraphs if p.style.name == "Heading 1")
    h2_count = sum(1 for p in final_doc.paragraphs if p.style.name == "Heading 2")
    # 5 sezioni × 1 H1 ciascuna; sezione 3 ha 2 H2
    assert h1_count == 5, f"Atteso 5 H1, trovati {h1_count}"
    assert h2_count >= 2, f"Atteso ≥ 2 H2 (dalla sezione 3), trovati {h2_count}"
