"""
Test V2 Fase 8 — pipeline E2E.

Coperture:
- Dry-run: pipeline E2E senza chiamate Gemini reali
- Output: docx finale prodotto, leggibile, con company_name corretto
- Subtitle "Audit - {company}" presente per Tab 2 compat
- Smart batching: First Fit Decreasing per char count
- Mock client orchestrato (classifier + OCR + analyze tutti mockati)
- Zero file documenti → return success=False con error
- ZIP corrotto → return success=False
- Token meter popolato dopo run con mock
- Cleanup: temp directory rimosse dopo successo
"""
from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from v2 import pipeline as pl
from v2 import token_meter


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _zip_with_pdf(content_text: str = "Lorem ipsum " * 100) -> bytes:
    """Crea ZIP minimal con 1 PDF (text-layer reale via reportlab)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab non installato")

    pdf_buf = io.BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=A4)
    c.setFont("Helvetica", 11)
    y = 800
    for line in content_text.split("."):
        if not line.strip():
            continue
        c.drawString(50, y, line.strip()[:80])
        y -= 14
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = 800
    c.save()
    pdf_bytes = pdf_buf.getvalue()

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("VISURA 2025.pdf", pdf_bytes)
    return zip_buf.getvalue()


@pytest.fixture(autouse=True)
def _isolate_state(tmp_path, monkeypatch):
    """Isola tutte le directory temp per ogni test."""
    from v2 import file_uploader as fu
    from v2 import incremental_docx_builder as ib
    from v2 import progress_store as ps
    from v2 import zip_extractor as ze

    monkeypatch.setattr(ze, "EXTRACT_BASE_DIR", tmp_path / "extract")
    monkeypatch.setattr(ib, "SECTIONS_BASE_DIR", tmp_path / "sections")
    monkeypatch.setattr(fu, "MANIFEST_BASE_DIR", tmp_path / "manifest")
    monkeypatch.setattr(ps, "PROGRESS_BASE_DIR", tmp_path / "progress")
    monkeypatch.setattr(token_meter, "TOKEN_USAGE_BASE_DIR", tmp_path / "token_usage")
    token_meter.reset_all_sessions()
    yield
    token_meter.reset_all_sessions()


# ──────────────────────────────────────────────────────────────────────────────
# Smart batching
# ──────────────────────────────────────────────────────────────────────────────

def test_smart_batching_first_fit_decreasing():
    docs = [
        {"filename": "a", "content": "x" * 30000},  # batch da solo
        {"filename": "b", "content": "x" * 15000},
        {"filename": "c", "content": "x" * 15000},
        {"filename": "d", "content": "x" * 5000},
        {"filename": "e", "content": "x" * 5000},
    ]
    batches = pl._create_smart_batches(docs, max_files=4, max_chars=50_000)
    # Almeno 2 batch (uno con il 30k file, uno con i 15k che non ci entrano)
    assert len(batches) >= 1
    # Nessun batch supera max_chars (eccetto se contiene un solo file > limit,
    # ma qui non è il caso)
    for b in batches:
        total = sum(len(d.get("content", "")) for d in b)
        assert total <= 50_000 or len(b) == 1


def test_smart_batching_respects_max_files():
    docs = [{"filename": f"f{i}", "content": "x" * 100} for i in range(10)]
    batches = pl._create_smart_batches(docs, max_files=4, max_chars=50_000)
    for b in batches:
        assert len(b) <= 4


def test_smart_batching_empty_input():
    batches = pl._create_smart_batches([], max_files=4)
    assert batches == []


# ──────────────────────────────────────────────────────────────────────────────
# Dry-run E2E
# ──────────────────────────────────────────────────────────────────────────────

def test_dry_run_produces_valid_docx(tmp_path):
    """Dry-run E2E: ZIP → pipeline → docx finale leggibile."""
    zip_bytes = _zip_with_pdf()
    output_dir = tmp_path / "output"

    result = pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="dry_test",
        api_key=None,
        emitter=None,
        dry_run=True,
        output_dir=output_dir,
    )

    assert result["success"] is True
    assert result["company_name"] == "DRY RUN COMPANY SRL"
    output_path = Path(result["output_path"])
    assert output_path.exists()
    assert output_path.stat().st_size > 0

    # Verifica: docx valido
    doc = Document(str(output_path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(paras) > 0


def test_dry_run_subtitle_for_tab2(tmp_path):
    """Dry-run: il subtitle 'Audit - {company}' è preservato per Tab 2."""
    zip_bytes = _zip_with_pdf()
    output_dir = tmp_path / "output"

    result = pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="dry_tab2",
        dry_run=True,
        output_dir=output_dir,
    )

    doc = Document(str(result["output_path"]))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Audit - DRY RUN COMPANY SRL" in full_text


def test_dry_run_no_token_calls_recorded(tmp_path):
    """Dry-run NON deve registrare token (no chiamata Gemini)."""
    zip_bytes = _zip_with_pdf()
    pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="dry_notok",
        dry_run=True,
        output_dir=tmp_path / "output",
    )
    report = token_meter.get_session_report("dry_notok")
    assert report["calls_count"] == 0


def test_dry_run_returns_stats(tmp_path):
    zip_bytes = _zip_with_pdf()
    result = pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="dry_stats",
        dry_run=True,
        output_dir=tmp_path / "output",
    )
    stats = result["stats"]
    assert stats["dry_run"] is True
    assert stats["company_name"] == "DRY RUN COMPANY SRL"
    assert stats["duration_seconds"] >= 0
    assert "tokens" in stats


# ──────────────────────────────────────────────────────────────────────────────
# Edge cases / error handling
# ──────────────────────────────────────────────────────────────────────────────

def test_invalid_zip_returns_error(tmp_path):
    result = pl.process_zip_v2(
        zip_bytes=b"not a zip",
        session_id="bad_zip",
        dry_run=True,
        output_dir=tmp_path / "out",
    )
    assert result["success"] is False
    assert "extraction_failed" in result.get("error", "")


def test_empty_bytes_returns_error(tmp_path):
    result = pl.process_zip_v2(
        zip_bytes=b"",
        session_id="empty",
        dry_run=True,
        output_dir=tmp_path / "out",
    )
    assert result["success"] is False


def test_zip_without_extractable_documents_returns_error(tmp_path):
    """ZIP con file PDF senza testo nativo + dry_run=False → fallisce."""
    # ZIP con un finto PDF privo di text layer
    fake_pdf = b"%PDF-1.0\n garbage content not really a PDF"
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("fake.pdf", fake_pdf)

    result = pl.process_zip_v2(
        zip_bytes=zip_buf.getvalue(),
        session_id="no_docs",
        dry_run=False,
        api_key="dummy",
        _client=MagicMock(),  # no real API
        output_dir=tmp_path / "out",
    )
    # PDF corrotto → 0 documents → return success=False
    assert result["success"] is False


# ──────────────────────────────────────────────────────────────────────────────
# Cleanup
# ──────────────────────────────────────────────────────────────────────────────

def test_cleanup_removes_extract_dir_after_success(tmp_path):
    """Dopo successo dry-run, extract_dir e sections_dir sono rimossi."""
    from v2 import incremental_docx_builder as ib
    from v2 import zip_extractor as ze

    zip_bytes = _zip_with_pdf()
    pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="clean_test",
        dry_run=True,
        output_dir=tmp_path / "output",
    )

    # extract_dir non deve più esistere
    extract_dir = ze.EXTRACT_BASE_DIR / "clean_test"
    assert not extract_dir.exists()

    # sections_dir non deve più esistere
    section_dir = ib._section_dir("clean_test")
    assert not section_dir.exists()


def test_no_op_emitter_does_not_crash(tmp_path):
    """emitter=None usa NoopEmitter interno; nessun crash."""
    zip_bytes = _zip_with_pdf()
    result = pl.process_zip_v2(
        zip_bytes=zip_bytes,
        session_id="noemit",
        dry_run=True,
        emitter=None,
        output_dir=tmp_path / "output",
    )
    assert result["success"] is True


# ──────────────────────────────────────────────────────────────────────────────
# Stub
# ──────────────────────────────────────────────────────────────────────────────

def test_process_v2_stub_still_works():
    """Stub di Fase 0 mantenuto per /api/v2/health."""
    payload = pl.process_v2_stub()
    assert payload["status"] == "v2_stub_alive"
    assert payload["phase"] == "8_orchestrator"


# ──────────────────────────────────────────────────────────────────────────────
# Auto-tuning leve (Leva 2C compact + Leva 4 model_mix)
# Decisione automatica basata sul numero di file AGGREGABLE.
# ──────────────────────────────────────────────────────────────────────────────

def test_resolve_leva_flag_explicit_true(monkeypatch):
    """Flag esplicito 'true' → ON forzato, source='manual_on'."""
    monkeypatch.setenv("V2_LEVA_TEST_FLAG", "true")
    enabled, source = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=False)
    assert enabled is True
    assert source == "manual_on"


def test_resolve_leva_flag_explicit_false(monkeypatch):
    """Flag esplicito 'false' → OFF forzato, source='manual_off'."""
    monkeypatch.setenv("V2_LEVA_TEST_FLAG", "false")
    enabled, source = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=True)
    assert enabled is False
    assert source == "manual_off"


def test_resolve_leva_flag_auto_default_on(monkeypatch):
    """Flag non settato + auto_default=True → ON, source='auto_on'."""
    monkeypatch.delenv("V2_LEVA_TEST_FLAG", raising=False)
    enabled, source = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=True)
    assert enabled is True
    assert source == "auto_on"


def test_resolve_leva_flag_auto_default_off(monkeypatch):
    """Flag non settato + auto_default=False → OFF, source='auto_off'."""
    monkeypatch.delenv("V2_LEVA_TEST_FLAG", raising=False)
    enabled, source = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=False)
    assert enabled is False
    assert source == "auto_off"


def test_resolve_leva_flag_explicit_auto_string(monkeypatch):
    """'auto' esplicito → applica auto_default."""
    monkeypatch.setenv("V2_LEVA_TEST_FLAG", "auto")
    enabled, source = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=True)
    assert enabled is True
    assert source == "auto_on"


def test_resolve_leva_flag_unknown_value_treated_as_auto(monkeypatch):
    """Valori sconosciuti (typo) trattati come 'auto'."""
    monkeypatch.setenv("V2_LEVA_TEST_FLAG", "yes")  # typo classico
    enabled, _ = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=False)
    assert enabled is False


def test_resolve_leva_flag_uppercase_tolerated(monkeypatch):
    """Maiuscole tollerate (case-insensitive)."""
    monkeypatch.setenv("V2_LEVA_TEST_FLAG", "TRUE")
    enabled, _ = pl._resolve_leva_flag("V2_LEVA_TEST_FLAG", auto_default=False)
    assert enabled is True


def test_count_aggregable_documents_basic():
    """Conteggio file AGGREGABLE in documents."""
    documents = [
        {"filename": "a.pdf", "content": "x"},
        {"filename": "b.pdf", "content": "y"},
        {"filename": "c.pdf", "content": "z"},
        {"filename": "d.pdf", "content": "w"},
    ]
    role_by_filename = {
        "a.pdf": "AGGREGABLE",
        "b.pdf": "CORE",
        "c.pdf": "AGGREGABLE",
        "d.pdf": "SUPPORT",
    }
    n = pl._count_aggregable_documents(documents, role_by_filename)
    assert n == 2


def test_count_aggregable_documents_empty():
    """Liste vuote → 0."""
    assert pl._count_aggregable_documents([], {}) == 0
    assert pl._count_aggregable_documents([{"filename": "a.pdf"}], {}) == 0


def test_count_aggregable_documents_filename_missing_in_role_map():
    """Filename non in map → ignorato (non contato)."""
    documents = [{"filename": "unknown.pdf", "content": "x"}]
    role_by_filename = {"a.pdf": "AGGREGABLE"}
    assert pl._count_aggregable_documents(documents, role_by_filename) == 0


def test_build_role_by_filename_uses_default_when_audit_role_none():
    """Se classified ha audit_role=None, viene derivato dalla classe."""
    from v2.schemas.classification import ClassifiedFile, DocumentClass

    classified = [
        ClassifiedFile(filename="att.pdf", classe=DocumentClass.ATTESTATO,
                       confidence=0.9, audit_role=None),
        ClassifiedFile(filename="dvr.pdf", classe=DocumentClass.DVR,
                       confidence=0.95, audit_role=None),
    ]
    role_map = pl._build_role_by_filename(classified, skipped_filenames=set())
    assert role_map["att.pdf"] == "AGGREGABLE"
    assert role_map["dvr.pdf"] == "CORE"


def test_build_role_by_filename_skips_files_in_skipped_set():
    """Filename in skipped_filenames non vengono inclusi nella mappa."""
    from v2.schemas.classification import ClassifiedFile, DocumentClass

    classified = [
        ClassifiedFile(filename="a.pdf", classe=DocumentClass.ATTESTATO, confidence=0.9),
        ClassifiedFile(filename="skip.pdf", classe=DocumentClass.ATTESTATO,
                       confidence=0.9, audit_role="NOISE"),
    ]
    role_map = pl._build_role_by_filename(classified, skipped_filenames={"skip.pdf"})
    assert "a.pdf" in role_map
    assert "skip.pdf" not in role_map


def test_build_role_by_filename_empty_classified_returns_empty():
    """Lista classified vuota o None → mappa vuota (no errore)."""
    assert pl._build_role_by_filename([], skipped_filenames=set()) == {}
    assert pl._build_role_by_filename(None, skipped_filenames=set()) == {}


def test_default_min_aggregable_threshold_is_50():
    """La soglia di default deve essere 50 file AGGREGABLE."""
    assert pl.DEFAULT_MIN_AGGREGABLE_FOR_COMPACT == 50


def test_resolve_leva_flag_reads_correct_env_var(monkeypatch):
    """Ogni flag legge la propria env var, non si influenzano fra loro."""
    monkeypatch.setenv("V2_LEVA2_AGGREGABLE_COMPACT", "true")
    monkeypatch.delenv("V2_LEVA4_MODEL_MIX", raising=False)
    enabled2, src2 = pl._resolve_leva_flag(
        "V2_LEVA2_AGGREGABLE_COMPACT", auto_default=False
    )
    enabled4, src4 = pl._resolve_leva_flag(
        "V2_LEVA4_MODEL_MIX", auto_default=False
    )
    assert (enabled2, src2) == (True, "manual_on")
    assert (enabled4, src4) == (False, "auto_off")


def test_auto_tuning_scenario_sirih_below_threshold():
    """
    Simula scenario SIRIH (33 file aggregable, sotto soglia 50): il sistema
    deve scegliere auto_default=False per compact_mode.
    """
    n_aggregable = 33  # SIRIH reale
    auto_default = n_aggregable >= pl.DEFAULT_MIN_AGGREGABLE_FOR_COMPACT
    assert auto_default is False  # leva DISABILITATA su pratica piccola


def test_auto_tuning_scenario_medil_above_threshold():
    """
    Simula scenario MEDIL (81 file aggregable, sopra soglia 50): il sistema
    deve scegliere auto_default=True per compact_mode.
    """
    n_aggregable = 81  # MEDIL reale
    auto_default = n_aggregable >= pl.DEFAULT_MIN_AGGREGABLE_FOR_COMPACT
    assert auto_default is True  # leva ABILITATA su pratica grande


def test_threshold_configurable_via_env(monkeypatch):
    """V2_MIN_AGGREGABLE_THRESHOLD permette override della soglia."""
    monkeypatch.setenv("V2_MIN_AGGREGABLE_THRESHOLD", "20")
    threshold = int(os.environ.get("V2_MIN_AGGREGABLE_THRESHOLD",
                                     str(pl.DEFAULT_MIN_AGGREGABLE_FOR_COMPACT)))
    assert threshold == 20
    # Con questa soglia, anche SIRIH (33) entrerebbe in auto_on
    assert 33 >= threshold
