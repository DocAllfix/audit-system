"""
Test V2 Fase 7 — docx_merger.

Coperture:
- list_session_partials: ordering deterministico per prefix numerico
- merge_session_sections: produce 1 file finale leggibile
- Skip di file con prefix non valido
- Empty session → success=False con error
- Force fallback path testato esplicitamente
- File finale contiene il merge di tutti i parziali
- Heading H1 di tutte le macroaree presenti nel finale
- Subtitle Tab 2 preservato come prima riga del finale
- Anti path-traversal su session_id
- merge_summary aggrega correttamente
"""
from __future__ import annotations

import shutil
from pathlib import Path

import pytest
from docx import Document

from v2 import docx_merger as dm
from v2 import incremental_docx_builder as ib


# ──────────────────────────────────────────────────────────────────────────────
# Fixtures
# ──────────────────────────────────────────────────────────────────────────────

SAMPLE_PARSED = {
    "meta": {
        "azienda": {"nome": "DEMO MERGE SRL"},
        "audit": {"data_estrazione": "01/05/2026"},
    },
    "sezioni": [
        {
            "id": "08",
            "nome": "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA",
            "documenti": [
                {"tipo": "Visura", "titolo": "Visura 2025", "data_doc": "15/03/2025"},
            ],
        },
        {
            "id": "10",
            "nome": "10 · SICUREZZA SUL LAVORO",
            "documenti": [
                {"tipo": "DVR", "titolo": "DVR Rev 05", "data_doc": "10/01/2025"},
            ],
        },
    ],
}


@pytest.fixture
def populated_session(tmp_path):
    """Crea una session con header + 2 macroaree popolate."""
    session_id = "merge_test"
    ib.build_all_sections(SAMPLE_PARSED, session_id, base_dir=tmp_path)
    return session_id, tmp_path


# ──────────────────────────────────────────────────────────────────────────────
# Listing
# ──────────────────────────────────────────────────────────────────────────────

def test_list_partials_ordered_by_prefix(populated_session):
    session_id, base = populated_session
    paths = dm.list_session_partials(session_id, base_dir=base)
    assert len(paths) >= 3  # 00 + 01 + 03 (sicurezza è la 3a in ordine)
    # Verifica ordinamento numerico
    prefixes = [int(p.name.split("_")[0]) for p in paths]
    assert prefixes == sorted(prefixes)
    # Primo deve essere 00 (header)
    assert prefixes[0] == 0


def test_list_partials_skips_invalid_filenames(tmp_path):
    """File con pattern non valido vengono ignorati."""
    section_dir = ib._section_dir("xyz", base_dir=tmp_path)
    section_dir.mkdir(parents=True, exist_ok=True)

    valid = section_dir / "01_legale.docx"
    invalid = section_dir / "random_name.docx"
    not_docx = section_dir / "01_test.txt"
    valid.write_bytes(b"\x50\x4b\x03\x04")  # ZIP header (docx-like fake)
    invalid.write_bytes(b"\x50\x4b\x03\x04")
    not_docx.write_bytes(b"")

    paths = dm.list_session_partials("xyz", base_dir=tmp_path)
    names = [p.name for p in paths]
    assert "01_legale.docx" in names
    assert "random_name.docx" not in names
    assert "01_test.txt" not in names


def test_list_partials_invalid_session_returns_empty(tmp_path):
    paths = dm.list_session_partials("../traversal", base_dir=tmp_path)
    assert paths == []


def test_list_partials_no_directory(tmp_path):
    paths = dm.list_session_partials("never_built", base_dir=tmp_path)
    assert paths == []


# ──────────────────────────────────────────────────────────────────────────────
# Merge
# ──────────────────────────────────────────────────────────────────────────────

def test_merge_produces_valid_docx(populated_session):
    session_id, base = populated_session
    output = base / f"{session_id}_final.docx"

    result = dm.merge_session_sections(session_id, output, base_dir=base)
    assert result.success is True
    assert result.output_path == output
    assert output.exists()
    assert result.sections_merged >= 3
    assert result.file_size_bytes > 0


def test_merged_docx_can_be_reopened(populated_session):
    session_id, base = populated_session
    output = base / "out.docx"
    dm.merge_session_sections(session_id, output, base_dir=base)
    # Riapri con python-docx
    doc = Document(str(output))
    assert len(doc.paragraphs) > 1


def test_merged_docx_preserves_subtitle_for_tab2(populated_session):
    """La prima riga significativa del finale deve essere il titolo del header."""
    session_id, base = populated_session
    output = base / "out.docx"
    dm.merge_session_sections(session_id, output, base_dir=base)

    doc = Document(str(output))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Prima riga = titolo, seconda = subtitle "Audit - DEMO MERGE SRL"
    assert any("RELAZIONE DI EVIDENZE" in p for p in paras[:5])
    assert any("Audit - DEMO MERGE SRL" in p for p in paras[:5])


def test_merged_docx_contains_all_macroaree(populated_session):
    session_id, base = populated_session
    output = base / "out.docx"
    dm.merge_session_sections(session_id, output, base_dir=base)

    doc = Document(str(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Le 2 macroaree popolate nel sample
    assert "DOCUMENTAZIONE LEGALE E SOCIETARIA" in full_text
    assert "SICUREZZA SUL LAVORO" in full_text
    # Documenti specifici
    assert "Visura 2025" in full_text
    assert "DVR Rev 05" in full_text


# ──────────────────────────────────────────────────────────────────────────────
# Empty/error cases
# ──────────────────────────────────────────────────────────────────────────────

def test_merge_empty_session_fails_gracefully(tmp_path):
    output = tmp_path / "empty.docx"
    result = dm.merge_session_sections("never_built", output, base_dir=tmp_path)
    assert result.success is False
    assert result.error == "no_partials_found"
    assert not output.exists()


def test_merge_invalid_session_id(tmp_path):
    output = tmp_path / "bad.docx"
    result = dm.merge_session_sections("../etc", output, base_dir=tmp_path)
    assert result.success is False
    assert "non valido" in (result.error or "")


# ──────────────────────────────────────────────────────────────────────────────
# Force fallback
# ──────────────────────────────────────────────────────────────────────────────

def test_force_fallback_marks_used_fallback_true(populated_session):
    session_id, base = populated_session
    output = base / "fallback.docx"

    result = dm.merge_session_sections(
        session_id, output, base_dir=base, force_fallback=True,
    )
    assert result.success is True
    assert result.used_fallback is True
    assert output.exists()


def test_fallback_produces_readable_docx(populated_session):
    """Anche il fallback atomico deve produrre un docx valido."""
    session_id, base = populated_session
    output = base / "fallback.docx"

    dm.merge_session_sections(
        session_id, output, base_dir=base, force_fallback=True,
    )
    # Riapri con python-docx
    doc = Document(str(output))
    # Almeno qualche paragrafo
    assert len(doc.paragraphs) >= 2
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Contenuto del primo file presente
    assert "RELAZIONE DI EVIDENZE" in full_text or "DEMO MERGE SRL" in full_text


# ──────────────────────────────────────────────────────────────────────────────
# Summary
# ──────────────────────────────────────────────────────────────────────────────

def test_merge_summary(populated_session):
    session_id, base = populated_session
    output = base / "out.docx"
    result = dm.merge_session_sections(session_id, output, base_dir=base)

    summary = dm.merge_summary(result)
    assert summary["success"] is True
    assert summary["sections_merged"] >= 3
    assert summary["size_kb"] > 0
    assert summary["used_fallback"] is False


def test_merge_summary_failure(tmp_path):
    output = tmp_path / "x.docx"
    result = dm.merge_session_sections("missing", output, base_dir=tmp_path)
    summary = dm.merge_summary(result)
    assert summary["success"] is False
    assert summary["error"] is not None
