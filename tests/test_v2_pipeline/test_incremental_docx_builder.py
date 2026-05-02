"""
Test V2 Fase 7 — incremental_docx_builder.

Coperture:
- build_header_section produce 00_header.docx con titolo + subtitle critico per Tab 2
- Subtitle "Audit - {company_name}" è la prima riga significativa (Tab 2 compat)
- Macroaree → file 01..10 con prefix numerico, slug deterministico
- Macroarea senza documenti → success ma niente file creato
- Documenti distribuiti correttamente per match nome sezione
- Mai eccezione: parsed_data invalido → success=False con error
- Cleanup rimuove tutta la directory
- V2_KEEP_PARTIAL_DOCX=true preserva i file
- Anti path-traversal su session_id
- File generato apre correttamente con python-docx (sanity check)
"""
from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document

from v2 import incremental_docx_builder as ib


# ──────────────────────────────────────────────────────────────────────────────
# Sample data
# ──────────────────────────────────────────────────────────────────────────────

SAMPLE_PARSED = {
    "meta": {
        "azienda": {
            "nome": "DEMO CONSULTING SRL",
            "piva": "02012345678",
            "sede": "Via Roma 1, 10100 Torino (TO)",
        },
        "audit": {
            "data_estrazione": "01/05/2026",
            "docs_estratti": 5,
            "docs_vuoti": 0,
            "docs_analizzati": 5,
        },
        "indice": [
            {"n": 1, "tipo": "Visura", "titolo": "Visura 2025", "categoria": "08 LEGALE"},
            {"n": 2, "tipo": "DVR", "titolo": "DVR Rev 05", "categoria": "10 SSL"},
        ],
    },
    "sezioni": [
        {
            "id": "08",
            "nome": "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA",
            "documenti": [
                {
                    "tipo": "Visura Camerale",
                    "titolo": "Visura ordinaria 2025",
                    "categoria": "08 · LEGALE",
                    "riferimento": "REA TO-12345",
                    "data_doc": "15/03/2025",
                    "emesso_da": "CCIAA Torino",
                    "soggetto": "DEMO CONSULTING SRL",
                    "extra_info": "Capitale sociale 10.000 EUR",
                },
            ],
        },
        {
            "id": "10",
            "nome": "10 · SICUREZZA SUL LAVORO",
            "documenti": [
                {
                    "tipo": "Documento Valutazione Rischi",
                    "titolo": "DVR Revisione 05",
                    "categoria": "10 · SSL",
                    "data_doc": "10/01/2025",
                    "data_scadenza": "10/01/2026",
                    "emesso_da": "RSPP Mario Rossi",
                },
                {
                    "tipo": "POS",
                    "titolo": "Piano Operativo Sicurezza",
                    "categoria": "10 · SSL",
                    "data_doc": "20/02/2025",
                },
            ],
        },
    ],
}


# ──────────────────────────────────────────────────────────────────────────────
# Slugify
# ──────────────────────────────────────────────────────────────────────────────

def test_slugify_lowercase_no_spaces():
    assert ib._slugify("REGOLARITÀ CONTRIBUTIVA E FISCALE") == "regolarita_contributiva_e_fiscale"
    assert ib._slugify("ALTRO") == "altro"
    assert ib._slugify("GESTIONE MEZZI E ATTREZZATURE") == "gestione_mezzi_e_attrezzature"


def test_slugify_max_length():
    long = "A" * 200
    s = ib._slugify(long)
    assert len(s) <= 60


# ──────────────────────────────────────────────────────────────────────────────
# Validation session_id
# ──────────────────────────────────────────────────────────────────────────────

def test_invalid_session_id_rejected(tmp_path):
    result = ib.build_header_section(SAMPLE_PARSED, "../etc", base_dir=tmp_path)
    assert result.success is False
    assert "non valido" in (result.error or "")

    result = ib.build_macroarea_section(
        SAMPLE_PARSED, "ALTRO", 10, "foo/bar", base_dir=tmp_path,
    )
    assert result.success is False


# ──────────────────────────────────────────────────────────────────────────────
# Header section
# ──────────────────────────────────────────────────────────────────────────────

def test_header_section_creates_file(tmp_path):
    result = ib.build_header_section(
        SAMPLE_PARSED, "sess1", docs_estratti=5, docs_vuoti=0, base_dir=tmp_path,
    )
    assert result.success is True
    assert result.output_path is not None
    assert result.output_path.exists()
    assert result.output_path.name == "00_header.docx"
    assert result.section_index == 0
    assert result.file_size_bytes > 0


def test_header_subtitle_critical_for_tab2(tmp_path):
    """Verifica che il subtitle 'Audit - {company_name}' sia prima riga significativa."""
    result = ib.build_header_section(SAMPLE_PARSED, "sess2", base_dir=tmp_path)
    assert result.success is True

    doc = Document(str(result.output_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(paragraphs) >= 2

    # Prima riga = titolo, seconda = subtitle con company_name
    assert "RELAZIONE DI EVIDENZE" in paragraphs[0]
    assert "Audit - DEMO CONSULTING SRL" in paragraphs[1]


def test_header_with_unknown_company(tmp_path):
    """Senza nome azienda → AZIENDA NON IDENTIFICATA."""
    data = {"meta": {"azienda": {}, "audit": {}, "indice": []}, "sezioni": []}
    result = ib.build_header_section(data, "sessunk", base_dir=tmp_path)
    assert result.success is True
    doc = Document(str(result.output_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AZIENDA NON IDENTIFICATA" in text


def test_header_invalid_parsed_data(tmp_path):
    """parsed_data non dict → success=False, no eccezione."""
    result = ib.build_header_section("not a dict", "x", base_dir=tmp_path)  # type: ignore
    assert result.success is False
    assert result.error == "parsed_data_not_dict"


# ──────────────────────────────────────────────────────────────────────────────
# Macroarea section
# ──────────────────────────────────────────────────────────────────────────────

def test_macroarea_with_documents_creates_file(tmp_path):
    result = ib.build_macroarea_section(
        SAMPLE_PARSED,
        "DOCUMENTAZIONE LEGALE E SOCIETARIA",
        section_index=1,
        session_id="sess3",
        base_dir=tmp_path,
    )
    assert result.success is True
    assert result.output_path is not None
    assert result.output_path.exists()
    assert result.output_path.name.startswith("01_")
    assert result.documents_count == 1


def test_macroarea_without_documents_no_file(tmp_path):
    """Macroarea senza documenti → success ma niente file.

    Uso "13 · ESG E RENDICONTAZIONE" (categoria V2) che non ha documenti
    nel SAMPLE_PARSED.
    """
    result = ib.build_macroarea_section(
        SAMPLE_PARSED,
        "13 · ESG E RENDICONTAZIONE",
        section_index=13,
        session_id="sess4",
        base_dir=tmp_path,
    )
    assert result.success is True
    assert result.output_path is None
    assert result.documents_count == 0


def test_macroarea_renders_all_documents(tmp_path):
    """SSL ha 2 documenti: entrambi devono apparire nel docx generato."""
    result = ib.build_macroarea_section(
        SAMPLE_PARSED,
        "SICUREZZA SUL LAVORO",
        section_index=3,
        session_id="sess5",
        base_dir=tmp_path,
    )
    assert result.success is True
    assert result.documents_count == 2

    doc = Document(str(result.output_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "DVR Revisione 05" in text
    assert "Piano Operativo Sicurezza" in text


def test_macroarea_extra_fields_rendered(tmp_path):
    """I campi extra (oltre header standard) finiscono in tabella secondaria."""
    result = ib.build_macroarea_section(
        SAMPLE_PARSED,
        "DOCUMENTAZIONE LEGALE E SOCIETARIA",
        section_index=1,
        session_id="sess6",
        base_dir=tmp_path,
    )
    doc = Document(str(result.output_path))
    # Estraggo testo dalle tabelle
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += "\n" + cell.text
    assert "Capitale sociale 10.000 EUR" in table_text


# ──────────────────────────────────────────────────────────────────────────────
# Build all sections
# ──────────────────────────────────────────────────────────────────────────────

def test_build_all_sections_includes_header_and_18_macroaree(tmp_path):
    """Allineato al prompt universale V2: 1 header + 18 macroaree = 19."""
    results = ib.build_all_sections(
        SAMPLE_PARSED, "sessall", docs_estratti=3, docs_vuoti=0, base_dir=tmp_path,
    )
    # 1 header + 18 macroaree = 19 risultati
    assert len(results) == 19
    assert results[0].section_index == 0
    assert results[0].section_name == "header"
    # I successivi sono in MACROAREA_ORDER (18 voci, formato "NN · NOME")
    assert len(ib.MACROAREA_ORDER) == 18
    for i, expected in enumerate(ib.MACROAREA_ORDER, start=1):
        assert results[i].section_name == expected


def test_build_all_sections_creates_only_populated_sections(tmp_path):
    results = ib.build_all_sections(SAMPLE_PARSED, "sessp", base_dir=tmp_path)
    files_created = [r for r in results if r.output_path is not None]
    # header + 2 macroaree popolate (LEGALE + SSL)
    assert len(files_created) == 3


# ──────────────────────────────────────────────────────────────────────────────
# Cleanup
# ──────────────────────────────────────────────────────────────────────────────

def test_cleanup_removes_all_partials(tmp_path):
    ib.build_all_sections(SAMPLE_PARSED, "cleanme", base_dir=tmp_path)
    section_dir = ib._section_dir("cleanme", base_dir=tmp_path)
    assert section_dir.exists()
    assert any(section_dir.glob("*.docx"))

    count = ib.cleanup_session_sections("cleanme", base_dir=tmp_path)
    assert count > 0
    assert not section_dir.exists()


def test_cleanup_idempotent_on_missing_dir(tmp_path):
    """Cleanup su session inesistente → 0, no crash."""
    count = ib.cleanup_session_sections("noexist", base_dir=tmp_path)
    assert count == 0


def test_cleanup_skipped_when_keep_env_true(tmp_path, monkeypatch):
    """Con V2_KEEP_PARTIAL_DOCX=true, cleanup è no-op."""
    monkeypatch.setenv("V2_KEEP_PARTIAL_DOCX", "true")
    ib.build_all_sections(SAMPLE_PARSED, "keepme", base_dir=tmp_path)
    section_dir = ib._section_dir("keepme", base_dir=tmp_path)

    count = ib.cleanup_session_sections("keepme", base_dir=tmp_path)
    assert count == 0
    assert section_dir.exists()  # Files ancora presenti


def test_cleanup_invalid_session_id(tmp_path):
    count = ib.cleanup_session_sections("../escape", base_dir=tmp_path)
    assert count == 0


# ──────────────────────────────────────────────────────────────────────────────
# Summary
# ──────────────────────────────────────────────────────────────────────────────

def test_builder_summary_aggregates(tmp_path):
    results = ib.build_all_sections(SAMPLE_PARSED, "summa", base_dir=tmp_path)
    summary = ib.builder_summary(results)
    # 1 header + 18 macroaree = 19 risultati; tutti success
    assert summary["total_sections"] == 19
    assert summary["success"] == 19
    assert summary["failed"] == 0
    # Solo 3 file creati: header + 08 LEGALE + 10 SSL (le altre macroaree
    # sono vuote per il SAMPLE_PARSED)
    assert summary["files_created"] == 3


def test_builder_summary_empty():
    s = ib.builder_summary([])
    assert s["total"] == 0


# ──────────────────────────────────────────────────────────────────────────────
# Sanity: il docx generato è ben formato
# ──────────────────────────────────────────────────────────────────────────────

def test_generated_docx_can_be_reopened(tmp_path):
    """Tutti i .docx generati devono poter essere riaperti senza errori."""
    results = ib.build_all_sections(SAMPLE_PARSED, "reopen", base_dir=tmp_path)
    files = [r.output_path for r in results if r.output_path is not None]
    assert len(files) >= 1

    for path in files:
        # Deve aprirsi senza eccezione
        doc = Document(str(path))
        # Almeno 1 paragrafo
        assert len(doc.paragraphs) >= 1


def test_no_mutation_of_input(tmp_path):
    """parsed_data non viene mutato dal builder."""
    import copy
    snapshot = copy.deepcopy(SAMPLE_PARSED)
    ib.build_all_sections(SAMPLE_PARSED, "nomut", base_dir=tmp_path)
    assert SAMPLE_PARSED == snapshot


# ──────────────────────────────────────────────────────────────────────────────
# Fix routing macroarea V2 — match per codice numerico + retro-compat
# ──────────────────────────────────────────────────────────────────────────────

def test_extract_category_code_canonical_format():
    """Formato canonico `NN · NOME` → codice zero-padded."""
    assert ib._extract_category_code("08 · DOCUMENTAZIONE LEGALE E SOCIETARIA") == "08"
    assert ib._extract_category_code("10 · SALUTE E SICUREZZA SUL LAVORO") == "10"
    assert ib._extract_category_code("01 · CONTESTO E PARTI INTERESSATE") == "01"
    assert ib._extract_category_code("18 · ALTRI") == "18"


def test_extract_category_code_variants_tolerated():
    """Varianti del modello (abbreviate, non zero-padded, separatori variabili)."""
    # Varianti tipo "08 · LEGALE/SOCIETARIA" o "08 · LEGALE" → comunque "08"
    assert ib._extract_category_code("08 · LEGALE/SOCIETARIA") == "08"
    assert ib._extract_category_code("08·LEGALE") == "08"
    assert ib._extract_category_code("08:LEGALE") == "08"
    # Numero senza zero-padding
    assert ib._extract_category_code("8 LEGALE") == "08"
    assert ib._extract_category_code("1 CONTESTO") == "01"
    # Spazi iniziali
    assert ib._extract_category_code("  10 · SSL") == "10"


def test_extract_category_code_invalid_returns_none():
    """Stringhe senza prefisso numerico o codice fuori range → None."""
    assert ib._extract_category_code("DOCUMENTAZIONE LEGALE E SOCIETARIA") is None
    assert ib._extract_category_code("LEGALE 08") is None  # numero non in testa
    assert ib._extract_category_code("19 · INVENTATA") is None  # fuori range 1-18
    assert ib._extract_category_code("99 · IMPOSSIBILE") is None
    assert ib._extract_category_code("0 · ZERO") is None
    assert ib._extract_category_code("") is None
    assert ib._extract_category_code(None) is None  # type: ignore


def test_docs_in_macroarea_match_by_code_when_model_writes_variant():
    """
    INVARIANTE CRITICA: se il modello scrive `08 · LEGALE/SOCIETARIA`
    (variante abbreviata), la VISURA deve essere comunque routata alla
    macroarea `08 · DOCUMENTAZIONE LEGALE E SOCIETARIA` via match codice.
    """
    parsed = {
        "sezioni": [
            {
                "id": "08",
                "nome": "08 · LEGALE/SOCIETARIA",  # variante abbreviata del modello
                "documenti": [
                    {
                        "tipo": "Visura Camerale",
                        "titolo": "Visura MEDIL",
                        "categoria": "08 · LEGALE/SOCIETARIA",
                        "data_doc": "10/08/2025",
                    },
                ],
            },
        ],
    }
    docs = ib._docs_in_macroarea(parsed, "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA")
    assert len(docs) == 1
    assert docs[0]["tipo"] == "Visura Camerale"


def test_docs_in_macroarea_match_by_code_with_short_number():
    """`8` invece di `08` → comunque match a 08."""
    parsed = {
        "sezioni": [
            {"nome": "8 · CIAO", "documenti": [{"tipo": "Visura"}]},
        ],
    }
    docs = ib._docs_in_macroarea(parsed, "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA")
    assert len(docs) == 1


def test_docs_in_macroarea_substring_fallback_retro_compat():
    """
    Retro-compat: macroarea senza prefisso (test legacy) usa substring match.
    """
    parsed = {
        "sezioni": [
            {
                "nome": "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA",
                "documenti": [{"tipo": "Visura"}],
            },
        ],
    }
    # macroarea senza prefisso → fallback substring → match
    docs = ib._docs_in_macroarea(parsed, "DOCUMENTAZIONE LEGALE E SOCIETARIA")
    assert len(docs) == 1


def test_docs_in_macroarea_does_not_cross_contaminate():
    """
    Documenti di codice 08 NON devono finire in macroarea 10 e viceversa.
    """
    parsed = {
        "sezioni": [
            {
                "nome": "08 · LEGALE",
                "documenti": [{"tipo": "Visura"}],
            },
            {
                "nome": "10 · SSL",
                "documenti": [{"tipo": "DVR"}],
            },
        ],
    }
    legale = ib._docs_in_macroarea(parsed, "08 · DOCUMENTAZIONE LEGALE E SOCIETARIA")
    ssl = ib._docs_in_macroarea(parsed, "10 · SALUTE E SICUREZZA SUL LAVORO")
    assert {d["tipo"] for d in legale} == {"Visura"}
    assert {d["tipo"] for d in ssl} == {"DVR"}


def test_docs_in_macroarea_invalid_section_code_falls_through():
    """
    Sezione con codice fuori range (es. "19 · INVENTATA") cade nel fallback
    substring. Se il nome non matcha, documento NON finisce in nessuna
    macroarea (verrà perso, ma è il modello che ha sbagliato).
    """
    parsed = {
        "sezioni": [
            {"nome": "19 · INVENTATA", "documenti": [{"tipo": "Strano"}]},
        ],
    }
    # Cerca in macroarea standard: nessun match
    for m in ("08 · DOCUMENTAZIONE LEGALE E SOCIETARIA",
              "10 · SALUTE E SICUREZZA SUL LAVORO",
              "18 · ALTRI"):
        docs = ib._docs_in_macroarea(parsed, m)
        assert docs == []


def test_category_code_to_macroarea_mapping_complete():
    """Tutte le 18 categorie 01-18 hanno un mapping macroarea."""
    assert len(ib.CATEGORY_CODE_TO_MACROAREA) == 18
    for n in range(1, 19):
        code = f"{n:02d}"
        assert code in ib.CATEGORY_CODE_TO_MACROAREA
        assert ib.CATEGORY_CODE_TO_MACROAREA[code].startswith(f"{code} · ")
