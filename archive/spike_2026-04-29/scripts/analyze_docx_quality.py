#!/usr/bin/env python3
"""
Analisi qualitativa dei docx prodotti dai 3 provider.
Conta schede, sezioni, parole, cluster, e pattern audit-critical.

Usage:
    python scripts/analyze_docx_quality.py
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document  # python-docx
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(2)

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore
except Exception:
    pass

_REPO_ROOT = Path(__file__).resolve().parent.parent
TEST_ROOT = _REPO_ROOT / "testllmnew"


def find_docx(provider: str) -> Path | None:
    """Ritorna il docx più recente (mtime) per il provider."""
    docx_dir = TEST_ROOT / provider / "docx_outputs"
    if not docx_dir.exists():
        return None
    files = sorted(docx_dir.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


# Pattern audit-critical che cerchiamo nei docx
AUDIT_PATTERNS = {
    "RSPP": re.compile(r"\bRSPP\b|\bResponsabile\s+Servizio\s+Prevenzione", re.IGNORECASE),
    "Medico_Comp": re.compile(r"\bMedico\s+Competente\b", re.IGNORECASE),
    "DVR": re.compile(r"\bDVR\b|\bValutazione\s+(?:dei\s+)?Rischi\b", re.IGNORECASE),
    "DUVRI": re.compile(r"\bDUVRI\b", re.IGNORECASE),
    "POS": re.compile(r"\bPOS\b|\bPiano\s+Operativo\s+(?:di\s+)?Sicurezza", re.IGNORECASE),
    "PSC": re.compile(r"\bPSC\b|\bPiano\s+(?:di\s+)?Sicurezza\s+(?:e\s+)?Coordinamento", re.IGNORECASE),
    "SOA": re.compile(r"\bSOA\b|\battestazione\s+SOA", re.IGNORECASE),
    "Visura": re.compile(r"\bVisura\s+Camerale|\bVisura\s+CCIAA|\bRegistro\s+Imprese\b", re.IGNORECASE),
    "ISO9001": re.compile(r"\bISO\s*9001\b|\bSistema\s+(?:di\s+)?Gestione\s+(?:per\s+la\s+)?Qualit", re.IGNORECASE),
    "ISO14001": re.compile(r"\bISO\s*14001\b|\bSistema\s+(?:di\s+)?Gestione\s+Ambient", re.IGNORECASE),
    "ISO45001": re.compile(r"\bISO\s*45001\b|\bOHSAS", re.IGNORECASE),
    "ISO37001": re.compile(r"\bISO\s*37001\b|\bAnticorruzione\b", re.IGNORECASE),
    "ISO50001": re.compile(r"\bISO\s*50001\b|\bAnalisi\s+Energetica\b|\bEnergy\s+Manager\b", re.IGNORECASE),
    "DURC": re.compile(r"\bDURC\b", re.IGNORECASE),
    "Fattura": re.compile(r"\bfattura\b", re.IGNORECASE),
    "Contratto": re.compile(r"\bcontratto\b", re.IGNORECASE),
    "Estintore": re.compile(r"\bestintor", re.IGNORECASE),
    "Manutenzione": re.compile(r"\bmanutenzion", re.IGNORECASE),
    "Carburante": re.compile(r"\bcarburant", re.IGNORECASE),
}


# Marker per identificare le "schede" nel docx
SCHEDA_MARKERS = [
    re.compile(r"^─+\s*DOC\s+\d+\s*─+\s*$"),     # ── DOC N ──
    re.compile(r"^Documento\s+\d+\s*[:\-]"),      # Documento 1:
    re.compile(r"^#+\s*DOC\s+\d+", re.IGNORECASE),
]


def analyze_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    full_text_with_tables = full_text + "\n"
    # Aggiungi anche il testo delle tabelle
    n_tables = len(doc.tables)
    n_table_rows = 0
    for tbl in doc.tables:
        n_table_rows += len(tbl.rows)
        for row in tbl.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    full_text_with_tables += ct + "\n"

    # Conta sezioni heading
    n_headings = 0
    heading_styles = {}
    for p in doc.paragraphs:
        if p.style and p.style.name and "Heading" in p.style.name:
            n_headings += 1
            heading_styles[p.style.name] = heading_styles.get(p.style.name, 0) + 1

    # Conta marker schede
    n_schede = 0
    for line in full_text_with_tables.splitlines():
        for marker in SCHEDA_MARKERS:
            if marker.search(line):
                n_schede += 1
                break

    # Audit-critical pattern matching
    audit_hits = {}
    for name, pat in AUDIT_PATTERNS.items():
        audit_hits[name] = len(pat.findall(full_text_with_tables))

    # Word count
    n_chars = len(full_text_with_tables)
    n_words = len(full_text_with_tables.split())

    return {
        "path": str(path),
        "size_kb": round(path.stat().st_size / 1024, 1),
        "n_paragraphs": len(paragraphs),
        "n_headings": n_headings,
        "heading_styles": heading_styles,
        "n_tables": n_tables,
        "n_table_rows": n_table_rows,
        "n_schede_markers": n_schede,
        "n_chars": n_chars,
        "n_words": n_words,
        "audit_hits": audit_hits,
    }


def main():
    providers = ["gemini-baseline", "deepseek-v4-flash", "gpt-4.1-mini"]
    results = {}

    print("=" * 90)
    print("ANALISI QUALITATIVA DOCX OUTPUT")
    print("=" * 90)
    print()

    for p in providers:
        docx_path = find_docx(p)
        if docx_path is None:
            print(f"  {p:<22}: (no docx)")
            continue
        try:
            stats = analyze_docx(docx_path)
            results[p] = stats
        except Exception as e:
            print(f"  {p:<22}: ERROR {e}")
            continue

    # Tabella struttura
    print("STRUTTURA DOCX")
    print("-" * 90)
    print(f"{'Provider':<22} {'Size KB':>8} {'Para':>6} {'Tables':>7} {'TblRows':>8} {'Headings':>9} {'Words':>8} {'Schede*':>8}")
    print("-" * 90)
    for p in providers:
        if p not in results:
            continue
        s = results[p]
        print(f"{p:<22} {s['size_kb']:>8.1f} {s['n_paragraphs']:>6} {s['n_tables']:>7} {s['n_table_rows']:>8} {s['n_headings']:>9} {s['n_words']:>8} {s['n_schede_markers']:>8}")
    print()
    print("(*) Schede = match esplicito di marker '── DOC N ──' o 'Documento N:'. Se 0, le schede sono in altro formato.")
    print()

    # Tabella audit-critical
    print("PATTERN AUDIT-CRITICAL — frequenza menzioni nel docx")
    print("-" * 90)
    pattern_names = list(AUDIT_PATTERNS.keys())
    header = f"{'Pattern':<14}" + "".join(f"{p[:18]:>20}" for p in providers if p in results)
    print(header)
    print("-" * len(header))
    for pat in pattern_names:
        line = f"{pat:<14}"
        for p in providers:
            if p not in results:
                continue
            v = results[p]['audit_hits'].get(pat, 0)
            mark = " ✓" if v > 0 else "  "
            line += f"{v}{mark}".rjust(20)
        print(line)
    print()

    # Sintesi
    print("SINTESI COMPARATIVA")
    print("-" * 90)
    if all(p in results for p in providers):
        gem = results["gemini-baseline"]
        ds = results["deepseek-v4-flash"]
        gpt = results["gpt-4.1-mini"]
        ratio_ds = ds['n_words'] / max(1, gem['n_words']) * 100
        ratio_gpt = gpt['n_words'] / max(1, gem['n_words']) * 100
        print(f"  Words: gemini={gem['n_words']:,}, deepseek={ds['n_words']:,} ({ratio_ds:.1f}% di gemini), gpt-4.1={gpt['n_words']:,} ({ratio_gpt:.1f}% di gemini)")

        gem_audit = sum(1 for v in gem['audit_hits'].values() if v > 0)
        ds_audit = sum(1 for v in ds['audit_hits'].values() if v > 0)
        gpt_audit = sum(1 for v in gpt['audit_hits'].values() if v > 0)
        n_pat = len(AUDIT_PATTERNS)
        print(f"  Audit pattern coverage (di {n_pat}): gemini={gem_audit}, deepseek={ds_audit}, gpt-4.1={gpt_audit}")
    print()

    # Salva summary JSON
    out_path = TEST_ROOT / "quality_analysis.json"
    out_path.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
