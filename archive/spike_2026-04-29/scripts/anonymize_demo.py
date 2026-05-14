"""
Anonimizzazione del file Audit_V2_NAR_TEST_9001_20260511_074019.docx
per creare una versione demo da mostrare a un cliente potenziale.

Sostituisce tutti i nomi, ragioni sociali, indirizzi, persone fisiche,
partite IVA, contatti email, riferimenti geografici riconducibili al
cliente reale CO.SAL. S.R.L. con dati fittizi neutri e coerenti.

Mantiene struttura, formattazione, ordine paragrafi, macroaree.
"""
from docx import Document
import re

SRC = r"C:\Users\user\AUDITORSEMI\Audit_V2_NAR_TEST_9001_20260511_074019.docx"
DST = r"C:\Users\user\AUDITORSEMI\Audit_V2_NAR_DEMO_CLIENT_20260511.docx"

# Coppie di sostituzione: ORIGINALE -> FITTIZIO
# Ordinate per LUNGHEZZA DECRESCENTE per evitare match parziali (es. "CO.SAL" prima di "COSAL").
REPLACEMENTS = [
    # === Azienda principale (tutte le varianti grafiche presenti nel docx) ===
    ("CO.SAL. S.R.L.",                "ALFA SERVIZI S.R.L."),
    ("CO.SAL. SRL",                   "ALFA SERVIZI S.R.L."),
    ("CO. SAL. S.r.l.",               "ALFA SERVIZI S.R.L."),
    ("CO.SAL S.R.L.",                 "ALFA SERVIZI S.R.L."),
    ("CO.SAL.",                       "ALFA SERVIZI"),
    ("COSAL",                         "ALFASERVIZI"),
    ("cosalservice",                  "alfaserviziconsulting"),

    # === Persone fisiche (titolari, soci, RSPP, RLS, formatori, dipendenti) ===
    ("GUGLIELMO PELUSO",              "MARIO ROSSI"),
    ("COSTANTINO PELUSO",             "GIOVANNI BIANCHI"),
    ("Guglielmo Peluso",              "Mario Rossi"),
    ("Costantino Peluso",             "Giovanni Bianchi"),
    ("CATERINA NOCERA",               "LAURA VERDI"),
    ("Dott.ssa CATERINA NOCERA",      "Dott.ssa LAURA VERDI"),
    ("ANNA SPADA",                    "ELENA NERI"),
    ("Anna Spada",                    "Elena Neri"),
    ("EGISTO FABBRICINI",             "MARCO GIALLI"),
    ("VALENTINA CASILLO",             "SOFIA ROSSI"),
    ("BHIYAN RATAN",                  "DIPENDENTE 01"),
    ("Daniela Candura",               "Operatore Selezione"),
    ("Anna Ragosta",                  "Tirocinante 1"),
    ("Domenico Scialo",               "Tirocinante 2"),
    ("valentina@cosalservice.it",     "amministrazione@alfaserviziconsulting.it"),
    ("medicinadellavoro@cosalservice.it",
                                      "medicinadellavoro@alfaserviziconsulting.it"),

    # === Indirizzi e località riconducibili ===
    ("Via Sodano, 46",                "Via dei Tigli, 10"),
    ("Via Sodano 46",                 "Via dei Tigli 10"),
    ("Via Sodano n°46",               "Via dei Tigli n°10"),
    ("Via Sodano",                    "Via dei Tigli"),
    ("Sarno (SA)",                    "Comune Esempio (XX)"),
    ("Sarno",                         "Comune Esempio"),
    ("84087",                         "00000"),
    ("VIA GAETANO MALASOMA 18, Pisa", "Via Esempio 1, Città Esempio"),
    ("VIA GAETANO MALASOMA 18",       "Via Esempio 1"),
    ("Pisa",                          "Città Esempio"),
    ("Via Boscofangone Zona Industriale ASI SNC",
                                      "Sede Formazione Esempio"),
    ("Via Boscofangone Zona Industriale ASI snc",
                                      "Sede Formazione Esempio"),
    ("Via Curti 8",                   "Via Esempio Trasporti 1"),
    ("Via Matteotti 68",              "Via Esempio Cartoleria 1"),
    ("Nola (NA)",                     "Sede Esempio (XX)"),
    ("Salerno",                       "Provincia Esempio"),
    ("SA-333626",                     "XX-000000"),
    ("Roma, Via Piave 24",            "Sede Esempio Ente Formazione"),

    # === Aziende terze citate (fornitori, clienti, organismi) ===
    ("Medialis srl",                  "Software Esempio S.R.L."),
    ("Medialis",                      "Software Esempio"),
    ("ACHILLE",                       "Gestionale Esempio"),
    ("Achille",                       "Gestionale Esempio"),
    ("StepOver",                      "Tavoletta Esempio"),
    ("STEP OVER",                     "Tavoletta Esempio"),
    ("ESQ Cert Ltd",                  "Organismo Certificazione Esempio"),
    ("ESQ Cert",                      "Organismo Certificazione Esempio"),
    ("KAELA S.R.L.",                  "BETA LOGISTICA S.R.L."),
    ("KAELA SRL",                     "BETA LOGISTICA S.R.L."),
    ("NU SA TRASPORTI S.R.L.",        "GAMMA TRASPORTI S.R.L."),
    ("NU.SA. TRASPORTI SRL",          "GAMMA TRASPORTI S.R.L."),
    ("NU.SA. Trasporti",              "Gamma Trasporti"),
    ("NU SA TRASPORTI",               "GAMMA TRASPORTI"),
    ("CARTUFFICIO SASDI FERRARA DONATELLA & C",
                                      "CARTOLERIA ESEMPIO S.A.S."),
    ("Cartufficio Sasdi Ferrara Donatella & C",
                                      "Cartoleria Esempio S.a.s."),
    ("EUROFIRE",                      "ANTINCENDIO ESEMPIO S.R.L."),
    ("DE BRICO CHIMICA SRL",          "CHIMICA ESEMPIO S.R.L."),
    ("Adecco Lifescience",            "Agenzia Selezione Esempio"),
    ("CERTITALY SRL",                 "ENTE CERT ESEMPIO S.R.L."),
    ("O.P.N. ITALIA LAVORO",          "ENTE FORMAZIONE ESEMPIO"),
    ("INTESA SANPAOLO SPA",           "ISTITUTO BANCARIO ESEMPIO S.P.A."),
    ("BCITITMMXXX",                   "XXXXITMMXXX"),
    ("Università degli Studi di Salerno",
                                      "Università Esempio"),
    ("Ministero della Pubblica Istruzione",
                                      "Pubblica Amministrazione Esempio"),
    ("Ordine degli Ingegneri di Salerno",
                                      "Ordine Professionale Esempio"),
    ("Ordine degli Ingegneri di Provincia Esempio",
                                      "Ordine Professionale Esempio"),
    ("Ordine degli Ingegneri",
                                      "Ordine Professionale Esempio"),
    ("Prodal S.c.ar.l.",              "Consorzio Esempio S.c.a r.l."),

    # === Riferimenti numerici/protocolli specifici ===
    ("5157/A",                        "0000/X"),
    ("CFAL/459629",                   "CFAL/000000"),
    ("prot. CFAL/459629",             "prot. CFAL/000000"),
    ("protocollo n. 459629",          "protocollo n. 000000"),
    ("n. 151",                        "n. XXX"),
    ("FD 000075",                     "FD XXXXXX"),
    ("n. 2/96",                       "n. XX/XX"),
    ("n. 1/2025",                     "n. XX/2025"),
    ("n. 1 del 28/01/2025",           "n. XX del 28/01/2025"),
    ("n. 2 del 12/02/2025",           "n. XX del 12/02/2025"),
    ("n. 277 del 29/03/2025",         "n. XXX del 29/03/2025"),
    ("2048/2025",                     "XXXX/2025"),

    # === Titolo del documento ===
    ("Audit - TEST 9001",             "Audit - DEMO CLIENT"),
]


def replace_in_text(text: str) -> str:
    if not text:
        return text
    out = text
    for old, new in REPLACEMENTS:
        if old in out:
            out = out.replace(old, new)
    return out


def replace_in_runs(paragraph):
    """Sostituisce nei run del paragrafo preservando la formattazione."""
    # Strategia robusta: concatena tutto il testo dei run, applica sostituzioni
    # e rimette tutto nel primo run azzerando gli altri. Funziona per testo
    # corrente; se un singolo paragrafo ha formattazioni miste su porzioni
    # diverse, perdiamo la granularità ma manteniamo lo stile del primo run.
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    new_full = replace_in_text(full)
    if new_full == full:
        return
    paragraph.runs[0].text = new_full
    for r in paragraph.runs[1:]:
        r.text = ""


def process_document(src, dst):
    doc = Document(src)
    # Paragrafi del corpo
    for p in doc.paragraphs:
        replace_in_runs(p)
    # Tabelle (a tutti i livelli annidati)
    def process_cells(cells):
        for cell in cells:
            for p in cell.paragraphs:
                replace_in_runs(p)
            for nested in cell.tables:
                for row in nested.rows:
                    process_cells(row.cells)
    for table in doc.tables:
        for row in table.rows:
            process_cells(row.cells)
    # Header e footer
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                replace_in_runs(p)
            for table in hf.tables:
                for row in table.rows:
                    process_cells(row.cells)
    doc.save(dst)


if __name__ == "__main__":
    process_document(SRC, DST)
    print(f"Salvato: {DST}")
