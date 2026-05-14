#!/usr/bin/env python3
"""
Analisi qualitativa DETTAGLIATA dei docx prodotti.
Estrae per ogni docx:
- Conteggio schede effettive
- Distribuzione categorie
- Distribuzione tipi documento
- Cross-check con file ZIP originale (quali file presenti / quali persi)
- Sample schede (prime 3) per ispezione manuale

Usage:
    python scripts/analyze_docx_detailed.py [zip_path]
"""
from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(2)

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore
except Exception:
    pass

_REPO_ROOT = Path(__file__).resolve().parent.parent
TEST_ROOT = _REPO_ROOT / "testllmnew"


def find_docx(provider: str) -> Path | None:
    docx_dir = TEST_ROOT / provider / "docx_outputs"
    if not docx_dir.exists():
        return None
    files = sorted(docx_dir.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def list_zip_files(zip_path: Path) -> list[str]:
    """Lista dei file analizzabili dentro lo ZIP (PDF, DOCX, XLSX, P7M)."""
    if not zip_path.exists():
        return []
    out = []
    with zipfile.ZipFile(zip_path, "r") as zf:
        for n in zf.namelist():
            low = n.lower()
            if low.endswith(("/", "\\")):
                continue
            if low.endswith((".pdf", ".docx", ".xlsx", ".p7m", ".doc", ".xls")):
                out.append(n)
    return out


def extract_text_from_docx(docx_path: Path) -> str:
    """Estrae tutto il testo (paragraphs + tables) come stringa unica."""
    doc = Document(docx_path)
    chunks = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    chunks.append(ct)
    return "\n".join(chunks)


def parse_schede(text: str) -> list[dict]:
    """
    Estrae le schede dal testo del docx.
    Una scheda è il blocco testo che inizia con un heading "tipo:" o pattern simili
    fino al successivo separator. Per V2 i docx hanno tipicamente sezioni con
    heading per ogni documento. Cerchiamo pattern multipli.
    """
    schede = []

    # Pattern 1: blocchi YAML con `tipo: "..."` e `categoria: "..."`
    # Cerco tutte le coppie tipo+categoria nel testo aggregato
    tipo_pattern = re.compile(
        r'tipo\s*[:\:]\s*[\'"]?([^\'"\n]{3,80})[\'"]?\s*(?:\n|\|)',
        re.IGNORECASE,
    )
    cat_pattern = re.compile(
        r'categoria\s*[:\:]\s*[\'"]?(\d{2}\s*[\-·]\s*[A-Za-z][^\'"\n]{2,60})[\'"]?\s*(?:\n|\|)',
        re.IGNORECASE,
    )
    titolo_pattern = re.compile(
        r'titolo\s*[:\:]\s*[\'"]?([^\'"\n]{3,200})[\'"]?\s*(?:\n|\|)',
        re.IGNORECASE,
    )

    tipi = [m.group(1).strip() for m in tipo_pattern.finditer(text)]
    categorie = [m.group(1).strip() for m in cat_pattern.finditer(text)]
    titoli = [m.group(1).strip() for m in titolo_pattern.finditer(text)]

    return {
        "tipi": tipi,
        "categorie": categorie,
        "titoli": titoli,
    }


def normalize_filename(s: str) -> str:
    """Normalizza per confronto fuzzy: lowercase, no estensione, no prefisso path."""
    name = Path(s).stem.lower()
    name = re.sub(r"[^a-z0-9]+", "", name)
    return name


def cross_check_files(zip_files: list[str], docx_text: str) -> dict:
    """
    Per ogni file nel ZIP, verifica se appare nel testo del docx.
    Match fuzzy: prendiamo i primi 6+ caratteri significativi del nome file.
    """
    docx_lower = docx_text.lower()
    found = []
    missing = []
    for zf in zip_files:
        stem = Path(zf).stem.lower()
        # Tokenizza: prendiamo le parole significative (>= 4 caratteri)
        tokens = [t for t in re.split(r"[\W_]+", stem) if len(t) >= 4]
        if not tokens:
            continue
        # Almeno 2 token devono apparire nel docx (o 1 se il file ha solo 1 token significativo)
        threshold = min(2, len(tokens))
        hits = sum(1 for t in tokens if t in docx_lower)
        if hits >= threshold:
            found.append(zf)
        else:
            missing.append(zf)
    return {
        "n_zip_files": len(zip_files),
        "n_found": len(found),
        "n_missing": len(missing),
        "missing_sample": missing[:20],  # primi 20 mancanti per debug
    }


def analyze(provider: str, zip_files: list[str]) -> dict:
    docx_path = find_docx(provider)
    if docx_path is None:
        return {"error": "no_docx"}

    text = extract_text_from_docx(docx_path)
    schede = parse_schede(text)
    cross = cross_check_files(zip_files, text)

    cat_dist = Counter(schede["categorie"])
    tipo_dist = Counter(schede["tipi"])

    return {
        "docx_path": str(docx_path),
        "docx_size_kb": round(docx_path.stat().st_size / 1024, 1),
        "n_chars": len(text),
        "n_words": len(text.split()),
        "n_tipi_extracted": len(schede["tipi"]),
        "n_categorie_extracted": len(schede["categorie"]),
        "n_titoli_extracted": len(schede["titoli"]),
        "n_tipi_unici": len(set(schede["tipi"])),
        "categoria_distribution": dict(cat_dist.most_common(20)),
        "tipo_top10": dict(tipo_dist.most_common(10)),
        "cross_check": cross,
        "sample_titoli": schede["titoli"][:5],
    }


def main():
    zip_path = Path(sys.argv[1]) if len(sys.argv) > 1 else _REPO_ROOT / "ALLEGATI MEDIL 37001_50001.zip"
    zip_files = list_zip_files(zip_path)
    print(f"ZIP: {zip_path.name} → {len(zip_files)} file analizzabili (PDF/DOCX/XLSX/P7M)")
    print()

    providers = ["gemini-baseline", "deepseek-v4-flash", "gpt-4.1-mini"]
    results = {}

    for p in providers:
        print(f"=== {p} ===")
        r = analyze(p, zip_files)
        results[p] = r
        if r.get("error"):
            print(f"  {r['error']}")
            continue

        print(f"  Docx           : {Path(r['docx_path']).name}")
        print(f"  Size KB        : {r['docx_size_kb']}")
        print(f"  Words          : {r['n_words']:,}")
        print(f"  Tipi estratti  : {r['n_tipi_extracted']} (unici: {r['n_tipi_unici']})")
        print(f"  Categorie estr : {r['n_categorie_extracted']}")
        print(f"  Titoli estratti: {r['n_titoli_extracted']}")
        print(f"  Cross-check ZIP: {r['cross_check']['n_found']}/{r['cross_check']['n_zip_files']} file trovati nel docx ({r['cross_check']['n_missing']} mancanti)")
        print(f"  Top 5 categorie:")
        for cat, n in list(r["categoria_distribution"].items())[:5]:
            print(f"    {n:3} × {cat}")
        print(f"  Top 5 tipi:")
        for t, n in list(r["tipo_top10"].items())[:5]:
            print(f"    {n:3} × {t}")
        print()

    # Cross-comparison: chi cattura più documenti?
    print("=== CROSS-COMPARISON CATTURA DOCUMENTI ===")
    for p in providers:
        r = results.get(p, {})
        if r.get("error"):
            continue
        cc = r.get("cross_check", {})
        pct = cc["n_found"] / max(1, cc["n_zip_files"]) * 100
        print(f"  {p:<22}: {cc['n_found']}/{cc['n_zip_files']} ({pct:.1f}%)")
    print()

    # Salva report
    out_path = TEST_ROOT / "quality_detailed.json"
    out_path.write_text(json.dumps(results, indent=2, ensure_ascii=False, default=str), encoding="utf-8")
    print(f"Report salvato: {out_path}")


if __name__ == "__main__":
    main()
