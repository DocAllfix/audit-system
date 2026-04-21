"""
==============================================================================
GENERATE_REPORT.PY - Generazione Report Word (FASE 4)
==============================================================================
Script per la generazione del file Word (.docx) della Relazione di Evidenze.
Riceve il contenuto strutturato dall'Agente e produce il documento finale.
Parte del framework DOE - Execution Layer.

Input:  Testo strutturato (via funzione o file JSON)
Output: /output/Relazione_Evidenze_YYYYMMDD_HHMMSS.docx

Autore: Agente DOE
Data: 2025-12-29
==============================================================================
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path

# Importa configurazione
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from config import PROJECT_ROOT

# Libreria per generazione Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None
    print("[ERROR] python-docx non installato. Impossibile generare Word.")
    sys.exit(1)


# ==============================================================================
# COSTANTI E PERCORSI
# ==============================================================================

OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
TEMP_DIR = os.path.join(PROJECT_ROOT, "temp")


# ==============================================================================
# CONFIGURAZIONE STILE DOCUMENTO
# ==============================================================================

DOCUMENT_CONFIG = {
    "title": "RELAZIONE DI EVIDENZE DI AUDIT",
    "font_name": "Times New Roman",
    "font_size_normal": 12,
    "font_size_title": 16,
    "font_size_heading": 14,
    "line_spacing": 1.5,
    "margin_cm": 2.5
}


# ==============================================================================
# FUNZIONI DI FORMATTAZIONE
# ==============================================================================

def setup_document_styles(doc):
    """
    Configura gli stili base del documento Word.
    
    Args:
        doc: Oggetto Document di python-docx.
    """
    # Stile Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = DOCUMENT_CONFIG["font_name"]
    font.size = Pt(DOCUMENT_CONFIG["font_size_normal"])
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = DOCUMENT_CONFIG["line_spacing"]
    paragraph_format.space_after = Pt(6)
    
    # Margini
    for section in doc.sections:
        section.top_margin = Cm(DOCUMENT_CONFIG["margin_cm"])
        section.bottom_margin = Cm(DOCUMENT_CONFIG["margin_cm"])
        section.left_margin = Cm(DOCUMENT_CONFIG["margin_cm"])
        section.right_margin = Cm(DOCUMENT_CONFIG["margin_cm"])


def add_title(doc, title_text):
    """
    Aggiunge il titolo principale al documento.
    
    Args:
        doc: Oggetto Document.
        title_text: Testo del titolo.
    """
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Formatta il titolo
    for run in title.runs:
        run.font.name = DOCUMENT_CONFIG["font_name"]
        run.font.size = Pt(DOCUMENT_CONFIG["font_size_title"])
        run.font.bold = True


def add_metadata(doc, metadata):
    """
    Aggiunge i metadati del documento (solo data).
    
    Args:
        doc: Oggetto Document.
        metadata: Dizionario con metadati.
    """
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # NOTA: Non inserire numero documenti analizzati (direttiva SOP Sezione 6.7)
    meta_text = f"Data: {metadata.get('data', datetime.now().strftime('%d/%m/%Y'))}"
    
    run = meta_para.add_run(meta_text)
    run.font.name = DOCUMENT_CONFIG["font_name"]
    run.font.size = Pt(DOCUMENT_CONFIG["font_size_normal"])
    run.font.italic = True
    
    # Linea separatrice
    doc.add_paragraph("_" * 60)


def add_statistics_header(doc, stats, data_redazione):
    """
    Aggiunge l'intestazione obbligatoria con data e statistiche documenti.
    
    Args:
        doc: Oggetto Document.
        stats: Dizionario con documenti_estratti, documenti_vuoti, documenti_analizzati.
        data_redazione: Data di redazione del report.
    """
    # Linea separatrice
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("─" * 50)
    sep_run.font.size = Pt(10)
    
    # Box statistiche
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_text = f"""Data di generazione: {data_redazione}
Documenti estratti: {stats.get('documenti_estratti', 'N/D')}
Documenti vuoti: {stats.get('documenti_vuoti', 'N/D')}
Documenti analizzati: {stats.get('documenti_analizzati', 'N/D')}"""
    
    run = header_para.add_run(header_text)
    run.font.name = DOCUMENT_CONFIG["font_name"]
    run.font.size = Pt(11)
    run.font.italic = True
    
    header_para.paragraph_format.space_after = Pt(18)
    
    # Linea separatrice inferiore
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep2_run = sep2.add_run("─" * 50)
    sep2_run.font.size = Pt(10)
    sep2.paragraph_format.space_after = Pt(24)


def add_category_heading(doc, category_name):
    """
    Aggiunge un'intestazione di categoria tematica.
    
    Args:
        doc: Oggetto Document.
        category_name: Nome della categoria.
    """
    heading = doc.add_heading(category_name.upper(), level=1)
    
    for run in heading.runs:
        run.font.name = DOCUMENT_CONFIG["font_name"]
        run.font.size = Pt(DOCUMENT_CONFIG["font_size_heading"])
        run.font.bold = True


def add_paragraph_with_subtitle(doc, numero, sottotitolo, contenuto):
    """
    Aggiunge un paragrafo con sottotitolo numerato e contenuto in prosa.
    
    Args:
        doc: Oggetto Document.
        numero: Numero progressivo del paragrafo.
        sottotitolo: Sottotitolo identificativo del documento.
        contenuto: Testo del paragrafo.
    """
    # Sottotitolo numerato
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    subtitle_run = subtitle_para.add_run(f"[{numero}] {sottotitolo}")
    subtitle_run.font.name = DOCUMENT_CONFIG["font_name"]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.bold = True
    subtitle_run.font.italic = True
    
    subtitle_para.paragraph_format.space_after = Pt(6)
    subtitle_para.paragraph_format.space_before = Pt(12)
    
    # Contenuto
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = para.add_run(contenuto)
    run.font.name = DOCUMENT_CONFIG["font_name"]
    run.font.size = Pt(DOCUMENT_CONFIG["font_size_normal"])
    
    # Spaziatura
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.first_line_indent = Cm(1)


def add_evidence_paragraph(doc, evidence_data):
    """
    Aggiunge un paragrafo di evidenza - supporta sia testo che dizionario.
    
    Args:
        doc: Oggetto Document.
        evidence_data: Testo dell'evidenza o dizionario con numero/sottotitolo/contenuto.
    """
    if isinstance(evidence_data, dict):
        # Nuovo formato con sottotitolo
        numero = evidence_data.get('numero', '')
        sottotitolo = evidence_data.get('sottotitolo', '')
        contenuto = evidence_data.get('contenuto', '')
        
        if sottotitolo and contenuto:
            add_paragraph_with_subtitle(doc, numero, sottotitolo, contenuto)
        elif contenuto:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(contenuto)
            run.font.name = DOCUMENT_CONFIG["font_name"]
            run.font.size = Pt(DOCUMENT_CONFIG["font_size_normal"])
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.first_line_indent = Cm(1)
    else:
        # Formato testo semplice
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = para.add_run(str(evidence_data))
        run.font.name = DOCUMENT_CONFIG["font_name"]
        run.font.size = Pt(DOCUMENT_CONFIG["font_size_normal"])
        
        # Spaziatura
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.first_line_indent = Cm(1)


# ==============================================================================
# FUNZIONE PRINCIPALE DI GENERAZIONE
# ==============================================================================

def generate_report(content, metadata=None):
    """
    Genera il documento Word della Relazione di Evidenze.
    
    Args:
        content: Contenuto strutturato. Può essere:
                 - str: Testo semplice da inserire
                 - dict: Struttura con categorie e evidenze
                 - list: Lista di sezioni/paragrafi
        metadata: Dizionario opzionale con metadati aggiuntivi.
    
    Returns:
        str: Percorso del file Word generato.
    """
    if Document is None:
        raise ImportError("python-docx non disponibile")
    
    # Assicura esistenza directory output
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Crea documento
    doc = Document()
    setup_document_styles(doc)
    
    # Metadata di default
    if metadata is None:
        metadata = {}
    
    metadata.setdefault('data', datetime.now().strftime('%d/%m/%Y'))
    
    # Titolo
    add_title(doc, DOCUMENT_CONFIG["title"])
    add_metadata(doc, metadata)
    
    # Contenuto
    if isinstance(content, str):
        # Testo semplice: dividi per doppi a-capo come sezioni
        sections = content.strip().split("\n\n")
        for section in sections:
            if section.strip():
                # Controlla se è un'intestazione di categoria (tutto maiuscolo)
                if section.strip().isupper() and len(section.strip()) < 100:
                    add_category_heading(doc, section.strip())
                else:
                    add_evidence_paragraph(doc, section.strip())
    
    elif isinstance(content, dict):
        # Struttura con categorie
        for category, evidences in content.items():
            if category.lower() not in ['metadata', 'meta']:
                add_category_heading(doc, category)
                
                if isinstance(evidences, list):
                    for ev in evidences:
                        add_evidence_paragraph(doc, str(ev))
                else:
                    add_evidence_paragraph(doc, str(evidences))
    
    elif isinstance(content, list):
        # Lista di sezioni
        for item in content:
            if isinstance(item, dict):
                # Sezione strutturata
                if 'categoria' in item:
                    add_category_heading(doc, item['categoria'])
                if 'contenuto' in item:
                    add_evidence_paragraph(doc, item['contenuto'])
                elif 'testo' in item:
                    add_evidence_paragraph(doc, item['testo'])
            else:
                add_evidence_paragraph(doc, str(item))
    
    # Genera nome file con timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Relazione_Evidenze_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    
    # Salva documento
    doc.save(filepath)
    
    print(f"[INFO] Documento generato: {filepath}")
    
    return filepath


def generate_from_json(json_path):
    """
    Genera il report da un file JSON strutturato.
    
    Args:
        json_path: Percorso del file JSON con il contenuto.
    
    Returns:
        str: Percorso del file Word generato.
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Gestisce il formato con categorie strutturate
    if 'categorie' in data:
        # Nuovo formato strutturato
        if Document is None:
            raise ImportError("python-docx non disponibile")
        
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc = Document()
        setup_document_styles(doc)
        
        # Titolo personalizzato se presente
        titolo = data.get('titolo', DOCUMENT_CONFIG["title"])
        add_title(doc, titolo)
        
        # Sottotitolo se presente
        if 'sottotitolo' in data:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub.add_run(data['sottotitolo'])
            run.font.name = DOCUMENT_CONFIG["font_name"]
            run.font.size = Pt(14)
            run.font.bold = True
        
        # Statistiche e intestazione obbligatoria
        stats = data.get('statistiche', {})
        data_redazione = data.get('data_redazione', datetime.now().strftime('%d/%m/%Y'))
        
        # Aggiungi intestazione con statistiche se disponibili
        if stats:
            add_statistics_header(doc, stats, data_redazione)
        else:
            # Fallback a metadata semplice
            metadata = {
                'data': data_redazione,
                'documenti_analizzati': len(data.get('categorie', []))
            }
            add_metadata(doc, metadata)
        
        # Categorie con paragrafi
        for categoria in data.get('categorie', []):
            nome_cat = categoria.get('nome', 'CATEGORIA')
            add_category_heading(doc, nome_cat)
            
            for paragrafo in categoria.get('paragrafi', []):
                add_evidence_paragraph(doc, paragrafo)
        
        # Salva con nome azienda nel filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Estrai nome azienda dal sottotitolo (es. "Audit ISO 14001 - ARNO IMPRESE S.R.L.")
        sottotitolo = data.get('sottotitolo', '')
        nome_azienda = ""
        if ' - ' in sottotitolo:
            nome_azienda = sottotitolo.split(' - ')[-1].strip()
            # Pulisci nome per filename valido
            nome_azienda = "".join(c for c in nome_azienda if c.isalnum() or c in (' ', '-', '_')).strip()
            nome_azienda = nome_azienda.replace(' ', '_')[:30]  # Max 30 char
        
        if nome_azienda:
            filename = f"Relazione_Evidenze_{nome_azienda}_{timestamp}.docx"
        else:
            filename = f"Relazione_Evidenze_{timestamp}.docx"
        
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        
        print(f"[INFO] Documento generato: {filepath}")
        return filepath
    else:
        # Formato originale
        content = data.get('content', data.get('evidenze', data))
        metadata = data.get('metadata', {})
        return generate_report(content, metadata)


def generate_from_jsonl(jsonl_path):
    """
    Genera il report da un file JSONL (JSON Lines) - formato incrementale.
    Legge riga per riga senza caricare tutto in memoria.
    
    Args:
        jsonl_path: Percorso del file JSONL con il contenuto.
    
    Returns:
        str: Percorso del file Word generato.
    """
    if Document is None:
        raise ImportError("python-docx non disponibile")
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    setup_document_styles(doc)
    
    # Variabili per costruire il documento
    header_data = None
    current_category = None
    paragrafo_count = 0
    
    print(f"[INFO] Parsing JSONL: {jsonl_path}")
    
    with open(jsonl_path, 'r', encoding='utf-8') as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            
            try:
                obj = json.loads(line)
                tipo = obj.get("type", "unknown")
                
                if tipo == "header":
                    # Prima riga: header con metadata
                    header_data = obj
                    titolo = obj.get('titolo', DOCUMENT_CONFIG["title"])
                    add_title(doc, titolo)
                    
                    # Sottotitolo se presente
                    if obj.get('sottotitolo'):
                        sub = doc.add_paragraph()
                        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = sub.add_run(obj['sottotitolo'])
                        run.font.name = DOCUMENT_CONFIG["font_name"]
                        run.font.size = Pt(14)
                        run.font.bold = True
                    
                    # Statistiche
                    stats = obj.get('statistiche', {})
                    data_redazione = obj.get('data_redazione', datetime.now().strftime('%d/%m/%Y'))
                    if stats:
                        add_statistics_header(doc, stats, data_redazione)
                
                elif tipo == "categoria":
                    # Marcatore di categoria
                    current_category = obj.get('nome', 'CATEGORIA')
                    add_category_heading(doc, current_category)
                
                elif tipo == "paragrafo":
                    # Paragrafo di evidenza
                    add_evidence_paragraph(doc, obj)
                    paragrafo_count += 1
                    
                    if paragrafo_count % 20 == 0:
                        print(f"[PROGRESS] {paragrafo_count} paragrafi processati...")
                
            except json.JSONDecodeError as e:
                print(f"[WARNING] Riga {line_num} JSON invalido: {e}")
                continue
    
    print(f"[INFO] Totale paragrafi: {paragrafo_count}")
    
    # Salva con nome azienda nel filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_azienda = ""
    
    if header_data and header_data.get('sottotitolo'):
        sottotitolo = header_data['sottotitolo']
        if ' - ' in sottotitolo:
            nome_azienda = sottotitolo.split(' - ')[-1].strip()
            nome_azienda = "".join(c for c in nome_azienda if c.isalnum() or c in (' ', '-', '_')).strip()
            nome_azienda = nome_azienda.replace(' ', '_')[:30]
    
    if nome_azienda:
        filename = f"Relazione_Evidenze_{nome_azienda}_{timestamp}.docx"
    else:
        filename = f"Relazione_Evidenze_{timestamp}.docx"
    
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    print(f"[INFO] Documento generato: {filepath}")
    return filepath


def generate_from_text_file(text_path):
    """
    Genera il report da un file di testo semplice.
    
    Args:
        text_path: Percorso del file di testo.
    
    Returns:
        str: Percorso del file Word generato.
    """
    with open(text_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return generate_report(content)


# ==============================================================================
# FUNZIONE DI UTILITÀ PER L'AGENTE
# ==============================================================================

def create_report_from_agent_output(agent_text, doc_count=None):
    """
    Funzione wrapper per uso dall'Agente Gemini.
    Riceve il testo finale redatto dall'Agente e genera il Word.
    
    Args:
        agent_text: Testo completo della relazione redatto dall'Agente.
        doc_count: Numero di documenti analizzati (opzionale).
    
    Returns:
        str: Percorso del file Word generato.
    """
    metadata = {
        'data': datetime.now().strftime('%d/%m/%Y'),
    }
    
    if doc_count:
        metadata['documenti_analizzati'] = doc_count
    
    return generate_report(agent_text, metadata)


# ==============================================================================
# ENTRY POINT
# ==============================================================================

def main():
    """
    Entry point per test o esecuzione standalone.
    Rileva automaticamente il formato di input: JSONL > JSON > TXT
    """
    print("=" * 60)
    print("FASE 4: GENERAZIONE REPORT")
    print("=" * 60)
    
    # Cerca file di input nella temp (ordine di priorità: JSONL > JSON > TXT)
    jsonl_input = os.path.join(TEMP_DIR, "agent_output.jsonl")
    json_input = os.path.join(TEMP_DIR, "agent_output.json")
    text_input = os.path.join(TEMP_DIR, "agent_output.txt")
    
    if os.path.exists(jsonl_input):
        # Formato JSONL (streaming/incrementale) - priorità massima
        print(f"[INFO] Rilevato formato JSONL (incrementale)")
        print(f"[INFO] Generazione da JSONL: {jsonl_input}")
        filepath = generate_from_jsonl(jsonl_input)
    elif os.path.exists(json_input):
        print(f"[INFO] Generazione da JSON: {json_input}")
        filepath = generate_from_json(json_input)
    elif os.path.exists(text_input):
        print(f"[INFO] Generazione da testo: {text_input}")
        filepath = generate_from_text_file(text_input)
    else:
        # Modalità demo/test
        print("[INFO] Nessun input trovato. Generazione documento di test.")
        demo_content = """
DOCUMENTAZIONE LEGALE

Visto il Certificato di Iscrizione alla Camera di Commercio della società Esempio S.r.l. n. REA TO-123456 del 15 gennaio 2024 rilasciato dalla Camera di Commercio di Torino. Il documento attesta che la società risulta iscritta al Registro delle Imprese con capitale sociale di euro 50.000,00 interamente versato, avente come oggetto sociale l'attività di consulenza aziendale e servizi alle imprese. La sede legale risulta ubicata in Via Roma n. 123, Torino (TO), CAP 10100. Il documento presenta firma digitale del Conservatore del Registro.

DOCUMENTAZIONE TECNICA

Visto il Manuale della Qualità versione 3.2 del 10 marzo 2024 redatto dal Responsabile Sistema Gestione Qualità. Il documento si compone di 45 pagine numerate progressivamente e descrive in modo dettagliato la struttura organizzativa aziendale, i processi primari e di supporto, le responsabilità delle figure chiave e le modalità operative per la gestione della qualità. Sono presenti riferimenti a 12 procedure operative e 8 istruzioni di lavoro collegate. Il documento riporta la data di approvazione e la firma della Direzione Generale.
"""
        filepath = generate_report(demo_content, {'documenti_analizzati': 2})
    
    print("\n" + "=" * 60)
    print("FASE 4 COMPLETATA")
    print("=" * 60)
    print(f"[INFO] Output salvato in: {filepath}")
    
    return filepath


if __name__ == "__main__":
    main()
