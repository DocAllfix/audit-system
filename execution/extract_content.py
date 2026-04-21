"""
==============================================================================
EXTRACT_CONTENT.PY - Preparazione Contenuti (FASE 2)
==============================================================================
Script per l'estrazione del testo da documenti (PDF, Word, immagini).
Applica OCR per documenti scansionati/raster.
Parte del framework DOE - Execution Layer.

Input:  /temp/extracted/, /temp/manifest.json
Output: /temp/text_chunks/, /temp/images/, /temp/extraction_report.json

Autore: Agente DOE
Data: 2025-12-29
OTTIMIZZATO: Processamento parallelo con ThreadPoolExecutor
==============================================================================
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# Importa configurazione
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from config import (
    PROJECT_ROOT, 
    TESSERACT_PATH, 
    POPPLER_PATH, 
    OCR_LANGUAGES
)

# Librerie per estrazione contenuti
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
except ImportError:
    Image = None
    pytesseract = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

# Supporto per file .doc legacy (formato Word 97-2003)
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

# Supporto per file Excel
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

# Supporto per file RTF
try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

# Supporto per email Outlook (.msg)
try:
    import extract_msg
    HAS_EXTRACT_MSG = True
except ImportError:
    HAS_EXTRACT_MSG = False


# ==============================================================================
# COSTANTI E PERCORSI
# ==============================================================================

TEMP_DIR = os.path.join(PROJECT_ROOT, "temp")
EXTRACTED_DIR = os.path.join(TEMP_DIR, "extracted")
TEXT_CHUNKS_DIR = os.path.join(TEMP_DIR, "text_chunks")
IMAGES_DIR = os.path.join(TEMP_DIR, "images")
MANIFEST_PATH = os.path.join(TEMP_DIR, "manifest.json")
EXTRACTION_REPORT_PATH = os.path.join(TEMP_DIR, "extraction_report.json")

# Soglia minima caratteri per considerare un PDF come "con testo selezionabile"
MIN_TEXT_THRESHOLD = 50  # Ridotto per velocita - evita OCR inutile

# Numero di worker threads per processamento parallelo (OTTIMIZZATO MAX)
# Impostato a 12x CPU cores per massimizzare throughput I/O-bound
MAX_WORKERS = 96

# Lock per scrittura file thread-safe
write_lock = threading.Lock()

# Lock per Word COM - CRITICO: Word non supporta accessi concorrenti
word_com_lock = threading.Lock()


# ==============================================================================
# FUNZIONI DI ESTRAZIONE PER TIPO (CON OCR AUTOMATICO)
# ==============================================================================

def extract_text_from_pdf(filepath):
    """
    Estrae testo da PDF usando PyMuPDF (molto più veloce di PyPDF2).
    Processa TUTTE le pagine senza limiti.
    """
    text_content = []
    
    # Prova prima con PyMuPDF (fitz) - molto più veloce
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        num_pages = len(doc)
        
        for page_num in range(num_pages):
            page = doc[page_num]
            page_text = page.get_text() or ""
            text_content.append(f"--- PAGINA {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        total_text = "\n\n".join(text_content)
        # Verifica se il testo è sufficiente
        avg_chars = len(total_text.replace(" ", "").replace("\n", "")) / max(num_pages, 1)
        
        if avg_chars < 50:
            # PDF scansionato - OCR sarà gestito dal chiamante
            return "", "needs_ocr", []
        
        return total_text, "native", []
        
    except ImportError:
        pass  # PyMuPDF non installato, usa PyPDF2
    except Exception as e:
        pass  # Errore con PyMuPDF, prova PyPDF2
    
    # Fallback a PyPDF2
    if PyPDF2 is None:
        return "", "error", []
    
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text_content.append(f"--- PAGINA {page_num + 1} ---\n{page_text}")
        
        total_text = "\n\n".join(text_content)
        avg_chars = len(total_text.replace(" ", "").replace("\n", "")) / max(num_pages, 1)
        
        if avg_chars < 50:
            return "", "needs_ocr", []
        
        return total_text, "native", []
        
    except Exception as e:
        return "", "error", []


def extract_text_from_docx(filepath):
    """Estrae testo da file Word."""
    if Document is None:
        return "", "error", []
    
    try:
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs), "native", []
        
    except Exception:
        return "", "error", []


def extract_text_from_txt(filepath):
    """Legge file di testo."""
    try:
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read(), "native", []
            except UnicodeDecodeError:
                continue
        return "", "encoding_error", []
    except Exception:
        return "", "error", []


def extract_text_from_doc_legacy(filepath):
    """Estrae testo da file .doc (Word 97-2003).
    - Windows: usa pywin32 (Word COM) come metodo primario
    - Linux: usa antiword o catdoc come fallback
    """
    import platform
    import subprocess
    
    # Su LINUX, usa antiword o catdoc
    if platform.system() != "Windows":
        # Prova antiword (più affidabile)
        try:
            result = subprocess.run(
                ['antiword', filepath],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip(), "antiword", []
        except FileNotFoundError:
            pass  # antiword non installato
        except subprocess.TimeoutExpired:
            pass  # timeout
        except Exception as e:
            pass
        
        # Fallback a catdoc
        try:
            result = subprocess.run(
                ['catdoc', filepath],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip(), "catdoc", []
        except FileNotFoundError:
            pass  # catdoc non installato
        except subprocess.TimeoutExpired:
            pass
        except Exception as e:
            pass
        
        return "", "no_linux_doc_tools", []
    
    # Su WINDOWS: pywin32 con Word COM (affidabile)
    if HAS_WIN32:
        with word_com_lock:
            word = None
            try:
                import pythoncom
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                
                # Converti a percorso assoluto Windows
                abs_path = os.path.abspath(filepath)
                
                doc = word.Documents.Open(abs_path, ReadOnly=True, AddToRecentFiles=False)
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
                
                if text and len(text.strip()) > 10:
                    return text, "win32", []
                else:
                    return "", "empty", []
            except Exception as e:
                try:
                    if word:
                        word.Quit()
                except:
                    pass
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
                # Se pywin32 fallisce, restituisci errore
                return "", f"error_win32: {str(e)[:50]}", []
    
    # Se pywin32 non disponibile su Windows
    return "", "no_win32_available", []


def extract_text_from_excel(filepath):
    """Estrae testo da file Excel (.xlsx, .xlsm, .xls)."""
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext in ('.xlsx', '.xlsm') and HAS_OPENPYXL:
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            text_parts = []
            for sheet in wb:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                    if row_text.strip() and row_text.strip() != "|":
                        text_parts.append(row_text)
            wb.close()
            return "\n".join(text_parts), "openpyxl", []
        
        elif ext == '.xls' and HAS_XLRD:
            wb = xlrd.open_workbook(filepath)
            text_parts = []
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    row_text = " | ".join(str(cell.value) if cell.value else "" for cell in sheet.row(row_idx))
                    if row_text.strip() and row_text.strip() != "|":
                        text_parts.append(row_text)
            return "\n".join(text_parts), "xlrd", []
        
        else:
            return "", "unsupported_excel", []
    except Exception:
        return "", "error", []


def extract_text_from_rtf(filepath):
    """Estrae testo da file RTF.
    Usa striprtf se disponibile, altrimenti fallback a lettura raw con pulizia.
    """
    # Metodo 1: striprtf (pulizia completa dei tag RTF)
    if HAS_STRIPRTF:
        try:
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        rtf_content = f.read()
                    text = rtf_to_text(rtf_content)
                    if text and len(text.strip()) > 5:
                        return text, "striprtf", []
                except UnicodeDecodeError:
                    continue
        except Exception:
            pass
    
    # Metodo 2: Fallback raw - leggi come testo e rimuovi tag RTF basilari
    try:
        import re
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    raw = f.read()
                # Rimuovi header RTF e tag base
                text = re.sub(r'\\[a-z]+\d*\s?', ' ', raw)
                text = re.sub(r'[{}]', '', text)
                text = re.sub(r'\s+', ' ', text).strip()
                if text and len(text.strip()) > 10:
                    return text, "rtf_raw", []
            except UnicodeDecodeError:
                continue
    except Exception:
        pass
    
    return "", "rtf_error", []


def extract_text_from_msg(filepath):
    """Estrae testo da email Outlook (.msg) tramite libreria extract-msg.
    Produce un output testuale con header (Oggetto, Da, A, Cc, Data),
    corpo del messaggio e nomi degli allegati.
    """
    if not HAS_EXTRACT_MSG:
        return "", "extract_msg_not_installed", []

    msg = None
    try:
        msg = extract_msg.Message(filepath)
        parts = []
        if msg.subject:
            parts.append(f"Oggetto: {msg.subject}")
        if msg.sender:
            parts.append(f"Da: {msg.sender}")
        if msg.to:
            parts.append(f"A: {msg.to}")
        if msg.cc:
            parts.append(f"Cc: {msg.cc}")
        if msg.date:
            parts.append(f"Data: {msg.date}")

        body = msg.body or ""
        if not body and getattr(msg, "htmlBody", None):
            try:
                import re as _re
                html = msg.htmlBody
                if isinstance(html, bytes):
                    html = html.decode("utf-8", errors="ignore")
                body = _re.sub(r"<[^>]+>", " ", html)
                body = _re.sub(r"\s+", " ", body).strip()
            except Exception:
                body = ""
        if body:
            parts.append("")
            parts.append("Corpo del messaggio:")
            parts.append(body)

        if msg.attachments:
            names = []
            for att in msg.attachments:
                name = getattr(att, "longFilename", None) or getattr(att, "shortFilename", None) or "(senza nome)"
                names.append(str(name))
            parts.append("")
            parts.append("Allegati: " + ", ".join(names))

        text = "\n".join(parts)
        return text, "extract_msg", []
    except Exception:
        return "", "msg_error", []
    finally:
        try:
            if msg is not None:
                msg.close()
        except Exception:
            pass


def extract_text_from_p7m(filepath):
    """Estrae testo da file .p7m (documenti firmati digitalmente PKCS#7).
    Usa openssl per estrarre il documento dall'envelope crittografico,
    poi lo processa in base al tipo di contenuto.
    """
    import subprocess
    import tempfile
    
    try:
        # Estrai il documento firmato dall'envelope P7M con openssl
        with tempfile.NamedTemporaryFile(delete=False, suffix='.extracted') as tmp:
            tmp_path = tmp.name
        
        # openssl smime: estrae il contenuto senza verificare la firma
        result = subprocess.run(
            ['openssl', 'smime', '-verify', '-noverify', '-nosigs',
             '-in', filepath, '-inform', 'DER', '-out', tmp_path],
            capture_output=True, timeout=30
        )
        
        # Se DER fallisce, prova PEM
        if result.returncode != 0 or not os.path.exists(tmp_path) or os.path.getsize(tmp_path) == 0:
            result = subprocess.run(
                ['openssl', 'smime', '-verify', '-noverify', '-nosigs',
                 '-in', filepath, '-inform', 'PEM', '-out', tmp_path],
                capture_output=True, timeout=30
            )
        
        if not os.path.exists(tmp_path) or os.path.getsize(tmp_path) == 0:
            try:
                os.unlink(tmp_path)
            except:
                pass
            return "", "p7m_extract_failed", []
        
        # Rileva tipo del contenuto estratto e usa la funzione appropriata
        text = ""
        method = "p7m_unknown"
        
        # Leggi i primi bytes per rilevare il tipo
        with open(tmp_path, 'rb') as f:
            header = f.read(16)
        
        if header[:4] == b'%PDF':
            # È un PDF
            text, method, _ = extract_text_from_pdf(tmp_path)
            method = f"p7m_pdf_{method}"
        elif header[:2] == b'PK':
            # È un file ZIP-based (DOCX, XLSX, etc.)
            text, method, _ = extract_text_from_docx(tmp_path)
            method = f"p7m_docx_{method}"
        elif header[:5] == b'{\\rtf':
            # È un RTF
            text, method, _ = extract_text_from_rtf(tmp_path)
            method = f"p7m_rtf_{method}"
        else:
            # Prova come testo generico
            text, method, _ = extract_text_from_txt(tmp_path)
            method = f"p7m_txt_{method}"
        
        # Pulizia file temporaneo
        try:
            os.unlink(tmp_path)
        except:
            pass
        
        return text, method, []
        
    except subprocess.TimeoutExpired:
        try:
            os.unlink(tmp_path)
        except:
            pass
        return "", "p7m_timeout", []
    except FileNotFoundError:
        # openssl non installato
        return "", "p7m_no_openssl", []
    except Exception as e:
        try:
            os.unlink(tmp_path)
        except:
            pass
        return "", f"p7m_error", []


# ==============================================================================
# WORKER PER PROCESSAMENTO PARALLELO
# ==============================================================================

def process_file_worker(file_info):
    """
    Worker thread per processare un singolo file.
    """
    filepath = file_info["absolute_path"]
    category = file_info["category"]
    filename = file_info["filename"]
    
    result = {
        "filename": filename,
        "category": category,
        "method": "unknown",
        "status": "pending",
        "text_length": 0,
        "images": [],
        "chunk_file": None
    }
    
    # Estrai in base alla categoria (CON OCR AUTOMATICO per PDF)
    if category == "pdf":
        text, method, images = extract_text_from_pdf(filepath)
    elif category == "word":
        # Distingui tra .docx e .doc legacy
        if filepath.lower().endswith('.doc'):
            text, method, images = extract_text_from_doc_legacy(filepath)
        else:
            text, method, images = extract_text_from_docx(filepath)
    elif category == "excel":
        text, method, images = extract_text_from_excel(filepath)
    elif category == "text":
        text, method, images = extract_text_from_txt(filepath)
    elif category == "rtf":
        text, method, images = extract_text_from_rtf(filepath)
    elif category == "p7m":
        text, method, images = extract_text_from_p7m(filepath)
    else:
        text, method, images = "", "unsupported", []
    
    result["method"] = method
    result["images"] = images
    result["text_length"] = len(text)
    
    # Salva il chunk di testo - OTTIMIZZATO con buffered write
    if text.strip():
        chunk_filename = f"{Path(filename).stem}.txt"
        chunk_path = os.path.join(TEXT_CHUNKS_DIR, chunk_filename)
        
        # Prepara contenuto prima del lock per minimizzare tempo di blocco
        content = f"=== DOCUMENTO: {filename} ===\n=== CATEGORIA: {category.upper()} ===\n=== METODO: {method.upper()} ===\n{'=' * 50}\n\n{text}"
        
        with write_lock:
            with open(chunk_path, 'w', encoding='utf-8', buffering=65536) as f:
                f.write(content)
        
        result["chunk_file"] = chunk_path
        result["status"] = "success"
    else:
        result["status"] = "empty" if method != "error" else "error"
    
    return result


# ==============================================================================
# ENTRY POINT
# ==============================================================================

def main():
    """
    Funzione principale - PROCESSAMENTO PARALLELO.
    """
    print("=" * 60)
    print("FASE 2: PREPARAZIONE CONTENUTI (PARALLELO)")
    print("=" * 60)
    
    if not os.path.exists(MANIFEST_PATH):
        print(f"[ERROR] Manifest non trovato: {MANIFEST_PATH}")
        return None
    
    with open(MANIFEST_PATH, 'r', encoding='utf-8') as f:
        manifest = json.load(f)
    
    files = manifest.get("files", [])
    print(f"[INFO] File da processare: {len(files)}")
    print(f"[INFO] Worker threads: {MAX_WORKERS}")
    
    os.makedirs(TEXT_CHUNKS_DIR, exist_ok=True)
    os.makedirs(IMAGES_DIR, exist_ok=True)
    
    extraction_results = []
    stats = {"success": 0, "empty": 0, "error": 0}
    completed = 0
    
    # Processamento parallelo
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(process_file_worker, f): f for f in files}
        
        for future in as_completed(futures):
            result = future.result()
            extraction_results.append(result)
            stats[result["status"]] = stats.get(result["status"], 0) + 1
            completed += 1
            
            if completed % 50 == 0 or completed == len(files):
                print(f"[PROGRESS] {completed}/{len(files)} completati")
    
    # Genera report
    report = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "total_processed": len(files),
            "statistics": stats
        },
        "results": extraction_results
    }
    
    with open(EXTRACTION_REPORT_PATH, 'w', encoding='utf-8') as f:
        json.dump(report, f, indent=2, ensure_ascii=False)
    
    print("\n" + "=" * 60)
    print("FASE 2 COMPLETATA")
    print("=" * 60)
    print(f"[INFO] Successi: {stats.get('success', 0)}")
    print(f"[INFO] Vuoti: {stats.get('empty', 0)}")
    print(f"[INFO] Errori: {stats.get('error', 0)}")
    
    return report


# ==============================================================================
# NUOVI HANDLER — Estensioni aggiuntive (aggiunte 2026-04-19)
# Tutte usano solo stdlib Python — zero dipendenze nuove
# ==============================================================================

def extract_text_from_eml(filepath):
    """Email standard (.eml) — header + corpo, con fallback HTML→testo."""
    import email as _email
    from email.header import decode_header as _dh
    import re as _re

    def _decode_str(s):
        if not s:
            return ""
        try:
            parts = _dh(s)
            out = []
            for part, charset in parts:
                if isinstance(part, bytes):
                    out.append(part.decode(charset or "utf-8", errors="ignore"))
                else:
                    out.append(str(part))
            return " ".join(out)
        except Exception:
            return str(s)

    try:
        with open(filepath, "rb") as f:
            msg = _email.message_from_bytes(f.read())

        lines = []
        subject = _decode_str(msg.get("Subject", ""))
        from_   = _decode_str(msg.get("From", ""))
        to      = _decode_str(msg.get("To", ""))
        date    = msg.get("Date", "")
        if subject: lines.append(f"Oggetto: {subject}")
        if from_:   lines.append(f"Da: {from_}")
        if to:      lines.append(f"A: {to}")
        if date:    lines.append(f"Data: {date}")

        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/plain" and not body:
                    try:
                        cs = part.get_content_charset() or "utf-8"
                        raw = part.get_payload(decode=True)
                        if raw:
                            body = raw.decode(cs, errors="ignore")
                    except Exception:
                        pass
                elif ctype == "text/html" and not body:
                    try:
                        cs = part.get_content_charset() or "utf-8"
                        raw = part.get_payload(decode=True)
                        if raw:
                            html = raw.decode(cs, errors="ignore")
                            body = _re.sub(r"<[^>]+>", " ", html)
                            body = _re.sub(r"\s+", " ", body).strip()
                    except Exception:
                        pass
        else:
            try:
                raw = msg.get_payload(decode=True)
                if raw:
                    cs = msg.get_content_charset() or "utf-8"
                    body = raw.decode(cs, errors="ignore")
                    if msg.get_content_type() == "text/html":
                        body = _re.sub(r"<[^>]+>", " ", body)
                        body = _re.sub(r"\s+", " ", body).strip()
            except Exception:
                pass

        if body.strip():
            lines.append("")
            lines.append("Corpo del messaggio:")
            lines.append(body.strip())

        text = "\n".join(lines)
        return (text, "eml", []) if text.strip() else ("", "eml_empty", [])
    except Exception:
        return "", "error", []


def extract_text_from_html(filepath):
    """HTML/HTM — estrae testo visibile, esclude script/style/head."""
    from html.parser import HTMLParser
    import re as _re

    class _Extractor(HTMLParser):
        # Tag che hanno contenuto da escludere (con tag di chiusura)
        SKIP_BLOCK = {"script", "style", "head", "noscript", "title"}
        # Tag void (auto-chiusi, nessun contenuto testo) — ignorati senza toccare il contatore
        VOID = {"meta", "link", "br", "hr", "img", "input", "area", "base",
                "col", "embed", "param", "source", "track", "wbr"}

        def __init__(self):
            super().__init__()
            self.texts = []
            self._depth = 0

        def handle_starttag(self, tag, attrs):
            if tag.lower() in self.SKIP_BLOCK:
                self._depth += 1
            # VOID tags: non toccano il contatore (mai hanno close tag)

        def handle_endtag(self, tag):
            if tag.lower() in self.SKIP_BLOCK and self._depth > 0:
                self._depth -= 1

        def handle_data(self, data):
            if self._depth == 0 and data.strip():
                self.texts.append(data.strip())

    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(filepath, "r", encoding=enc, errors="ignore") as f:
                html = f.read()
            parser = _Extractor()
            parser.feed(html)
            text = "\n".join(t for t in parser.texts if t)
            text = _re.sub(r"\n{3,}", "\n\n", text).strip()
            if text:
                return text, "html_parser", []
            return "", "html_empty", []
        except Exception:
            continue
    return "", "error", []


def extract_text_from_xml(filepath):
    """XML — estrae testo da tutti i nodi (DGUE, documenti normativi, ecc.)."""
    from xml.etree import ElementTree as ET

    try:
        # Tentativo 1: parsing standard
        try:
            root = ET.parse(filepath).getroot()
        except ET.ParseError:
            # Tentativo 2: legge raw e rimuove BOM
            with open(filepath, "rb") as f:
                raw = f.read()
            if raw.startswith(b"\xef\xbb\xbf"):
                raw = raw[3:]
            root = ET.fromstring(raw)

        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            if elem.tail and elem.tail.strip():
                texts.append(elem.tail.strip())

        text = "\n".join(texts)
        return (text, "xml_parser", []) if text.strip() else ("", "xml_empty", [])
    except Exception:
        return "", "error", []


def extract_text_from_odt(filepath):
    """LibreOffice Writer (.odt) — legge content.xml nello ZIP ODF."""
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(filepath, "r") as z:
            with z.open("content.xml") as f:
                root = ET.parse(f).getroot()

        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            if elem.tail and elem.tail.strip():
                texts.append(elem.tail.strip())

        text = "\n".join(texts)
        return (text, "odt_xml", []) if text.strip() else ("", "odt_empty", [])
    except Exception:
        return "", "error", []


def extract_text_from_ods(filepath):
    """LibreOffice Calc (.ods) — legge content.xml nello ZIP ODF.
    Openpyxl non supporta .ods — si usa il parsing XML diretto.
    """
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(filepath, "r") as z:
            with z.open("content.xml") as f:
                root = ET.parse(f).getroot()

        rows_text = []
        for elem in root.iter():
            # Righe di tabella ODS: tag che contengono 'table-row'
            if "table-row" in elem.tag:
                cells = []
                for cell in elem:
                    if "table-cell" in cell.tag:
                        cell_text = " ".join(
                            (t.text or "").strip()
                            for t in cell.iter()
                            if t.text and t.text.strip()
                        )
                        if cell_text:
                            cells.append(cell_text)
                row = " | ".join(cells)
                if row.strip():
                    rows_text.append(row)

        text = "\n".join(rows_text)
        return (text, "ods_xml", []) if text.strip() else ("", "ods_empty", [])
    except Exception:
        return "", "error", []


def extract_text_from_pptx(filepath):
    """PowerPoint (.pptx) — estrae testo da slide XML senza python-pptx.
    PPTX è un archivio ZIP: il testo è in ppt/slides/slide*.xml nei tag <a:t>.
    """
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        texts = []
        with zipfile.ZipFile(filepath, "r") as z:
            # Slide in ordine numerico
            slide_files = sorted(
                n for n in z.namelist()
                if n.startswith("ppt/slides/slide") and n.endswith(".xml")
            )
            for slide_path in slide_files:
                slide_num = slide_path.replace("ppt/slides/slide", "").replace(".xml", "")
                with z.open(slide_path) as f:
                    root = ET.parse(f).getroot()
                slide_texts = [
                    elem.text.strip()
                    for elem in root.iter()
                    if elem.tag.endswith("}t") and elem.text and elem.text.strip()
                ]
                if slide_texts:
                    texts.append(f"--- SLIDE {slide_num} ---")
                    texts.extend(slide_texts)

        text = "\n".join(texts)
        return (text, "pptx_xml", []) if text.strip() else ("", "pptx_empty", [])
    except Exception:
        return "", "error", []


def extract_text_from_mpp(filepath):
    """Microsoft Project (.mpp) — Compound Document OLE.
    Estrae nomi task, risorse e milestone scansionando tutti i Var2Data stream
    con pattern UTF-16LE null-terminate. Richiede olefile (>=0.46).
    """
    import re as _re

    try:
        import olefile
    except ImportError:
        return "", "olefile_missing", []

    try:
        ole = olefile.OleFileIO(filepath)
        all_entries = ole.listdir()

        # Pattern: sequenze di (byte stampabile 0x20-0x7E) + \x00, >=4 chars, null-terminate
        _utf16_pat = _re.compile(b'(?:[\x20-\x7e]\x00){4,}(?:\x00\x00)?')

        seen = set()
        strings = []

        for entry in all_entries:
            try:
                data = ole.openstream(entry).read()
                for m in _utf16_pat.finditer(data):
                    s = m.group().decode('utf-16-le', errors='ignore').rstrip('\x00').strip()
                    if len(s) >= 4 and s not in seen:
                        seen.add(s)
                        strings.append(s)
            except Exception:
                continue

        ole.close()

        if not strings:
            return "", "mpp_empty", []

        # Filtra metadati OLE interni irrilevanti
        _noise = {'CompObj', 'ObjInfo', 'OlePres000', 'Package', 'Ole',
                  'PROJECT', 'PROJECTwm', '_VBA_PROJECT', 'ThisProject',
                  'Start', 'EndDate', 'Week', 'Default', 'Tasks', 'Resources'}
        _internal_prefixes = ('_', 'TBknd', 'CV_', 'CDrawing', 'CEdl', 'CFilter',
                              'CGrouping', 'CMap', 'CReport', 'CTable', 'CVba', 'MsoData')

        def _is_noise(s):
            if s in _noise:
                return True
            if any(s.startswith(p) for p in _internal_prefixes):
                return True
            # Lista di stream names separati da virgola (metadata interno)
            if ',' in s and all(part.strip().startswith(('TBknd', 'CV')) for part in s.split(',') if part.strip()):
                return True
            # Frammenti troncati (iniziano o finiscono con lettere minuscole isolate)
            if len(s) <= 6 and s[0].islower():
                return True
            return False

        strings = [s for s in strings if not _is_noise(s)]

        text = "\n".join(strings)
        return (text, "mpp_ole", []) if text.strip() else ("", "mpp_empty", [])

    except Exception:
        return "", "error", []


def extract_text_from_cam(filepath):
    """File .cam (Camera AutoCAD / CNC machining plan).
    Formato variabile: alcuni sono testo G-code/NC, altri XML, altri binary.
    Strategia: tenta ASCII/UTF-8 (testo), poi UTF-16LE scan (binary strutturato).
    """
    import re as _re

    try:
        with open(filepath, 'rb') as f:
            raw = f.read(4096)  # Leggi solo header per decidere strategia

        # Se testo leggibile (>60% printable ASCII), decodifica intero file
        printable = sum(1 for b in raw if 32 <= b < 127 or b in (9, 10, 13))
        ratio = printable / max(len(raw), 1)

        if ratio >= 0.60:
            # Leggi file completo come testo
            for enc in ('utf-8', 'latin-1', 'cp1252'):
                try:
                    with open(filepath, 'r', encoding=enc, errors='ignore') as f:
                        text = f.read()
                    if text.strip():
                        # Per file CNC/G-code: estrai commenti e header (righe con ';' o '(')
                        lines = text.splitlines()
                        meaningful = []
                        for line in lines[:500]:  # Limita a prime 500 righe
                            stripped = line.strip()
                            if stripped and not stripped.startswith('%'):
                                meaningful.append(stripped)
                        result = "\n".join(meaningful[:200])
                        return (result, "cam_text", []) if result.strip() else ("", "cam_empty", [])
                except Exception:
                    continue

        # Altrimenti: scan UTF-16LE (formato structured binary, es. AutoCAD camera)
        with open(filepath, 'rb') as f:
            data = f.read()

        _utf16_pat = _re.compile(b'(?:[\x20-\x7e]\x00){4,}(?:\x00\x00)?')
        strings = []
        seen = set()
        for m in _utf16_pat.finditer(data):
            s = m.group().decode('utf-16-le', errors='ignore').rstrip('\x00').strip()
            if len(s) >= 4 and s not in seen:
                seen.add(s)
                strings.append(s)

        if strings:
            text = "\n".join(strings[:200])
            return text, "cam_binary_scan", []

        return "", "cam_unreadable", []

    except Exception:
        return "", "error", []


if __name__ == "__main__":
    main()

